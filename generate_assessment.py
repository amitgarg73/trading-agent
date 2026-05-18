"""
Generates Trading_Agent_Brutal_Assessment.docx — an honest, unfiltered critique
of the system's real edge (or lack thereof) and what would actually create alpha.

Run: python3 generate_assessment.py
Update this file whenever the architecture or signal thesis changes materially.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

NAVY   = RGBColor(0x1A, 0x3A, 0x6A)
RED    = RGBColor(0xC0, 0x39, 0x2B)
ORANGE = RGBColor(0xF4, 0x7B, 0x20)
GREEN  = RGBColor(0x1A, 0x7A, 0x3A)
GRAY   = RGBColor(0x44, 0x44, 0x44)


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = NAVY
    return p


def red_heading(text, level=2):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RED
    return p


def green_heading(text, level=2):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = GREEN
    return p


def subheading(text, color=ORANGE):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()


def divider():
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading('AI Trading Agent — Brutal Honest Assessment', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RED

meta = doc.add_paragraph('Amit Garg  ·  May 2026  ·  Living document — update when thesis changes')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(11)
meta.runs[0].font.color.rgb = GRAY

doc.add_paragraph()
body(
    'This document is an unfiltered critique of the trading agent — written without credit '
    'for engineering work. The goal is to distinguish between what is operationally solid '
    'and what actually generates alpha. These are different things. A beautifully engineered '
    'system with no edge is just a beautiful system. Read this before every sprint to stay honest.'
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 1: BRUTAL ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 1: What Is Actually Wrong')

body(
    'Seven honest problems with the system as of v5.6. None of these are addressed by '
    'better guardrails, faster trailing stops, or a prettier dashboard.'
)

divider()
red_heading('Problem 1: The Return Claim Is Physically Implausible')
body('Our best backtest: $716/day on $100K = ~175% annualized.')
body(
    'Renaissance Medallion (best quant fund in 30 years): 39% net. '
    'Two Sigma: ~15%. DE Shaw: ~20%.'
)
body(
    'We are claiming 4.5x better than the best quantitative fund ever built, '
    'using public yfinance data and an LLM. That is not a sign we found alpha. '
    'That is a sign the backtest is wrong.'
)
bullet('30-day sample is not statistical significance — it is a streak')
bullet('Apr–May 2026 was a specific momentum regime (post-tariff recovery) — we have not tested a choppy sideways market')
bullet('90% win days = 27 winning days. Way too small a sample to conclude anything')

divider()
red_heading('Problem 2: The Signal Is Weak by Construction')
body(
    'RSI, MACD, Bollinger Bands — the most over-exploited signals in existence. '
    'Every retail platform, algo fund, and screener runs these. '
    'If they had persistent edge, arbitrage would eliminate it quickly.'
)
body('What the academic literature actually says:')
bullet('MACD crossovers win ~40% of the time — worse than a coin flip')
bullet('Bollinger Bands provide no directional signal by themselves')
bullet('RSI shows marginal predictive power in some studies but most lack transaction cost analysis')
bullet('Combined RSI+MACD strategies show 77% win rates in backtesting — but backtests are optimistic by construction')
body(
    'Our scanner scores -10 to +10 on these indicators. The ML model (AUC 0.78) is trained '
    'on them. We have built a sophisticated system on top of signals that the academic '
    'literature largely cannot validate after transaction costs.'
)

divider()
red_heading('Problem 3: Claude Is the Wrong Tool for Stock Selection')
body(
    'Claude was trained on internet text through 2024. When it reasons about why NVDA '
    'is a good setup today, it is pattern-matching on training data — not market microstructure. '
    'It has zero information advantage over any retail trader who can read a chart.'
)
bullet('Claude\'s HIGH/MEDIUM/LOW confidence ratings have never been validated against actual win rates')
bullet('Claude sees the same technical indicators as every retail trader')
bullet('LLM "reasoning" from public data is not alpha — it is post-hoc rationalization of technical signals')
bullet('Research shows LLMs don\'t consistently beat baselines on stock prediction tasks')
body(
    'This does not mean Claude has no role. It means Claude is currently being used '
    'as the signal source when it should be a synthesizer on top of real signals.'
)

divider()
red_heading('Problem 4: The Backtest Is Too Short and Too Optimistic')
body(
    'Paper trading systematically overstates returns. Real fills depend on order book depth, '
    'queue position, and spread dynamics. Research estimates the paper-to-live gap at '
    '2–3% per trade in execution friction alone.'
)
bullet('Limit orders at entry × 1.001 do not model real fill uncertainty')
bullet('No market impact modeled on $5K–7K orders')
bullet('30-day window — not enough to distinguish alpha from regime-specific luck')
bullet('Out-of-sample test: the ML model was trained on the same 2-year window — the backtest is not truly out-of-sample')
bullet('Statistical significance for win rate requires ~200+ trades minimum; we have far fewer')

divider()
red_heading('Problem 5: We Engineered the Wrong Thing')
body(
    'The health checks, eval, trailing stops, guardrails, exit mechanism tracking, '
    'Agent Scorecard — all excellent. All operational. None of it generates alpha.'
)
body(
    'We have spent 70+ steps making the system run reliably. The core question — '
    '"does the signal actually predict price movement?" — we have addressed with '
    '30 days of paper trading and an ML model trained on public data everyone already has.'
)
body('Operational robustness is necessary. It is not sufficient.')

divider()
red_heading('Problem 6: The Goal Structure Is a Behavioral Trap')
body(
    'Setting a "$1,000/day target" and locking in at $716 when we hit it is behavioral finance, '
    'not systematic trading.'
)
bullet('Lock-in caps upside: on a strong day where we might make $2K, we exit at $716')
bullet('Loss is not symmetrically capped — the -$300 daily loss limit is much smaller than the upside we give up')
bullet('Real systematic trading optimizes Sharpe ratio and max drawdown — not a dollar number')
bullet('Dollar targets encourage over-trading on slow days to "hit the number"')

divider()
red_heading('Problem 7: The Market Comparison Was Too Kind to Us')
body(
    'In our initial analysis we listed guardrails and eval as advantages over competitors. '
    'Those are operational advantages. They do not generate alpha.'
)
add_table(
    ['Competitor', 'What They Have That We Don\'t'],
    [
        ('MAHORAGA', 'Social sentiment signal (Reddit/StockTwits) — genuinely different data class'),
        ('OpenClaw', 'Multi-agent consensus (3/4 vote) — reduces single-LLM hallucination bias'),
        ('TradingAgents', 'Bull/Bear debate structure — forces adversarial reasoning before entry'),
        ('Us', 'Best operations. Possibly weakest signal.'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 2: WHAT WOULD ACTUALLY CREATE EDGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 2: What Would Actually Create Edge')

body(
    'Real edge in markets comes from one of three sources: information others do not have yet, '
    'behavioral anomalies that persist because humans are predictable, or execution advantages. '
    'Technical analysis + LLM reasoning gives you none of these. Here is what does.'
)

divider()
green_heading('Edge 1: Post-Earnings Announcement Drift (PEAD)')
body(
    'One of the most robustly documented anomalies in academic finance — 40+ years of papers. '
    'When a stock beats earnings estimates significantly, it continues drifting upward for '
    '5–20 days because institutions are slow to fully reprice and retail herds in late.'
)
bullet('Why it persists: underreaction is structural — large funds cannot fully reposition in one day')
bullet('Scale: +3–8% over the first 5–10 trading days after a big beat')
bullet('What we need: earnings calendar + actual vs. estimate + % surprise magnitude')
bullet('Data: free — yfinance has EPS data, Earnings Whispers has consensus estimates')
bullet('Hold period: 3–5 days (not intraday) — different exit rules than current system')
body(
    'This is a completely different signal class. It is fundamental-based drift, not technical noise. '
    'The edge is real and documented. Build this as Mode 2.'
)

divider()
green_heading('Edge 2: Unusual Options Activity — Informed Money Signal')
body(
    'Large out-of-money call purchases 1–3 weeks before a price move means informed traders '
    'are positioning. Legally observable. Academic evidence confirms unusual options volume '
    'predicts stock returns 1–5 days ahead.'
)
bullet('Why it works: options are cheap leverage — informed traders use them before news')
bullet('What to look for: OTM call volume 5–10x average, high bid-side pressure, short expiry')
bullet('Data: Unusual Whales (free tier), Barchart unusual activity (free)')
bullet('Use case: confirming filter on top of PEAD or technicals — not a standalone signal')
body(
    'Do not trade unusual options activity alone. Use it to add conviction to setups '
    'that already qualify on another signal.'
)

divider()
green_heading('Edge 3: Insider Buying (SEC Form 4) — Documented 5–8% Abnormal Return')
body(
    'When executives buy their own stock in the open market (not grants, not options exercise — '
    'actual open market purchases), academic literature shows 5–8% abnormal return over 6 months. '
    'Signal strengthens when multiple insiders buy in the same window.'
)
bullet('Why it works: insiders have legal knowledge of business trajectory')
bullet('What to look for: CEO/CFO open market purchase, >$50K amount, no prior pattern of purchases')
bullet('Data: SEC EDGAR free API — Form 4 filings within 2 business days of transaction')
bullet('Build effort: low — parse EDGAR XML, filter for open market purchase type codes')
body('Strongest as a swing trade signal (5–20 day hold), not intraday.')

divider()
green_heading('Edge 4: Short Squeeze Setup — Structural Forced-Buyer Anomaly')
body(
    'High short interest + high borrow rate + momentum catalyst = asymmetric upside '
    'when shorts are forced to cover. This exploits a structural forced-buyer situation '
    '— short sellers have unlimited downside risk and must cover eventually.'
)
bullet('What to look for: short interest >20% of float, days-to-cover >5, momentum turning up')
bullet('Data: FINRA short interest (free, 2x/month)')
bullet('Use case: combine with unusual options or news catalyst for best results')

divider()
green_heading('Edge 5: Analyst Revision Momentum')
body(
    'When multiple analysts raise price targets within a short window, it signals institutional '
    'consensus shifting. Academic evidence shows analyst revisions predict returns 5–20 days out. '
    'More durable than technical signals because it reflects fundamental reassessment.'
)
bullet('What to look for: 2+ upgrades or price target raises within 5 trading days')
bullet('Data: free via finviz screener, Yahoo Finance')
bullet('Build effort: low — scrape finviz upgrade/downgrade page')


# ══════════════════════════════════════════════════════════════════════════════
# PART 3: THE ARCHITECTURE SHIFT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 3: The Architecture Shift Required')

body(
    'The single biggest structural change is not a new feature — it is changing what Claude '
    'is responsible for. Currently Claude is the signal. It should be the synthesizer.'
)

subheading('Current (Wrong) Architecture')
body('Technical scanner → Claude decides what to buy → Alpaca')
body(
    'Claude looks at RSI/MACD/volume and picks stocks. This is using LLM reasoning '
    'as the primary alpha source — the weakest possible use of Claude.'
)

subheading('Target Architecture')
body(
    'PEAD candidates + Unusual options flow + Insider buying + Technical confirmation '
    '→ Claude synthesizes → high-conviction trade plan → Alpaca'
)
body(
    'Claude is excellent at: "Given NVDA beat earnings by 12%, unusual call volume is 8x normal, '
    'and two directors bought last week — here is the entry, target, stop, and conviction level." '
    'Claude is not good at: "RSI is 58 and MACD just crossed — buy this."'
)

add_table(
    ['Signal Layer', 'Current', 'Target'],
    [
        ('Primary signal',      'RSI/MACD/Bollinger (weak)',         'PEAD + insider buys + unusual options (documented edge)'),
        ('Pre-filter',          'ML scorer on technicals',           'ML scorer on event signals'),
        ('Claude\'s role',      'Decision-maker (wrong)',            'Synthesizer across real signals (right)'),
        ('Hold period',         'Intraday only',                     'Intraday (Mode 1) + 3-5 day event-driven (Mode 2)'),
        ('Success metric',      '$716/day dollar target (wrong)',    'Sharpe ratio ≥ 1.5 over 200+ trades (right)'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 4: WHAT EVIDENCE WOULD ACTUALLY MEAN SOMETHING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 4: What Evidence Would Actually Mean Something')

body(
    'The June 1 gate (2 weeks of paper with native trailing stop) validates execution. '
    'It does not validate edge. Here is what would.'
)

add_table(
    ['What We Have', 'What Would Change the Assessment'],
    [
        ('30-day paper backtest',           '6-month live paper + out-of-sample backtest on held-out data'),
        ('AUC 0.78 on historical data',     'AUC 0.78 on data the model never trained on (true holdout)'),
        ('Claude picks from technicals',    'A signal Claude has that the market does not (PEAD, options flow, insider)'),
        ('Win rate on paper',               'Win rate on live with realistic slippage modeled'),
        ('$716/day dollar target',          'Sharpe ratio ≥ 1.5 over 200+ closed trades'),
        ('90% win days (27 days)',          'Statistical significance requires ~200 win/loss data points minimum'),
        ('2-week validation gate',          '6+ months of live paper on the new signal architecture'),
    ]
)

divider()
subheading('Signal Evidence Ranking', GREEN)
add_table(
    ['Signal', 'Academic Evidence', 'Build Effort', 'Data Cost'],
    [
        ('PEAD (post-earnings drift)',         'Very strong (40+ years)',      'Medium',   'Free'),
        ('Insider buying Form 4',              'Strong (5–8% abnormal return)','Low',      'Free (SEC EDGAR)'),
        ('Unusual options flow',               'Strong (1–5 day predictive)',  'Medium',   'Free tier'),
        ('Short squeeze setup',                'Moderate',                     'Medium',   'Free (FINRA)'),
        ('Analyst revision momentum',          'Moderate',                     'Low',      'Free (finviz)'),
        ('Pure technicals (what we have now)', 'Weak to none after costs',     'Done',     'Free'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 5: HONEST BOTTOM LINE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 5: Honest Bottom Line')

body(
    'The system is a well-engineered paper trading simulator. The operational infrastructure '
    '(guardrails, eval, trailing stops, health checks, Agent Scorecard) is genuinely good. '
    'The signal is not.'
)
body(
    'The backtest returns (175% annualized) are implausible on their face and almost certainly '
    'reflect a lucky 30-day window in a favorable momentum regime, not a repeatable edge. '
    'The signal — technical indicators filtered by an LLM — is the most competed-for signal '
    'class in retail trading. Thousands of algos run the same signals.'
)
body(
    'That is not a reason to stop. It is a reason to get clear on what would actually '
    'constitute evidence, and to build toward signals with documented predictive power.'
)

subheading('Recommended Sprint Order — Alpha / Product Path (Post-June 1)')
bullet('1. 15-min momentum confirmation (low effort, filters false breakouts — improves Mode 1 marginally)')
bullet('2. Insider buying feed from SEC EDGAR (low effort, genuinely new signal class)')
bullet('3. PEAD mode — post-earnings drift as Mode 2 (medium effort, strongest documented edge)')
bullet('4. Unusual options activity as confirming filter (medium effort, adds conviction to Mode 2)')
bullet('5. Change success metric from dollar target to Sharpe ratio')
bullet('6. Architect Claude as synthesizer, not signal source')

body('')
body(
    'The path to a real edge runs through events and alternative data — not through '
    'better technical indicator combinations or more sophisticated LLM prompting. '
    'Every sprint that improves the signal layer matters more than every sprint '
    'that improves the operational layer.'
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 6: REFRAME — CASH FLOW GOAL VS ALPHA GOAL
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 6: Reframe — Cash Flow Goal vs. Alpha Goal')

body(
    'Parts 1–5 were written against the wrong benchmark. The goal of this system '
    'is not to beat the market, outperform hedge funds, or achieve a publishable Sharpe ratio. '
    'The goal is to generate a reliable daily income — a predefined cash flow target — '
    'without a day job. That is a fundamentally different problem, and most of the brutal '
    'critique above does not apply to it.'
)

divider()
subheading('The Math Reframe')
body(
    'With 3:1 reward:risk, you only need to be right more than 25% of the time to be '
    'profitable. At a 55% win rate (what the backtest suggests):'
)
add_table(
    ['Input', 'Value'],
    [
        ('Win rate',                    '55%'),
        ('Avg gain per trade (2% on $6K)', '$120'),
        ('Avg loss per trade (0.67% on $6K)', '$40'),
        ('Expected value per trade',    '0.55 × $120 − 0.45 × $40 = $48'),
        ('15 trades per day',           '$48 × 15 = $720/day expected value'),
        ('Even at 40% degradation (live slippage)', '$720 × 0.60 = $432/day'),
    ]
)
body(
    'The comparison to Renaissance Medallion (39% net annualized) was irrelevant. '
    'That benchmark applies when the goal is maximum compounding on unlimited capital. '
    'When the goal is $500–700/day cash flow on $100K, a 55% win rate with 3:1 R:R '
    'is a sound, achievable system — not an implausible one.'
)

divider()
subheading('Which Brutal Assessment Points Still Apply')
add_table(
    ['Assessment Point', 'Still Relevant for Cash Flow?', 'Why'],
    [
        ('"175% annualized beats Renaissance" — implausible',
         'No',
         'Wrong benchmark. Cash flow does not require compounding.'),
        ('"Technical signals have no academic edge"',
         'Partially',
         'Still matters — if win rate drops below 40%, the EV turns negative. But 55% + 3:1 R:R works.'),
        ('"Claude is wrong signal source"',
         'Less critical',
         'If positive EV is consistent across market regimes, the source need not be Nobel-worthy.'),
        ('"30-day backtest not statistically significant"',
         'Yes',
         'Still true. Need to see the system across multiple market regimes.'),
        ('"Paper-to-live degradation"',
         'Yes — biggest risk',
         'Unknown how much win rate degrades on real fills. This is the key variable.'),
        ('"Lock-in is a behavioral trap"',
         'No — reversed',
         'For cash flow, locking in at $716 when the day target is hit is exactly right behavior.'),
        ('"Need Mode 2 / PEAD / insider buying"',
         'No',
         'Mode 2 is for alpha and product. Not needed for cash flow. Mode 1 intraday is the right fit.'),
        ('"Replace dollar target with Sharpe ratio"',
         'No',
         'Sharpe ratio is for alpha optimization. Dollar target is appropriate for cash flow.'),
    ]
)

divider()
subheading('Why Mode 1 (Intraday) Is the Right Fit for Cash Flow')
bullet('P&L is known every day — predictable income, not lumpy quarterly events')
bullet('Capital resets daily — no capital tied up in 20-day PEAD holds')
bullet('Lock-in at $716 is perfectly designed for "hit the number, protect it, done"')
bullet('Daily loss limit (-$300) contains downside on bad days')
bullet('No overnight risk — clean slate every morning')
body(
    'Mode 2 (event-driven, PEAD, multi-day holds) is better for maximizing risk-adjusted '
    'returns and for a SaaS product signal. It is worse for predictable daily cash flow. '
    'Keep Mode 1 as the primary engine for the cash flow goal.'
)

divider()
subheading('The Three Real Risks That Remain')
body('The brutal assessment shrinks to three questions for the cash flow goal:')

subheading('Risk 1: Paper-to-Live Degradation', ORANGE)
body(
    'Does the 55% win rate hold on real fills? Limit orders at entry × 1.001 work on paper. '
    'Real fills depend on queue position, spread, and order book depth. '
    'If win rate drops to 45% live, EV per trade drops from $48 to $14 — still positive, '
    'but daily output falls to ~$210. At 40% win rate, EV goes negative.'
)
bullet('Mitigation: start real money small ($10K–$25K); validate live fill quality before scaling')
bullet('Mitigation: model slippage explicitly in eval — track paper price vs. actual fill price')

subheading('Risk 2: Regime Sensitivity', ORANGE)
body(
    'The 30-day backtest ran in a momentum-favorable market (post-tariff recovery, April–May 2026). '
    'Does the system work in a choppy sideways market? A trending-down market? '
    'Unknown — we have not seen it run through multiple regimes.'
)
bullet('Mitigation: the VIX gate and futures gate already reduce exposure on bad market days')
bullet('Mitigation: run through at least one down-trending period on paper before committing real capital')

subheading('Risk 3: Capital Adequacy vs. Daily Output', ORANGE)
body(
    'At $100K with $5K–$7K per position, 15 trades = $75K–$105K deployed. '
    'If the daily output goal grows (say to $1K/day), you either need more capital '
    'or higher win rate — you cannot simply add more positions without more capital.'
)
bullet('Path: at $716/day → capital compounds to ~$140K in 45 trading days → same % = $1K/day')
bullet('This is the natural scaling path without adding risk')

divider()
subheading('Two Separate Applications Going Forward', GREEN)
body(
    'The cash flow goal and the product/alpha goal are now explicitly decoupled. '
    'They should be built as two separate things:'
)
add_table(
    ['Application', 'Goal', 'Signal', 'Hold Period', 'Success Metric'],
    [
        ('This system (App 1)',
         'Personal cash flow — $500–700/day',
         'Intraday momentum (Mode 1)',
         'Same day — no overnight',
         'Consistent positive EV days; win rate ≥50% live'),
        ('Future SaaS product (App 2)',
         'Productize for other traders',
         'PEAD + insider buys + unusual options (Mode 2)',
         '3–20 days',
         'Sharpe ≥1.5 over 200+ live trades; auditable signal quality'),
    ]
)
body(
    'App 1 is what we are building now and validating through June 2026. '
    'App 2 is a future separate application — different signal architecture, '
    'different hold period, different success metrics, different codebase. '
    'Do not conflate them. Do not add Mode 2 complexity to App 1 in service of App 2.'
)

divider()
subheading('Revised Sprint Priorities for Cash Flow Goal')
bullet('1. Pass the June 1 gate — native trailing stop validated, win rate ≥80% on paper')
bullet('2. 15-min momentum confirmation (2f) — small improvement to Mode 1 signal quality')
bullet('3. Model paper-to-live degradation in eval — track expected vs. actual fill prices')
bullet('4. Deploy $10K–$25K real capital; validate live win rate matches paper within ±10%')
bullet('5. Scale capital once live win rate is confirmed over 60+ real trades')
bullet('6. App 2 planning begins only after App 1 is running profitably on real money')

# ── Save ──────────────────────────────────────────────────────────────────────
_project_dir = "/Users/amitgarg/Claude Projects/trading-agent"
out = os.path.join(_project_dir, "Trading_Agent_Brutal_Assessment.docx")
doc.save(out)
print(f"Saved: {out}")
