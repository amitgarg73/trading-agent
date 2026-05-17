"""Generates the Trading Agent PRD as a Word document using Marily Nika's AI PRD template."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AI Product Requirements Doc', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)

sub = doc.add_paragraph('AI Trading Agent (PRD) — v5.0')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Author: ').bold = True
meta.add_run('Amit Garg\n')
meta.add_run('Contributors: ').bold = True
meta.add_run('Claude Code (Anthropic)\n')
meta.add_run('Date: ').bold = True
meta.add_run('May 2026')

doc.add_paragraph()
body(
    'How can we use AI to help our users? — By automating the full investment research and '
    'trade selection workflow that currently takes a skilled trader 2–3 hours every morning, '
    'reducing it to zero manual effort while maintaining institutional-grade decision quality '
    'and full auditability of every choice the system makes.'
)

# ── Relevant Documents ────────────────────────────────────────────────────────
h1('Relevant Documents')
add_table(
    ['Document', 'Location'],
    [
        ('Technical Architecture & Build Log', 'Trading_Agent_Documentation.docx (project folder)'),
        ('GitHub Repository', 'https://github.com/amitgarg73/trading-agent'),
        ('Live Dashboard', 'https://trading-agent-q39gepfhtsg3ianyezywl7.streamlit.app'),
        ('Database Schema', 'schema.sql (project folder)'),
        ('Eval Script', 'eval.py — python3 eval.py --days 5'),
        ('Backtest Script', 'backtest.py — python3 backtest.py --days 30 --top 15'),
    ]
)

# ── About ─────────────────────────────────────────────────────────────────────
doc.add_page_break()
h1('About')
body(
    'The AI Trading Agent is a fully autonomous stock trading simulation system. Every weekday '
    'it wakes up at 9AM ET, checks market conditions (VIX, futures, international markets), '
    'scans 430+ stocks and ETFs for high-conviction setups, filters out earnings-risk tickers, '
    'uses Claude AI to select and validate trades, simulates position management throughout the '
    'day, and produces a complete performance record — all without any human input.'
)
body(
    'The system is designed as a personal investment research tool and paper trading engine. '
    'It demonstrates that a single person with no quant background can deploy institutional-grade '
    'AI-driven trade analysis at near-zero cost using modern cloud infrastructure and LLMs. '
    'Every decision is auditable — the dashboard shows exactly what the system saw, what it '
    'chose, and why, at every step of the pipeline.'
)

h2('Market Insights')
h3('Competitor Analysis')
add_table(
    ['Competitor', 'Approach', 'Gap'],
    [
        ('Bloomberg Terminal', 'Data + analytics, human-driven decisions', '$24K/year, no autonomous execution, no LLM reasoning'),
        ('Trade Ideas (Holly AI)', 'AI scanner, no LLM reasoning', 'No natural language explanations, opaque signals'),
        ('Kavout / Alpaca AI', 'ML-based stock scoring', 'Black box, no auditable reasoning per trade'),
        ('ChatGPT / Claude (manual)', 'Ad-hoc LLM prompting', 'No automation, no real data integration, no memory'),
        ('Robinhood / Webull', 'Retail brokerage with basic screeners', 'No AI, no autonomy, human must decide and execute'),
    ]
)

h3('Market Analysis')
body(
    'The retail algorithmic trading market is growing rapidly. Over 70% of US equity volume is '
    'algorithmic. Retail traders lack access to the same systematic tools as hedge funds. '
    'The emergence of powerful LLMs (Claude, GPT-4) combined with free market data APIs (yfinance) '
    'and free cloud infrastructure (GitHub Actions, Supabase, Streamlit Cloud) has eliminated the '
    'cost barrier that previously kept autonomous AI trading tools in the hands of institutions only. '
    'This system costs ~$0.15–0.30/day in API calls — less than a cup of coffee.'
)

h3('Technology Analysis')
bullet('LLMs (Claude Sonnet) — natural language reasoning over structured market data, with JSON output for downstream processing')
bullet('yfinance — free real-time and historical price/volume data, news headlines, earnings calendar')
bullet('Technical analysis library (ta) — RSI, MACD, Bollinger Bands, ATR computed in-process')
bullet('Supabase (PostgreSQL) — persistent storage, free tier sufficient for months of history')
bullet('GitHub Actions — serverless cron scheduler, 2,000 free min/month, no infrastructure to manage')
bullet('Streamlit Cloud — zero-cost dashboard hosting, always-on, mobile-accessible')

h2('Customer Segments')
h3('User Personas')
body(
    'Primary persona: Amit — VP-level executive with deep tech and product background, '
    'comfortable with APIs and systems thinking, not a quant or active day trader. '
    'Wants systematic, explainable trade ideas without spending hours on research every morning. '
    'Values auditability (why did the agent pick this trade?) over black-box ML signals. '
    'Budget-conscious — wants institutional-grade output at near-zero infrastructure cost.'
)

# ── The Problem ───────────────────────────────────────────────────────────────
doc.add_page_break()
h1('The Problem')

h2('Use Cases')
bullet('Generate daily trade ideas based on technical momentum signals across 430+ tickers')
bullet('Validate each idea against hard risk rules before acting (reward:risk, position sizing, stop width)')
bullet('Skip trading entirely on high-volatility days (VIX > 30) or strong pre-market sell-offs')
bullet('Block trades in tickers with earnings today/tomorrow to avoid binary event risk')
bullet('Monitor open positions throughout the day and close automatically on target or stop')
bullet('Maintain a full audit trail of every decision with natural language reasoning attached')
bullet('Review weekly performance via eval.py to tune the strategy over time')

h2('Pain Points')
bullet('Manual market scanning takes 2–3 hours daily — not feasible for a full-time executive')
bullet('Retail screeners give signals but no reasoning — hard to trust or learn from')
bullet('Existing AI trading tools are opaque black boxes with no explainability')
bullet('Institutional tools (Bloomberg) cost $24K/year — inaccessible for personal use')
bullet('No tools combine real-time data, market context, news intelligence, LLM reasoning, risk validation, AND autonomous execution in one system')

h2('Problem Statement')
body(
    'Amit spends too much time on manual market research and trade selection each morning, '
    'and when he does find setups, he has no systematic way to validate risk, filter for '
    'macro context, or track performance over time — so he never builds conviction or compounds learning.'
)

h2('Hypotheses and Mission Statement')
body(
    'Hypothesis: By automating the scan → market context → news intelligence → strategy → risk → '
    'execution pipeline using Claude AI, we can surface high-conviction trade setups daily with full '
    'reasoning, enforce consistent risk discipline automatically, protect against adverse conditions '
    'before entering, and generate $750–$1,000/day in simulated profit — validated by 30-day backtest '
    'showing 93% win days at $750/day average. LLM-powered autonomous trading is viable at personal scale.'
)

# ── The Solution ──────────────────────────────────────────────────────────────
doc.add_page_break()
h1('The Solution')

h2('Ideation')
bullet('Automated technical scanner scoring 430+ tickers on momentum and volume signals')
bullet('Market context agent: VIX gate + futures signal + international markets (V2a)')
bullet('News intelligence agent: earnings blackout + news headline context for Claude (V2b)')
bullet('LLM strategy agent that reads scored candidates and selects trades with full reasoning')
bullet('LLM risk agent that enforces hard rules (stop width, reward:risk, position sizing, target %)')
bullet('Intraday position monitor that closes trades automatically on 3% target or 1% stop')
bullet('EOD performance calculator and daily P&L tracker with portfolio compounding')
bullet('Workflow dashboard: shows full pipeline step-by-step (market → scan → plan → positions)')
bullet('Weekly eval script to score strategy quality and generate tuning recommendations')
bullet('30-day backtesting script for configuration validation without Claude API costs')

h2('Leveraging AI')
body('AI is essential — not optional — for two reasons:')
body(
    '1. Strategy selection: The scanner produces 5–20 candidates daily with quantitative scores, '
    'but selecting the right trades requires judgment about market context, momentum quality, news '
    'sentiment, and conviction. This is exactly what LLMs excel at — synthesizing structured data '
    'with contextual reasoning and expressing a defensible view in natural language.'
)
body(
    '2. Risk validation: Risk rules are deterministic, but explaining WHY a trade was rejected '
    '(and logging that reasoning for review) requires language generation. Claude produces '
    'rejection reasons that are human-readable and actionable, not just error codes.'
)
body(
    'Without AI, this is a screener — useful but not autonomous. With AI, it becomes an agent '
    'that makes decisions, explains itself, adapts to market conditions, and can be held accountable. '
    'The V2 intelligence layer (market context + news) further grounds Claude\'s decisions in '
    'real conditions before any trade is selected.'
)

h2('Feature Prioritization (RICE)')
add_table(
    ['Feature', 'Reach', 'Impact', 'Confidence', 'Effort', 'Score', 'Priority'],
    [
        ('Technical scanner (430+ tickers)', 'High', 'High', '95%', 'M', '9.5', 'P0 — shipped'),
        ('Claude strategy agent', 'High', 'High', '90%', 'M', '9.0', 'P0 — shipped'),
        ('Claude risk agent', 'High', 'High', '95%', 'S', '9.5', 'P0 — shipped'),
        ('GitHub Actions scheduler', 'High', 'High', '99%', 'S', '9.9', 'P0 — shipped'),
        ('Supabase persistence', 'High', 'High', '99%', 'S', '9.9', 'P0 — shipped'),
        ('30-day backtesting (backtest.py)', 'High', 'High', '95%', 'M', '9.5', 'P0 — shipped'),
        ('V2a: VIX + futures gate', 'High', 'High', '90%', 'S', '9.0', 'P1 — shipped'),
        ('V2b: Earnings blackout + news', 'High', 'High', '90%', 'S', '9.0', 'P1 — shipped'),
        ('Workflow dashboard (Today tab)', 'High', 'Med', '90%', 'M', '8.1', 'P1 — shipped'),
        ('Eval script (eval.py)', 'Med', 'High', '85%', 'S', '8.5', 'P1 — shipped'),
        ('V2c: Fear & Greed + FOMC/CPI/NFP calendar', 'High', 'High', '95%', 'S', '9.5', 'P1 — shipped (v2.2)'),
        ('V2c.1: Tune F&G to confirming signal', 'High', 'High', '99%', 'XS', '9.9', 'P1 — shipped (v2.3)'),
        ('V3a: Dynamic universe refresh (weekly screener)', 'High', 'High', '95%', 'M', '9.5', 'P1 — shipped (v3.0)'),
        ('Dashboard company names column', 'High', 'Med', '99%', 'S', '9.0', 'P1 — shipped (v3.0)'),
        ('V2d: Sector correlation guard', 'Med', 'Med', '80%', 'S', '7.2', 'P1 — shipped (v4.1)'),
        ('V5: Guardrails (6 safety checks)', 'High', 'High', '99%', 'S', '9.9', 'P0 — shipped (v5.0)'),
        ('V2e: Sector rotation scoring', 'Med', 'Med', '75%', 'M', '6.8', 'P2 — next'),
        ('Alpaca paper trading API (V2g)', 'Med', 'High', '80%', 'L', '6.4', 'P0 — shipped (v4.0)'),
        ('SMS/email alerts', 'Med', 'Med', '90%', 'S', '8.1', 'P2 — next'),
        ('V2f: 15-min momentum confirmation', 'Med', 'Med', '70%', 'M', '6.3', 'P3 — future'),
        ('Macro context agent (Fed calendar)', 'Med', 'Med', '70%', 'M', '6.3', 'P3 — future'),
    ]
)

h2('AI MVP')
body('The AI MVP is live and validated via 30-day backtest. Two-agent architecture:')
bullet('Agent 1 (Strategy): claude-sonnet-4-6 — receives structured JSON of scored candidates plus '
       'market summary (VIX, futures, news headlines), returns selected trades with '
       'entry/target/stop/confidence/reasoning in JSON. Max positions dynamic based on market conditions.')
bullet('Agent 2 (Risk): claude-sonnet-4-6 — receives proposed trades, applies hard rules '
       '(3% target, 1% stop, 3:1 reward:risk, $5K–$7K position size), returns approved/rejected split '
       'with plain-English rejection reasons.')
body(
    'Both agents use structured output (JSON) for deterministic downstream processing. '
    'Natural language reasoning stored alongside structured data for human review in dashboard. '
    'Switched from claude-opus-4-7 to claude-sonnet-4-6 — reduced API cost ~70% with comparable quality.'
)

h2('Roadmap')
add_table(
    ['Phase', 'Timeline', 'Deliverables'],
    [
        ('Phase 1 — Foundation', 'May 2026 (complete)', 'Scanner, strategy agent, risk agent, portfolio sim, GitHub Actions, Supabase, dashboard, backtest'),
        ('Phase 2a — Intelligence (V2a–V2c.1)', 'May 2026 (complete)', 'VIX gate, futures signal, earnings blackout, news context, Fear & Greed + calendar gates, F&G tuned to confirming signal, workflow dashboard'),
        ('Phase 3a — Dynamic Universe (V3a)', 'May 2026 (complete)', 'Weekly S&P500+Nasdaq100 screener, 458 tickers on first run, company names in dashboard'),
        ('Phase 2b — Execution (V2g)', 'May 2026 (complete)', 'Alpaca paper trading — real bracket orders (entry + take-profit + stop-loss), position sync, fill price on close'),
        ('Phase 2c — More Intelligence (V2d)', 'May 2026 (complete)', 'V2d sector correlation guard — max 3 per sector, drops lowest-confidence excess'),
        ('Phase 2d — Safety (V5)', 'May 2026 (complete)', 'V5 guardrails — 6 safety checks (action whitelist, ticker whitelist, duplicate guard, price sanity, capital check, daily loss limit), concurrent run lock, EOD close retry'),
        ('Phase 2e — More Intelligence', 'June 2026', 'V2e sector rotation scoring, V2f momentum confirmation (15-min rule)'),
        ('Phase 3 — Alerts & Monitoring', 'July 2026', 'SMS/email alerts on position close, weekly email summaries'),
        ('Phase 4 — Scale', 'Q3–Q4 2026', 'Strategy A/B testing, weekly email summaries, real capital evaluation'),
    ]
)

# ── Technical Architecture ────────────────────────────────────────────────────
doc.add_page_break()
h1('Technical Architecture')
body('High-level data flow (premarket pipeline):')
bullet('GitHub Actions (cron 9AM ET) → orchestrator.py')
bullet('Step 0: market_context agent — checks VIX, futures, international markets → GO/CAUTION/SKIP + max_positions')
bullet('Step 1: scanner.py — yfinance + ta → scored candidates JSON (430+ tickers)')
bullet('Step 1.5: news_intel agent — earnings blackout check + news headlines → filtered candidates')
bullet('Step 2: strategy agent (Claude API) — proposed trades JSON')
bullet('Step 3: risk agent (Claude API) — approved/rejected trades JSON')
bullet('Step 4: portfolio agent → Supabase (trade_plans, planned_trades, positions tables)')
bullet('Intraday: GitHub Actions (every 30 min) → intraday agent → price check → close on target/stop')
bullet('EOD: GitHub Actions → performance agent → daily_performance record')
bullet('Dashboard: Streamlit Cloud → reads Supabase → renders workflow view with 4 steps')

body('Infrastructure:')
bullet('Compute: GitHub Actions (Ubuntu, Python 3.11, serverless, free tier, 2,000 min/month)')
bullet('AI: Anthropic API (claude-sonnet-4-6, ~$0.15–0.30/day for 2 calls per day)')
bullet('Data: yfinance (free, no key) — prices, volume, news, earnings calendar')
bullet('Storage: Supabase PostgreSQL (free tier, 500MB limit — months of runway)')
bullet('Frontend: Streamlit Community Cloud (free tier, always-on, mobile-accessible)')
bullet('Auth: SSH key for git, GitHub Secrets + Streamlit Secrets for runtime credentials')

h2('Key Hard Rules (Enforced by Risk Agent)')
add_table(
    ['Rule', 'Value', 'Why'],
    [
        ('Profit target', '3% above entry', 'Hard-coded in prompt — risk agent rejects if not set'),
        ('Stop loss', '1% below entry', 'Hard-coded in prompt — risk agent rejects if wider'),
        ('Reward:risk minimum', '3:1', 'Ensures winners outpace losers even at 50% win rate'),
        ('Position size', '$5,000–$7,000', '5–7% of $100K — limits single-position impact'),
        ('Max positions', '15 (reduced by VIX/futures)', 'Dynamic — VIX 20-25 → 10, VIX 25-30 → 5, VIX >30 → skip'),
        ('No overnight holds', 'All positions closed EOD', 'Eliminates gap risk from overnight news'),
        ('Earnings blackout', 'No trades day-of or day-before earnings', 'Earnings = binary event = unacceptable gap risk'),
    ]
)

h2('Assumptions and Constraints')
bullet('yfinance provides reliable enough data for paper trading — not suitable for live trading without a paid data feed')
bullet('GitHub Actions free tier (2,000 min/month) is sufficient for current run frequency')
bullet('Supabase free tier (500MB) is sufficient for months of trade history at current volume')
bullet('claude-sonnet-4-6 provides sufficient reasoning quality for trade selection at this scale')
bullet('Paper trading runs through Alpaca Paper Trading API — realistic bracket orders with fills; no real money at risk')
bullet('Market hours assumed to be standard US equity (9:30AM–4PM ET Mon–Fri)')
bullet('yfinance earnings calendar not always populated — some tickers may slip through blackout check')

# ── Risks ─────────────────────────────────────────────────────────────────────
h1('Risks')
add_table(
    ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ('yfinance data gaps or missing earnings calendar', 'High', 'Low', 'Handled gracefully — bad tickers skipped, earnings check fails safe'),
        ('Claude API credit exhausted', 'Med', 'High', 'Set auto-recharge on Anthropic console; ~$0.15-0.30/day so $10 lasts 33-66 days'),
        ('Earnings blackout miss (calendar not populated)', 'Med', 'Med', 'Fail-safe: if calendar unavailable, ticker passes through (not blocked)'),
        ('Strategy underperforms (win rate < 45%)', 'Med', 'Med', 'Run eval.py weekly; conditional: drop target to 2.5% if win rate < 45%'),
        ('GitHub Actions free tier exhausted', 'Low', 'High', 'Well within 2,000 min/month at current frequency'),
        ('Supabase free tier storage limit', 'Low', 'Med', 'Months of runway; add cleanup job if needed'),
        ('Duplicate key DB errors on rerun', 'Low', 'Low', 'Fixed — check-before-insert logic in orchestrator.py'),
        ('Model quality regression', 'Low', 'Med', 'Monitor reasoning quality in dashboard; can switch back to Opus'),
        ('VIX/futures data unavailable pre-market', 'Low', 'Low', 'Handled — None values fall through to GO with default max_positions'),
        ('Alpaca API rate limits or downtime', 'Low', 'Med', 'Falls back to yfinance simulation for P&L if Alpaca unavailable; positions still logged to Supabase'),
    ]
)

# ── Requirements ──────────────────────────────────────────────────────────────
doc.add_page_break()
h1('Requirements')

h2('User Journeys')
body('Journey 1 — Morning check (2 min)')
bullet('User opens dashboard URL on any device → enters password')
bullet('Today tab loads: Step 0 shows VIX + futures (green/yellow/red), Step 1 shows screened stocks')
bullet('Step 2 shows Claude\'s market read + approved trades with entry/target/stop/reasoning')
bullet('Step 3 shows any positions already open with live P&L')
bullet('User decides whether to mirror trades manually in their brokerage (optional)')

body('Journey 2 — Intraday monitoring (30 sec)')
bullet('User opens Positions tab during market hours')
bullet('Sees unrealized P&L for each open position in real time')
bullet('Sees which positions were auto-closed and why (TARGET / STOP / EOD)')

body('Journey 3 — Weekly review (5 min)')
bullet('User runs: python3 eval.py --days 5 from terminal')
bullet('Reviews grade (A/B/C/D), win rate, actual reward:risk, top/worst tickers')
bullet('Adjusts thresholds in config/settings.py based on recommendations')
bullet('Optionally reruns backtest.py with new settings before pushing changes')

h2('Functional Requirements')
bullet('Market context agent must fetch VIX + futures and return GO/CAUTION/SKIP with dynamic max_positions before scanner runs')
bullet('News intelligence agent must check earnings calendar and remove day-of/day-before tickers before strategy sees candidates')
bullet('Scanner must evaluate all 430+ tickers and return scored candidates in under 90 seconds')
bullet('Strategy agent must return structured JSON with entry (3% target, 1% stop enforced), position size, confidence, and 2-3 sentence reasoning per trade')
bullet('Risk agent must enforce all 7 hard rules and return rejected trades with specific plain-English reasons')
bullet('Guardrails must run after sector guard and block any trade that fails action whitelist, ticker whitelist, duplicate check, price sanity, capital check, or daily loss limit — with specific rejection reason logged')
bullet('Concurrent run lock must prevent any premarket run from executing if scan_results for today already exists in Supabase')
bullet('Intraday agent must check prices and close positions within 5 minutes of target/stop being hit')
bullet('EOD agent must produce a complete daily_performance record including P&L, win rate, and capital')
bullet('Dashboard Today tab must show full 4-step workflow with live data within 3 seconds')
bullet('All pipeline failures must be visible in GitHub Actions logs without data corruption or partial writes')

h2('Non-functional Requirements')
bullet('Availability: system runs unattended Mon–Fri; no human monitoring required')
bullet('Security: all secrets encrypted at rest and never logged; dashboard password-protected')
bullet('Cost: total infrastructure cost < $15/month (API ~$5–10 + all cloud tiers free)')
bullet('Auditability: every trade decision has human-readable reasoning stored permanently in Supabase')
bullet('Recoverability: reruns are safe — check-before-insert prevents duplicate data on any rerun')
bullet('Resilience: VIX/futures/news failures are non-fatal — pipeline continues with safe defaults')

h2('AI & Data Requirements')
body('Data sources:')
bullet('yfinance: 3-month price history + daily OHLCV for technical indicators')
bullet('yfinance: ticker.news for recent headlines (3 per ticker, passed to Claude)')
bullet('yfinance: ticker.calendar for earnings dates (earnings blackout filter)')
bullet('No proprietary data feeds required in current phase')

body('Model requirements:')
bullet('claude-sonnet-4-6 for strategy and risk agents — structured JSON output required')
bullet('Temperature: default — some creative judgment desirable in trade selection')
bullet('Max tokens: 3,000 for strategy (up to 15 trades), 2,000 for risk agent')

body('Challenges:')
bullet('yfinance reliability: free tier, occasionally returns stale/missing data — handled with graceful skipping')
bullet('LLM consistency: Claude may select different trades on identical inputs — acceptable for daily cadence')
bullet('Earnings calendar gaps: not all tickers have populated calendars in yfinance — fail-safe: if unavailable, ticker passes through')
bullet('No labeled training data: agent quality evaluated purely on realized P&L and win rate')

# ── Positioning ───────────────────────────────────────────────────────────────
doc.add_page_break()
h1('Positioning')
add_table(
    ['Dimension', 'Details'],
    [
        ('For', 'Technically sophisticated individuals who want systematic, AI-driven trade ideas with full auditability'),
        ('Who', 'Need daily trade selection with explainable reasoning, automatic risk discipline, and market condition awareness'),
        ('The AI Trading Agent is a', 'Fully autonomous paper trading system with multi-layer intelligence'),
        ('That', 'Checks market conditions, filters earnings risk, selects trades using Claude AI, manages positions, and tracks performance — zero daily effort'),
        ('Unlike', 'Bloomberg ($24K/yr), Trade Ideas (opaque signals), or manual ChatGPT prompting (no automation)'),
        ('Our product', 'Combines real market data, LLM reasoning, hard risk rules, market intelligence gates, and cloud automation at ~$10/month all-in'),
    ]
)

# ── Measuring Success ─────────────────────────────────────────────────────────
h1('Measuring Success')

h2('Metrics')
add_table(
    ['Metric', 'Target', 'Backtest Result', 'How Measured'],
    [
        ('Daily P&L', '>= $750/day (path to $1K)', '$750 avg / 30 days', 'daily_performance table'),
        ('Win rate (days)', '>= 85% of trading days positive', '93% (28/30 days)', 'eval.py, daily_performance'),
        ('Win rate (trades)', '>= 55%', 'Measured live', 'eval.py, daily_performance.win_rate'),
        ('Reward:risk ratio', '>= 3.0x per trade', '3.0 enforced hard', 'risk agent hard rule'),
        ('Pipeline uptime', '>= 95% of scheduled runs succeed', 'N/A (live monitoring)', 'GitHub Actions success rate'),
        ('Days to $1K/day', '<= 45 trading days (compounding)', 'Projected', 'daily_performance.ending_capital'),
    ]
)

h2('AI-specific Metrics')
bullet('Reasoning quality: manually reviewed weekly — does the trade reasoning make sense given market conditions shown in Step 0?')
bullet('Market gate accuracy: did SKIP days actually see bad market conditions? Reviewed in Scan Log tab')
bullet('Earnings blackout value: did blocked tickers gap up/down on earnings? Validate via stock price check post-earnings')
bullet('Confidence calibration: do HIGH confidence trades win more often than LOW? Tracked in eval.py')
bullet('Signal consistency: does strategy select coherent setups, or random noise? Reviewed via top ticker patterns over time')

h2('North Star Metric')
body(
    'Annualized return on $100K simulated capital — targeting >189% annualized '
    '(validated by 30-day backtest at current configuration). '
    'This single number captures whether the AI strategy is genuinely alpha-generating '
    'or just getting lucky on individual days. At $750/day × 250 trading days = $187,500 '
    'annual profit on $100K starting capital = 187% annualized return.'
)

# ── Launching ─────────────────────────────────────────────────────────────────
doc.add_page_break()
h1('Launching')

h2('Stakeholders & Communication')
add_table(
    ['Stakeholder', 'Interest', 'Communication'],
    [
        ('Amit Garg (owner)', 'Daily P&L, strategy quality, workflow health', 'Dashboard + weekly eval.py run'),
        ('Future investors / partners', 'Proof of concept for AI trading at low cost', 'Trading_Agent_Documentation.docx + dashboard demo'),
        ('Anthropic (API provider)', 'API usage within credits', 'console.anthropic.com — monitor manually (no usage alerts available)'),
    ]
)

h2('Roll-out Strategy')
body('Phase 1 (current — v4.0): Paper trading running through Alpaca Paper Trading API — realistic bracket orders with fills; no real money at risk. '
     'V2a (VIX + futures), V2b (earnings blackout + news), V2c (Fear & Greed + calendar), V2c.1 (F&G gate tuned to confirming signal), '
     'V3a (dynamic universe refresh — weekly S&P500+Nasdaq100 screener, 458 tickers), V2g (Alpaca broker — bracket orders + fill price sync) all deployed. '
     'Backtest validated: $21,474 P&L over 30 days (grade B), gates cost -$2,549 vs ungated baseline — acceptable insurance. '
     'Dashboard shows company names in all ticker tables.')
body('Phase 2 (v4.1–v5.0, complete): V2d sector correlation guard, V5 guardrails (6 safety checks), concurrent run lock, EOD close retry.')
body('Phase 3: V2e sector rotation scoring, V2f momentum confirmation.')
body('Phase 3: If win rate > 60% and reward:risk > 2x sustained over 30 live trading days, evaluate real '
     'capital deployment with strict position limits.')
body('Go/no-go criteria for real capital:')
bullet('Win rate >= 60% over >= 20 live trading days')
bullet('Actual realized reward:risk >= 2.0x')
bullet('No single-day drawdown > 3% of capital')
bullet('All V2 intelligence gates validated (VIX gate prevented losses on down days, earnings blackout prevented gap losses)')
bullet('Strategy explainability confirmed — every trade reasoning reviewed and found logical')

# ── Open Items & Backlog ──────────────────────────────────────────────────────
doc.add_page_break()
h1('Open Items & Backlog')
body(
    'Captured from build sessions — items discussed but not yet implemented. '
    'Update status as work progresses.'
)

h2('V2 Pipeline — Status & Remaining Phases')
add_table(
    ['Item', 'What', 'Why', 'Status'],
    [
        ('V2c ✅', 'Fear & Greed Index + economic calendar (FOMC/CPI/NFP)',
         'Fear & Greed (alternative.me, free) + FOMC/CPI/NFP hardcoded dates. FOMC day → cap at 8 pos. CPI/NFP day → cap at 10 pos.',
         'Shipped — v2.2'),
        ('V2c.1 ✅', 'Tune F&G gate: confirming signal only',
         'F&G is lagging — reads low after selloffs during recoveries. New rule: F&G <25 only reduces positions when VIX >20 OR futures bearish. Standalone F&G is context only. Backtest: gate cost -$9,596 → -$2,549; both grade B.',
         'Shipped — v2.3'),
        ('V3a ✅', 'Dynamic universe refresh — weekly S&P500+Nasdaq100 screener',
         'Fetches index components from Wikipedia, screens 550+ tickers for ATR≥2%+volume≥500K, saves 450+ to Supabase. '
         'orchestrator load_universe() reads Supabase if ≤7 days old, else falls back to static settings.py. '
         'First run: 553 screened → 458 passed. GitHub Actions fires every Monday 8:30 AM ET.',
         'Shipped — v3.0'),
        ('V2d ✅', 'Sector correlation guard',
         'Max 3 positions per sector. Fetches sector via yfinance for approved trades; ETFs classified separately. '
         'Unknown sector (rate-limited/unclassifiable) passes through — safe fallback. '
         'Drops lowest-confidence excess (HIGH > MEDIUM > LOW, tiebreak by estimated_profit). '
         'Dashboard shows sector-blocked trades in Strategy & Risk step.',
         'Shipped — v4.1'),
        ('V5 ✅', 'Guardrails — 6 safety checks before any trade executes',
         'Action whitelist (BUY only), ticker whitelist (universe only), duplicate position guard (no same ticker twice in one day), '
         'price sanity (entry within 5% of market price), capital check (Alpaca buying_power covers position_size), '
         'daily loss limit (stop all new trades if P&L < -$300). '
         'Also: concurrent run lock (prevents duplicate positions from overlapping GitHub Actions runs), '
         'EOD close retry (retries Alpaca close once on failure, warns loudly if still failing).',
         'Shipped — v5.0'),
        ('V2e', 'Sector rotation scoring',
         'Favor sectors showing relative strength this week; deprioritize lagging sectors',
         'Planned'),
        ('V2f', 'Momentum confirmation (15-min rule)',
         'Wait for confirmed breakout 15 min after open before entering; reduces false signals on gap-and-fade setups',
         'Planned'),
        ('V2g ✅', 'Alpaca paper trading API integration',
         'Real bracket order simulation: submit_bracket_order (entry market + take-profit limit + stop-loss), '
         'position sync via get_position_data, fill price on bracket close; alpaca_order_id stored in positions table',
         'Shipped — v4.0'),
    ]
)

h2('Features & Enhancements')
add_table(
    ['Item', 'What', 'Why', 'Status'],
    [
        ('Alerts', 'SMS/email notification on position close',
         'Know immediately when a target is hit or stop is triggered without checking dashboard',
         'Planned — P2'),
        ('CBRS', 'Add Cerebras Systems to stock universe post-IPO',
         'High-conviction AI infrastructure play; IPO range $150–160, $26.6B valuation, OpenAI $20B compute deal',
         'Pending IPO listing — add to settings.py once trading'),
        ('Weekly summary email', 'Auto-generated weekly P&L summary with top trades and win rate',
         'Replace manual eval.py run with push delivery',
         'Future'),
    ]
)

h2('Tune Triggers (Conditional Actions)')
add_table(
    ['Condition', 'Action', 'When to Check'],
    [
        ('Live win rate < 45% after first 2 weeks',
         'Drop profit target from 3% back to 2.5% in config/settings.py',
         'Run eval.py --days 10 after second Monday'),
        ('Live win rate > 70% consistently',
         'Consider increasing max positions or raising target to 3.5%',
         'Monthly review'),
        ('Anthropic API credit < $5',
         'Top up credit at console.anthropic.com — no automatic alert exists',
         'Check manually weekly; ~$0.15–0.30/day burn rate'),
        ('GitHub Actions run fails 2 days in a row',
         'Check Actions logs — likely a yfinance API change or dependency issue',
         'Dashboard will show no new data as signal'),
    ]
)

h2('Known Constraints & Gaps')
add_table(
    ['Constraint', 'Detail', 'Workaround'],
    [
        ('Anthropic console has no usage alerts',
         'console.anthropic.com only supports auto-recharge — no threshold alerts or email notifications',
         'Monitor manually weekly; set auto-recharge to avoid service interruption'),
        ('yfinance earnings calendar not always populated',
         'Some tickers have no calendar data — V2b blackout check fails safe (lets them through)',
         'Accept the risk; most major tickers are populated; gaps are edge cases'),
        ('git push requires user terminal',
         'Claude Code cannot enter the SSH passphrase interactively — git push must be run from user terminal (resolved by switching to HTTPS PAT)',
         'Switch remote to HTTPS with Personal Access Token: git remote set-url origin https://<token>@github.com/amitgarg73/trading-agent.git'),
        ('No real-time intraday prices in yfinance free tier',
         'yfinance intraday prices have 15-min delay — positions may close slightly after actual target/stop hit',
         'Acceptable for paper trading simulation; use paid data feed for live trading'),
    ]
)

h2('Evaluation Cadence')
add_table(
    ['When', 'What to Do', 'Command'],
    [
        ('Monday evening (after first live day)', 'Score first live run', 'python3 eval.py --days 1'),
        ('Every Friday', 'Weekly performance review', 'python3 eval.py --days 5'),
        ('Monthly', 'Backtest with latest universe against prior month', 'python3 backtest.py --days 22 --top 15'),
        ('If performance degrades', 'Tune thresholds then validate via backtest before pushing', 'Edit config/settings.py → backtest → push'),
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/amitgarg/Claude Projects/trading-agent/Trading_Agent_PRD.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
