"""
Generate Trading_Agent_Design.docx from DESIGN.md.
Run: python3 generate_design.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

VERSION = "v5.12"
DATE    = "2026-05-20"

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x6A)
ORANGE = RGBColor(0xF4, 0x7B, 0x20)
SLATE  = RGBColor(0x44, 0x55, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF2, 0xF5, 0xF9)


def _shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _style_table(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def _header_row(table, *labels):
    row = table.rows[0]
    for i, label in enumerate(labels):
        cell = row.cells[i]
        cell.text = label
        _shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9)


def _data_row(table, row_idx, *values, shade=False):
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        if shade:
            _shade_cell(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = SLATE


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else SLATE
    return p


def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color or SLATE
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = SLATE


def title_block(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Trading Agent")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("System Design Document")
    run2.font.size  = Pt(14)
    run2.font.color.rgb = ORANGE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"{VERSION}  ·  {DATE}  ·  Alpaca Paper Trading")
    run3.font.size  = Pt(9)
    run3.font.italic = True
    run3.font.color.rgb = SLATE

    doc.add_paragraph()


def section_overview(doc):
    add_heading(doc, "1. What It Is")
    add_body(doc,
        "An autonomous day-trading system that scans 600+ stocks and ETFs daily, "
        "selects high-probability intraday setups using Claude AI, manages risk "
        "through five independent guard layers, and executes bracket orders on "
        "Alpaca paper trading.")
    add_body(doc, "Daily objective: $500–$1,000 realized P&L.", bold=True)
    doc.add_paragraph()


def section_architecture(doc):
    add_heading(doc, "2. Architecture")

    add_code(doc,
        "cron-job.org  →  GitHub Actions  →  orchestrator.py\n"
        "                                         │\n"
        "                    ┌────────────────────┼────────────────────┐\n"
        "                    ▼                    ▼                    ▼\n"
        "                Scanner           Strategy Agent        Risk + Guards\n"
        "                ML Scorer         (Claude Sonnet)       Sector Guard\n"
        "                    │                    │              Guardrails\n"
        "                    └────────────────────┼────────────────────┘\n"
        "                                         ▼\n"
        "                              Portfolio Agent\n"
        "                            /              \\\n"
        "                    Alpaca Paper        Supabase DB\n"
        "                    (bracket orders)    (positions, P&L)\n"
        "                                             │\n"
        "                                    Streamlit Dashboard"
    )

    add_heading(doc, "Stack", level=2)
    items = [
        "Python 3.11 · Claude claude-sonnet-4-6 (Anthropic)",
        "Alpaca Markets API — paper trading, bracket orders, live quotes",
        "Supabase (PostgreSQL) — positions, P&L, scan results",
        "Streamlit Cloud — 5-tab dashboard (Summary, Today, Positions, Performance, Scan Log)",
        "GitHub Actions — cron scheduler (premarket · intraday · EOD)",
        "cron-job.org — external trigger for reliability (GitHub Actions cron can lag 5–15 min)",
    ]
    for item in items:
        add_bullet(doc, item)
    doc.add_paragraph()


def section_pipeline(doc):
    add_heading(doc, "3. Daily Pipeline")

    # Premarket
    add_heading(doc, "3.1  Premarket — 10:00 AM ET", level=2)
    add_body(doc,
        "Runs once before significant intraday volume develops. Delayed from 9:00 AM "
        "to avoid unreliable Alpaca ask quotes in the first 30 minutes after open "
        "(spreads can be 4–5% wide, causing immediate stop-outs on bracket fills).")
    doc.add_paragraph()

    table = doc.add_table(rows=11, cols=2)
    _style_table(table)
    _set_col_widths(table, [4.5, 12])
    _header_row(table, "Step", "What Happens")
    steps = [
        ("0. Market Context",
         "Fetches VIX, Fear & Greed, US futures, intl markets, economic calendar. "
         "Sets max_positions and quiet_day flag. Hard skip if futures < −1.5%."),
        ("1. Scan",
         "Scores 600+ tickers on RSI, MACD, Bollinger Bands, volume ratio, SMA20/50. "
         "Candidates with |score| ≥ 3 pass. Alpaca daily bars used as yfinance fallback."),
        ("1.5 News Filter",
         "Removes earnings-day tickers (blackout). Adds news sentiment to market summary."),
        ("1.75 Pre-filter",
         "Drops bearish candidates (score < 3) before the Claude call. Cuts input tokens ~60%."),
        ("1.76 ML Scoring",
         "XGBoost model predicts P(stock hits +2% intraday). Re-ranks candidates by probability."),
        ("1.8 Live Prices",
         "Refreshes entry prices from Alpaca ask quotes. Rejects if ask is >10% from scanner price."),
        ("1.85 VWAP Signals",
         "Enriches candidates: VWAP level, today's % change, relative strength vs SPY."),
        ("2. Strategy (Claude)",
         "Selects trades, assigns confidence (HIGH/MEDIUM/LOW), writes reasoning. "
         "Sets entry/target/stop using fixed formulas only."),
        ("3–3.75 Risk Guards",
         "Risk Agent (R:R, size bounds) → Sector Guard (≤3 per sector) → Guardrails (duplicates, price sanity, daily loss limit)."),
        ("4. Execute",
         "Opens two bracket orders per trade (partial profit design). Records positions in Supabase."),
    ]
    for i, (step, desc) in enumerate(steps):
        _data_row(table, i + 1, step, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()

    # Intraday
    add_heading(doc, "3.2  Intraday — Every 15 min, 10:30 AM–3:45 PM ET", level=2)
    for item in [
        "Reconcile: detect positions closed by Alpaca bracket (stop/target fired), record real exit price and P&L.",
        "Refresh: sync current price and unrealized P&L for open positions.",
        "Trailing stop: manual high-watermark trail — fires if price drops 1% from peak.",
        "Lock-in: Tier 1 ($716 realized) — let winners ride with tighter 0.5% trail. Tier 2 ($1,000 total) — close everything.",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # EOD
    add_heading(doc, "3.3  EOD — 4:30 PM ET", level=2)
    for item in [
        "Records daily performance (P&L, win rate, best/worst trade) to daily_performance table.",
        "Runs eval against 30-day rolling window.",
        "Generates daily summary for review.",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()


def section_trading_logic(doc):
    add_heading(doc, "4. Trading Logic")

    # Position sizing
    add_heading(doc, "4.1  Position Sizing", level=2)
    table = doc.add_table(rows=4, cols=3)
    _style_table(table)
    _set_col_widths(table, [3, 2.5, 11])
    _header_row(table, "Confidence", "Size", "Trigger Criteria")
    rows = [
        ("HIGH",   "$7,000", "Technical score ≥ 7  AND  volume ratio > 1.8  AND  3+ confirming signals"),
        ("MEDIUM", "$6,000", "Score 4–6  OR  (score 3–4  AND  above VWAP  AND  RS vs SPY ≥ 1.5)"),
        ("LOW",    "$5,000", "Score 3–4 with weak or absent VWAP/RS signals"),
    ]
    for i, (conf, size, trigger) in enumerate(rows):
        _data_row(table, i + 1, conf, size, trigger, shade=(i % 2 == 1))
    doc.add_paragraph()

    # Trade formulas
    add_heading(doc, "4.2  Trade Formulas", level=2)
    add_body(doc, "All values set by fixed formula — Claude cannot deviate:")
    add_code(doc,
        "entry_price    = Alpaca ask price (live) or scanner close\n"
        "target_price   = round(entry × 1.0200, 2)   # +2.0% full target\n"
        "partial_target = round(entry × 1.0100, 2)   # +1.0% partial exit\n"
        "stop_loss      = round(entry × 0.9933, 2)   # −0.67% stop\n"
        "shares         = int(position_size / entry)\n"
        "\n"
        "Reward:Risk  = 2.00% / 0.67% = 2.99 ≈ 3:1  (normal days)\n"
        "Reward:Risk  = 2.00% / 0.67% = 2.99 ≈ 3:1  (quiet day floor: 2.0)"
    )
    doc.add_paragraph()

    # Partial profit
    add_heading(doc, "4.3  Partial Profit Design", level=2)
    add_body(doc,
        "Each trade opens as two independent bracket orders with the same stop price:")
    add_code(doc,
        "Leg A  →  shares // 2   ·  target = +1%  (locks in early)\n"
        "Leg B  →  shares - A    ·  target = +2%  (rides the full move)\n"
        "Both   →  stop = entry × 0.9933"
    )
    add_body(doc,
        "Why: Converts all-or-nothing outcomes into graduated P&L. On quiet days where 2% "
        "moves are uncommon, Leg A frequently closes profitable while Leg B stops out — net "
        "positive vs. net zero under a single-bracket design.")
    doc.add_paragraph()

    # Trailing stop
    add_heading(doc, "4.4  Trailing Stop", level=2)
    add_body(doc,
        "Manual high-watermark trail checked every 15 min (Alpaca native trailing stops "
        "are not used — not supported in bracket order legs):")
    add_code(doc,
        "effective_stop = max(stop_loss, high_watermark × (1 − 1.0%))\n"
        "\n"
        "After Tier 1 lock-in ($716 realized):\n"
        "effective_stop = max(stop_loss, high_watermark × (1 − 0.5%))")
    doc.add_paragraph()


def section_risk(doc):
    add_heading(doc, "5. Risk Controls")
    add_body(doc, "Five independent layers applied in sequence — any one can block a trade:")
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    _style_table(table)
    _set_col_widths(table, [4.5, 12])
    _header_row(table, "Layer", "What It Blocks")
    guards = [
        ("Market Context",  "Trading when futures are down >1.5% (crash day skip)"),
        ("News Filter",     "Earnings-day tickers, negative catalyst stocks"),
        ("Risk Agent",      "R:R below floor, position size out of bounds, stop loss too wide (>0.67%)"),
        ("Sector Guard",    "More than 3 positions in any single sector"),
        ("Guardrails",      "Duplicate tickers (already traded today), price sanity >5% from market, daily loss limit −$300 breached"),
    ]
    for i, (layer, desc) in enumerate(guards):
        _data_row(table, i + 1, layer, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()

    add_heading(doc, "5.1  Quiet Day Mode", level=2)
    add_body(doc,
        "Triggered automatically when Fear & Greed Index < 35. Market sentiment is "
        "in Fear territory — setups are weaker but contrarian opportunities exist.")
    add_code(doc,
        "quiet_day = fear_greed_value < 35\n"
        "\n"
        "When quiet_day=True:\n"
        "  MIN_REWARD_RISK  →  2.0  (from 3.0)\n"
        "  MEDIUM confidence  →  also qualifies: score 3-4 + above VWAP + RS ≥ 1.5"
    )
    add_body(doc,
        "Rationale: At 80%+ win rate (validation target), 2:1 R:R has strongly positive "
        "expected value (EV = 0.8×2 − 0.2×1 = +1.4). The 3.0 floor was correctly "
        "rejecting borderline trades on strong days but over-filtering on quiet days "
        "where near-miss trades (R:R 2.9x) are the norm, not the exception.")
    doc.add_paragraph()


def section_config(doc):
    add_heading(doc, "6. Key Configuration")

    params = [
        ("TOTAL_CAPITAL",              "$100,000",  "Simulated account size"),
        ("TARGET_PCT",                 "2.0%",      "Full profit target per trade"),
        ("PARTIAL_PROFIT_PCT",         "1.0%",      "Partial exit target (Leg A)"),
        ("MAX_LOSS_PER_TRADE",         "0.67%",     "Stop loss depth"),
        ("MIN_REWARD_RISK",            "3.0",       "Normal day R:R floor"),
        ("QUIET_DAY_MIN_REWARD_RISK",  "2.0",       "Quiet day R:R floor (F&G < 35)"),
        ("QUIET_DAY_FG_THRESHOLD",     "35",        "Fear & Greed threshold for quiet day mode"),
        ("TRAIL_PCT",                  "1.0%",      "Trailing stop from high watermark"),
        ("LOCK_IN_TRAIL_PCT",          "0.5%",      "Tighter trail after Tier 1 lock-in"),
        ("DAILY_LOCK_IN_TARGET",       "$716",      "Tier 1: realized P&L — let winners ride"),
        ("DAILY_BONUS_TARGET",         "$1,000",    "Tier 2: realized+unrealized — close everything"),
        ("DAILY_LOSS_LIMIT",           "−$300",     "Stop trading for the day"),
        ("MAX_POSITIONS",              "15",        "Max concurrent positions"),
        ("MAX_PER_SECTOR",             "3",         "Sector concentration cap"),
        ("MIN_AVG_VOLUME",             "1,000,000", "Liquidity floor (shares/day)"),
        ("SCORE_THRESHOLD",            "3",         "Min scanner score |score| ≥ 3"),
        ("STRATEGY_MIN_SCORE",         "3",         "Pre-filter before Claude call (bullish only)"),
    ]

    table = doc.add_table(rows=len(params) + 1, cols=3)
    _style_table(table)
    _set_col_widths(table, [5.5, 2.5, 8.5])
    _header_row(table, "Parameter", "Value", "Purpose")
    for i, (param, val, desc) in enumerate(params):
        _data_row(table, i + 1, param, val, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
    doc.add_paragraph()


def section_validation(doc):
    add_heading(doc, "7. Validation Gate")
    add_body(doc, "Run on June 1, 2026:  python3 eval.py --days 14", bold=True)
    add_body(doc, "Do not deploy real capital until all criteria pass.")
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    _style_table(table)
    _set_col_widths(table, [6, 10.5])
    _header_row(table, "Criterion", "Threshold")
    criteria = [
        ("Win rate",              "≥ 80%"),
        ("Avg daily P&L",         "≥ $500"),
        ("Double-sell events",    "None — bracket reconciliation confirmed clean"),
        ("Integrity flags",       "None — no UNFILLED/CLEANUP anomalies"),
        ("Confidence score",      "≥ 7/10"),
    ]
    for i, (crit, threshold) in enumerate(criteria):
        _data_row(table, i + 1, crit, threshold, shade=(i % 2 == 1))
    doc.add_paragraph()


def section_roadmap(doc):
    add_heading(doc, "8. Post-June-1 Roadmap")

    table = doc.add_table(rows=3, cols=3)
    _style_table(table)
    _set_col_widths(table, [2, 4.5, 10])
    _header_row(table, "Priority", "Feature", "Description")
    items = [
        ("P0", "P&L Reconciliation",
         "Pull Alpaca account.equity at EOD as source of truth. Add commission/spread/slippage breakdown."),
        ("P1", "Intraday Trade Entries",
         "Second scan at 12:30 PM. Momentum-only. Guardrailed by open position count and daily P&L sign."),
    ]
    for i, (pri, feat, desc) in enumerate(items):
        _data_row(table, i + 1, pri, feat, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = ORANGE
    doc.add_paragraph()


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    title_block(doc)
    section_overview(doc)
    section_architecture(doc)
    section_pipeline(doc)
    section_trading_logic(doc)
    section_risk(doc)
    section_config(doc)
    section_validation(doc)
    section_roadmap(doc)

    out = "Trading_Agent_Design.docx"
    doc.save(out)
    print(f"✅  Saved: {out}")


if __name__ == "__main__":
    main()
