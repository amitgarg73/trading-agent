"""
Generates Trading_Agent_Vision.docx — where the system could go, what the
competitive landscape looks like, and what Mode 2 (event-driven) would be.

Run: python3 generate_vision.py
Update this document when the product thesis or architecture evolves materially.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

NAVY   = RGBColor(0x1A, 0x3A, 0x6A)
ORANGE = RGBColor(0xF4, 0x7B, 0x20)
GREEN  = RGBColor(0x1A, 0x7A, 0x3A)
GRAY   = RGBColor(0x44, 0x44, 0x44)
PURPLE = RGBColor(0x6A, 0x1A, 0x6A)


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = NAVY
    return p


def subheading(text, color=ORANGE):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()


def divider():
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading('AI Trading Agent — Vision & Where Next', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = NAVY

meta = doc.add_paragraph('Amit Garg  ·  May 2026  ·  Living document — update when product thesis changes')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(11)
meta.runs[0].font.color.rgb = GRAY

doc.add_paragraph()
body(
    'This document covers three things: (1) where the system fits in the current market '
    'landscape and what gap exists, (2) what Mode 2 (event-driven trading) looks like in '
    'full design, and (3) three scenarios for where this goes as a product. Read alongside '
    'Trading_Agent_Brutal_Assessment.docx — the assessment explains the problem, '
    'this document charts the path forward.'
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 1: COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 1: The Market Landscape — A Gap Nobody Owns')

body(
    'The retail algorithmic trading market has two distinct camps. '
    'Neither of them is doing what a properly-built version of this system does.'
)

divider()
subheading('Camp 1 — No-Code Strategy Runners')
add_table(
    ['Product', 'Target User', 'AI Layer', 'Signal Source', 'Price'],
    [
        ('Composer',      'Non-technical traders', 'None',         'Rule-based strategies', 'Free tier (1 live strategy)'),
        ('Streak',        'Visual builders',       'None',         'Rule-based',            'Free (Alpaca-integrated)'),
        ('Autopilot',     'Copy traders',          'None',         'Expert portfolios',     'Subscription'),
        ('Magnifi',       'Retail investors',      'GPT wrapper',  'Fund selection only',   'Paid'),
    ]
)
body(
    'These platforms let users run strategies without code. None have an LLM reasoning '
    'layer. None synthesize multiple signal classes. They are strategy executors, '
    'not signal generators.'
)

divider()
subheading('Camp 2 — Developer Platforms')
add_table(
    ['Product',        'Target User',     'AI Layer',        'Signal Source',     'Price'],
    [
        ('QuantConnect',  'Quant developers', 'None',           'Custom Python',     'Free / $20/mo+'),
        ('Alpaca',        'Developers',       'None (MCP beta)', 'Custom via API',    'Free (commission-free)'),
        ('Interactive Brokers', 'Active traders', 'None',       'Custom via API',    'Commissions'),
    ]
)
body(
    'QuantConnect + Alpaca is the institutional-grade free stack for developers. '
    'Powerful backtest infrastructure, no LLM reasoning. Alpaca launched an official '
    'MCP server (Model Context Protocol) in early 2026 — Claude can now place trades '
    'via natural language. But nobody has wired it to real alternative data signals.'
)

divider()
subheading('Camp 3 — LLM Trading Experiments (Emerging)')
add_table(
    ['Project',        'Signal Source',               'Stage',              'Gap'],
    [
        ('TradingAgents',   'Technical + fundamental',    'Research/paper',     'No real alternative data'),
        ('MAHORAGA',        'Social sentiment',           'Open source',        'Single signal, no synthesis'),
        ('OpenClaw',        'Multi-agent consensus',      'Paper trading',      'Still technical signals'),
        ('Alpaca MCP',      'Natural language → orders',  'Early commercial',   'No signal layer at all'),
        ('This system',     'Technical (v5.6)',           'Paper, 2-week gate', 'Signal is weak — see Assessment doc'),
    ]
)
body(
    'The LLM trading space is in early formation. Everyone is using LLMs to make decisions '
    'on technical signals or natural language commands. Nobody has built the combination: '
    'alternative data signals (PEAD + insider buys + unusual options) synthesized by an LLM '
    'into high-conviction trade plans. That gap is real and unoccupied.'
)

divider()
subheading('The Gap We Can Own', GREEN)
body(
    'The defensible position is not "we run technical strategies with an LLM." '
    'Every competitor can copy that in a weekend. The defensible position is:'
)
bullet('Multiple genuine signal classes (earnings surprise, insider conviction, options flow) — not technicals')
bullet('LLM as synthesizer across those signals — not LLM as signal source')
bullet('Full auditability of every trade decision (eval, exit mechanism tracking, Agent Scorecard)')
bullet('Operational infrastructure that competitors don\'t have (health checks, guardrails, integrity checks)')
body(
    'The last two we already have. The first two are what we need to build. '
    'That combination — real signals + LLM synthesis + operational rigor — is genuinely '
    'differentiated from everything currently in the market.'
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 2: MODE 2 — EVENT-DRIVEN ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 2: Mode 2 — Event-Driven Trading Design')

body(
    'Mode 1 (current) is intraday momentum: scan → Claude picks → bracket order → close same day. '
    'Mode 2 is event-driven: a catalyst occurs → signals confirm → Claude synthesizes → '
    'hold 3–20 days. Different signal class, different hold period, different exit rules. '
    'This is where the documented edge lives.'
)

divider()
subheading('Signal 1 — Post-Earnings Announcement Drift (PEAD)')
body('Academic evidence: 40+ years of papers. One of the most robust documented anomalies in finance.')
body('Why it persists: institutions cannot fully reprice in one session; retail herds in over days.')

add_table(
    ['Parameter', 'Rule'],
    [
        ('Entry timing',    'Day 1 close OR Day 2 open — never earnings day (gap-fill risk, extreme volatility)'),
        ('Entry filter',    'Gap must hold through Day 1 session — if filled, skip'),
        ('Surprise threshold', 'EPS beat ≥ 5% above consensus estimate (larger surprise = stronger drift)'),
        ('Stop loss',       'Below Day 1 close — gap fill = thesis broken, exit immediately'),
        ('Time stop',       'Exit after 20 trading days if no drift, or on reversal of earnings-day move'),
        ('Target',          'No fixed target — trail stop ratchets as drift continues'),
        ('Position size',   '1–2% portfolio risk per trade; max 10% total in PEAD positions simultaneously'),
        ('Hold period',     '5–20 trading days (not intraday)'),
    ]
)

body('Data sources (all free):')
bullet('EPS actual vs. estimate: yfinance earnings data (quarterly)')
bullet('Earnings calendar: yfinance calendar, Earnings Whispers (free tier)')
bullet('Surprise magnitude: (actual EPS - estimate EPS) / abs(estimate EPS) × 100')

body('Implementation path:')
bullet('agents/earnings_scanner.py — queries next 5 trading days earnings calendar')
bullet('filters for tickers already in our universe with significant beat (≥5% positive surprise)')
bullet('orchestrator mode=PEAD runs next morning after earnings announcement')
bullet('Claude synthesizes: earnings beat + technical context + any confirming signals → go/no-go')

divider()
subheading('Signal 2 — Insider Buying (SEC Form 4)')
body('Academic evidence: 5–8% abnormal return over 6 months. Strongest when multiple insiders buy.')
body('Only counts: open market purchases by CEO/CFO/COO/Director. Not grants, not options exercise.')

add_table(
    ['Parameter', 'Rule'],
    [
        ('Filing source',     'SEC EDGAR Form 4 — free, filed within 2 business days of transaction'),
        ('Transaction filter', 'Transaction code "P" only (open market purchase) — exclude grants/exercises'),
        ('Role filter',       'CEO, CFO, COO, President, Director — exclude VP and below'),
        ('Size filter',       '> $50K purchase amount — filters symbolic buys'),
        ('Cluster signal',    '2+ insiders at same company buying within 30-day window = strong signal'),
        ('Hold period',       '20–60 trading days (longer than PEAD — fundamental conviction play)'),
        ('Entry',             'Within 5 trading days of filing date'),
        ('Stop',              '8–10% below entry — wider than intraday (longer hold = more room)'),
    ]
)

body('Free data sources:')
bullet('Securities Database REST API — free, no API key: securitiesdb.com/developers/insider-trading-api')
bullet('OpenInsider screener — free real-time: openinsider.com')
bullet('SEC EDGAR direct — free XML API: efts.sec.gov/LATEST/search-index?q=%22form+4%22')

body('Use as confirming filter: insider buying on a PEAD ticker = highest conviction setup.')

divider()
subheading('Signal 3 — Unusual Options Activity')
body('Evidence: informed options volume predicts stock returns 1–5 days ahead.')
body(
    'Informed traders use options (cheap leverage) before expected moves. '
    'Legally observable. Not insider trading to monitor it.'
)

add_table(
    ['Parameter', 'Rule'],
    [
        ('Volume threshold',  'Call volume ≥ 5x 30-day average on that strike/expiry'),
        ('Strike filter',     'Out-of-money calls (OTM) — not protective puts or covered calls'),
        ('Expiry filter',     '1–4 weeks out — short-dated = conviction, not hedging'),
        ('Bid-side pressure', 'Calls bought on ask (aggressor) — not passive limit orders'),
        ('Use case',          'Confirming signal — trade stock, not the options'),
        ('Hold period',       '1–5 trading days (short-term catalyst signal)'),
    ]
)

body('Free data sources:')
bullet('Unusual Whales — free tier: unusualwhales.com')
bullet('Barchart Unusual Options Activity — free: barchart.com/options/unusual-activity')
bullet('Market Chameleon — free tier: marketchameleon.com')

divider()
subheading('Claude\'s Role in Mode 2', GREEN)
body(
    'In Mode 2, Claude is NOT the signal source. Claude is the final synthesizer '
    'that decides whether a confluence of signals justifies a trade and at what sizing.'
)

body('Example Claude prompt for Mode 2:')
body(
    '"NVDA reported earnings yesterday. EPS beat by 12% (actual $0.89, estimate $0.79). '
    'The stock gapped up 6% and held through the session. Unusual call volume was 8x normal '
    'at the $140 strike 2 weeks out. Two directors filed Form 4 open market purchases '
    'totaling $380K last week. Current price is $138. Our trailing stop is 1%. '
    'Assess conviction level (HIGH/MEDIUM/LOW), recommended entry, target, and stop '
    'for a 10-day PEAD hold. Consider: is this a real drift setup or a gap-and-fade?"'
)
body('This is Claude doing what it is actually good at: synthesizing structured evidence into a reasoned trade decision.')

divider()
subheading('Mode 2 Architecture')
add_table(
    ['Step', 'Agent', 'What It Does'],
    [
        ('0',   'market_context.py',   'VIX + futures gate (same as Mode 1)'),
        ('1',   'earnings_scanner.py', 'NEW: find tickers with significant earnings beat in last 2 days'),
        ('2',   'insider_feed.py',     'NEW: check Securities DB API for cluster insider buys in universe'),
        ('3',   'options_watcher.py',  'NEW: check Unusual Whales for unusual call activity in universe'),
        ('4',   'strategy.py (Mode 2)','Claude synthesizes all 3 signals → ranked trade plans with hold period'),
        ('5',   'risk.py',             'Validates sizing (wider stops for multi-day holds)'),
        ('6',   'guardrails.py',       'Existing 6 checks — unchanged'),
        ('7',   'alpaca_broker.py',    'Submit bracket orders (wider target/stop for multi-day hold)'),
        ('EOD', 'performance.py',      'Tracks Mode 1 vs Mode 2 P&L separately in daily_performance'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# PART 3: THREE SCENARIOS FOR WHERE THIS GOES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 3: Three Scenarios for Where This Goes')

body(
    'The system is currently a personal paper trading research tool. '
    'Three distinct paths forward — each with different ambition, complexity, and risk.'
)

divider()
subheading('Scenario A — Personal Alpha Tool (Lowest Complexity)')
body('Keep it exactly as is. Build Mode 2. Validate with real Sharpe tracking. Deploy real capital when confidence is ≥8/10.')
add_table(
    ['Dimension', 'Detail'],
    [
        ('Capital', '$100K paper → real money → compound to $500K over 3–5 years'),
        ('Target',  'Sharpe ≥ 1.5 sustained over 12 months of live trading'),
        ('Regulatory', 'None — trading your own account'),
        ('Complexity', 'Low — just keep building what we have'),
        ('Upside',  'Personal wealth compounding; validated trading system'),
        ('Risk',    'Never prove the edge; plateau at paper trading indefinitely'),
        ('When to choose', 'If the June 2026 eval shows strong performance; path of least resistance'),
    ]
)

divider()
subheading('Scenario B — Small Fund / LP Structure (Medium Complexity)')
body(
    'Once the system has 12+ months of live audited returns, raise a small pool from '
    'friends/family under the SEC\'s exemption thresholds. Below $150M AUM, you do not '
    'need to register as an investment adviser if limited to accredited investors.'
)
add_table(
    ['Dimension', 'Detail'],
    [
        ('Capital',     '$100K own → $500K–$2M LP pool from accredited investors'),
        ('Target',      'Audited returns ≥ 20% annualized net — credible, not 175%'),
        ('Regulatory',  'Exempt reporting adviser (ERA) under Dodd-Frank — below $150M AUM threshold'),
        ('Complexity',  'Medium — legal structure, audited books, LP reporting required'),
        ('Upside',      'Carry on profits (typically 20%); real AUM = real business'),
        ('Risk',        'Regulatory exposure; losing other people\'s money is a different game'),
        ('When to choose', 'If Sharpe ≥ 1.5 with 200+ live trades and 12+ months of audited history'),
    ]
)

divider()
subheading('Scenario C — SaaS Signal Product (Highest Upside, Different Business)', PURPLE)
body(
    'The LLM-as-synthesizer over real signals architecture is genuinely differentiated from '
    'anything currently in the market. Build a subscription product that delivers daily '
    'trade research — not execution. Pure signal and reasoning product.'
)
body(
    'Key insight: no broker-dealer registration needed if you provide research, not execution. '
    'Investment adviser registration required above $100M or 15 clients in some states — '
    'but a SaaS product with disclaimers and no personalized advice is a different category.'
)
add_table(
    ['Dimension', 'Detail'],
    [
        ('Product',     'Daily briefing: top 3–5 PEAD/insider/options setups with Claude\'s full reasoning'),
        ('Format',      'Web app or email digest — "Here are today\'s highest-conviction setups and why"'),
        ('Pricing',     '$49–$99/month per user; 1,000 users = $600K–$1.2M ARR'),
        ('Regulatory',  'Research/education product — not personalized advice; add standard disclaimers'),
        ('Moat',        'Signal quality + Claude reasoning + audit trail — nobody else has this combination'),
        ('Complexity',  'High — product, distribution, customer support, legal review'),
        ('Upside',      'Scalable SaaS revenue; plays directly to PM background'),
        ('When to choose', 'After Mode 2 is validated on paper; if the signal quality is demonstrably better than competitors'),
    ]
)

body('The Scenario C product stack would look like:')
bullet('Signal engine: earnings_scanner + insider_feed + options_watcher (same as Mode 2)')
bullet('Reasoning layer: Claude synthesizes → structured trade brief (ticker, catalyst, conviction, entry zone, risk)')
bullet('Delivery: Streamlit web app (already have) → email digest → eventually API')
bullet('Differentiation from existing screeners: Claude\'s reasoning is auditable — you can read WHY each setup qualifies')


# ══════════════════════════════════════════════════════════════════════════════
# PART 4: RECOMMENDED PATH
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('Part 4: Recommended Path Forward')

body(
    'The scenarios are not mutually exclusive. A builds B builds C. '
    'The right move is to validate the signal before deciding which scenario to pursue.'
)

add_table(
    ['Phase', 'Timeline', 'Goal', 'Decision Gate'],
    [
        ('Phase 0: Validate Mode 1',      'Now → Jun 1 2026',    '2-week paper gate: native trailing stop, win rate ≥80%',            'Pass → Phase 1. Fail → fix Mode 1 first.'),
        ('Phase 1: Build Mode 2 signals', 'Jun–Aug 2026',        'earnings_scanner + insider_feed + options_watcher on paper',         'Run 60 trading days; track Sharpe separately from Mode 1'),
        ('Phase 2: Validate Mode 2',      'Aug–Nov 2026',        '200+ Mode 2 trades; Sharpe ≥1.0; real slippage modeled',            'Pass → real money. Fail → rethink signal architecture.'),
        ('Phase 3: Real money (small)',   'Late 2026',           '$10K–$25K real capital; validate live fills match paper',           'Pass → scale. Fail → size down and diagnose.'),
        ('Phase 4: Choose scenario',      '2027',                'A (compound), B (LP), or C (SaaS) based on validated returns',      'Scenario choice depends on Sharpe, personal goals, appetite'),
    ]
)

divider()
subheading('Immediate Next Sprint (Post-June 1)', GREEN)
body('In priority order — each one is buildable in a single session:')
bullet('1. 15-min momentum confirmation (2f) — quick win on Mode 1, low effort')
bullet('2. insider_feed.py — Securities DB API, filter for cluster buys in universe, free data')
bullet('3. earnings_scanner.py — yfinance earnings calendar + EPS surprise calculation')
bullet('4. options_watcher.py — Unusual Whales free API integration')
bullet('5. Mode 2 orchestrator — chain 1–3 into a separate daily run (premarket, different mode flag)')
bullet('6. Retire dollar target; replace with Sharpe ratio tracking in eval.py')

divider()
subheading('What Would Change the Scenario Choice')
add_table(
    ['Metric', 'Scenario A', 'Scenario B', 'Scenario C'],
    [
        ('Sharpe ratio (live)',   '≥ 1.0',          '≥ 1.5 over 12 months', '≥ 1.5 over 12 months'),
        ('Trade count (live)',    '200+',            '500+',                 '500+'),
        ('Max drawdown',          '< 15%',           '< 10%',               '< 10%'),
        ('Signal quality proof',  'Paper sufficient', 'Audited live required', 'Audited live required'),
        ('Personal time budget',  'Low',             'Medium (investor relations)', 'High (product + customers)'),
    ]
)


# ── Save ──────────────────────────────────────────────────────────────────────
_project_dir = "/Users/amitgarg/Claude Projects/trading-agent"
out = os.path.join(_project_dir, "Trading_Agent_Vision.docx")
doc.save(out)
print(f"Saved: {out}")
