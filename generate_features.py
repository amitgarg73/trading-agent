"""Generates the Trading Agent Feature Roadmap as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AI Trading Agent — Feature Roadmap', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)

sub = doc.add_paragraph('Potential Features, Ideas & Enhancements')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph('Amit Garg  ·  May 2026  ·  v1.0')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── Legend ────────────────────────────────────────────────────────────────────
heading('How to Read This Doc')
body(
    'Features are grouped by category. Each has an effort estimate, priority, and status. '
    'Effort and priority are rough — revisit before committing to any item. '
    'Status: SHIPPED = live in production. PLANNED = next iteration. IDEA = worth exploring. BACKLOG = lower priority.'
)
add_table(
    ['Label', 'Meaning'],
    [
        ('Effort XS', '< 1 hour — config change or 10-line tweak'),
        ('Effort S',  '1–4 hours — single file change, well-understood'),
        ('Effort M',  '4–16 hours — multi-file, some design needed'),
        ('Effort L',  '16+ hours — new agent, schema change, external integration'),
        ('P0', 'Critical — affects core correctness or safety'),
        ('P1', 'High value — direct P&L or reliability impact'),
        ('P2', 'Medium value — UX, observability, operational efficiency'),
        ('P3', 'Nice-to-have — future state, low urgency'),
    ]
)

# ── 1. Signal Quality ─────────────────────────────────────────────────────────
doc.add_page_break()
heading('1. Signal Quality — Better Inputs for Claude')
body(
    'These features improve the quality of candidates and context passed to the strategy agent. '
    'Better signals = higher win rate without changing position sizing or risk rules.'
)

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('Sector rotation scoring (V2e)',
         'Score sectors by relative strength this week vs. S&P 500; pass sector bias to Claude',
         'Claude currently picks sectors randomly; momentum sectors outperform in trending markets',
         'M', 'P1', 'PLANNED'),

        ('Momentum confirmation — 15-min rule (V2f)',
         'Delay entry until 9:45 AM; skip candidates that reversed in the first 15 min',
         'First-15-min noise causes many false signals; waiting for confirmation improves entry quality',
         'M', 'P1', 'PLANNED'),

        ('Pre-market gap filter',
         'Exclude candidates already up >2% pre-market (target nearly hit before entry)',
         'Entering a stock already at +2% gives only 1% upside to the 3% target — asymmetric risk',
         'S', 'P1', 'IDEA'),

        ('Pre-market volume spike',
         'Flag candidates with pre-market volume >3x normal as high-momentum candidates',
         'Pre-market volume surge often predicts continuation; currently scanner only uses regular hours volume',
         'S', 'P2', 'IDEA'),

        ('RSI divergence detection',
         'Flag when price makes new high but RSI does not (bearish divergence) or vice versa',
         'Divergence is a leading reversal signal — reduces entries on exhausted moves',
         'M', 'P2', 'IDEA'),

        ('Multi-timeframe RSI alignment',
         'Only enter if RSI is bullish on both daily AND 15-min timeframe',
         'Daily and intraday alignment filters out counter-trend noise; higher-quality setups',
         'M', 'P2', 'IDEA'),

        ('VWAP + Relative Strength enrichment vs. SPY (Thread 1)',
         'Step 1.85: StockSnapshotRequest fetches above_vwap, vwap, today_pct_change, rs_vs_spy '
         'per candidate; re-sorts above-VWAP candidates first before Claude call; '
         'Claude SYSTEM prompt explains interpretation; eval.py validates signal quality',
         'Above-VWAP = institutional buying pressure confirmed; RS vs SPY identifies market leaders. '
         'Candidates sorted so Claude sees live confirmed setups first — not just overnight technicals',
         'S', 'P1', 'SHIPPED'),

        ('Relative strength vs. sector index',
         'Score each candidate by performance relative to its sector ETF (e.g., NVDA vs. SOXX)',
         'Sector leaders outperform laggards; relative strength is a proven momentum factor',
         'S', 'P2', 'IDEA'),

        ('Options flow — put/call ratio',
         'Pull put/call ratio for each candidate; flag unusual options activity as a signal',
         'Institutional positioning shows up in options before price; smart money signal',
         'L', 'P3', 'BACKLOG'),

        ('Earnings momentum plays',
         'Target stocks 2–5 days after a positive earnings surprise, post-blackout period',
         'Post-earnings momentum is well-documented; currently we skip earnings entirely',
         'M', 'P2', 'IDEA'),

        ('Short interest data',
         'Integrate Finviz or similar free source for short interest %; flag high-short tickers',
         'High short interest + price strength = squeeze potential; asymmetric upside',
         'M', 'P3', 'BACKLOG'),
    ]
)

# ── 2. Position Management ────────────────────────────────────────────────────
heading('2. Position Management — Better Exits and Sizing')
body('These features control how positions are sized, managed, and exited once open.')

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('Trailing stops',
         'effective_stop = max(original_stop, high_watermark × 0.99); ratchets as stock rises',
         'Prevents giving back gains on reversals; shipped in v5.3',
         'S', 'P1', 'SHIPPED'),

        ('Confidence-weighted sizing',
         'HIGH=$7K, MEDIUM=$6K, LOW=$5K per trade; risk agent enforces',
         'Higher conviction = more capital; improves P&L on best setups',
         'XS', 'P1', 'SHIPPED'),

        ('Partial profit taking',
         'Close 50% of position at +1.5%, let remaining 50% run to +3% target',
         'Locks in partial gain on strong movers while keeping upside exposure; reduces STOP losses',
         'M', 'P1', 'IDEA'),

        ('Scale-in entries',
         'Enter 50% at open, add 50% more if price pulls back 0.5% within first 30 min',
         'Better average entry price; confirms momentum before full deployment',
         'M', 'P2', 'IDEA'),

        ('Time-based stops',
         'Close any position flat (no P&L) if it has not moved ±1% after 2 hours',
         'Stuck positions tie up capital and tend not to recover; free capital for better setups',
         'S', 'P2', 'IDEA'),

        ('Volatility-adjusted sizing',
         'Smaller position size for high-ATR tickers (ATR >5% → reduce to $5K max)',
         'High-volatility names have wider swings; same $5K position has bigger dollar risk on volatile names',
         'S', 'P2', 'IDEA'),

        ('Correlation cap across open positions',
         'Block a trade if it is highly correlated with 2+ already-open positions (e.g., 5 semiconductors)',
         'Sector guard handles the cap but not intra-sector correlation; prevents concentrated exposure',
         'M', 'P2', 'IDEA'),

        ('Tiered lock-in — let winners ride (Thread 2)',
         'Tier 1 ($716 realized): stop closing positions, tighten trail to 0.5%; '
         'Tier 2 ($1,000 realized+unrealized): close everything, protect exceptional day. '
         'DAILY_LOCK_IN_TARGET, DAILY_BONUS_TARGET, LOCK_IN_TRAIL_PCT settings; '
         'effective_trail in portfolio.py; tiered logic in intraday.py; '
         'tailwind banners + VWAP badges in dashboard; tailwind analysis in eval.py',
         'Old system closed all positions at $716 even when winners were still running. '
         'On strong momentum days, the top trades are still moving at Tier 1 — '
         'tiered logic captures the additional $284+ toward $1,000 without adding risk',
         'S', 'P1', 'SHIPPED'),

        ('Raise DAILY_LOCK_IN_TARGET as capital compounds',
         'Auto-scale DAILY_LOCK_IN_TARGET proportionally as ending_capital grows past thresholds',
         'Currently fixed at $716; as capital grows to $140K the same % profit = $1K — threshold should grow too',
         'XS', 'P1', 'IDEA'),
    ]
)

# ── 3. Strategy Expansion ─────────────────────────────────────────────────────
doc.add_page_break()
heading('3. Strategy Expansion — New Trade Types')
body('These features add entirely new strategy modes beyond the current long-only intraday setup.')

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('Short selling / mean reversion',
         'RSI overbought (>70) + price extended above Bollinger upper band → short candidate; update guardrails to allow SELL_SHORT',
         'Doubles opportunity set; short side works on down days when longs struggle; captures both directions',
         'M', 'P2', 'IDEA'),

        ('Pairs / long-short trades',
         'Long sector leader + short sector laggard in same trade; delta-neutral within sector',
         'Market-neutral; profits from relative performance, not direction; lower drawdown days',
         'L', 'P3', 'BACKLOG'),

        ('Swing trading mode',
         'Hold 2–3 days; larger targets (8–10%); separate logic for overnight risk management',
         'Intraday targets of 3% are often hit but more could be captured on strong trends',
         'L', 'P2', 'IDEA'),

        ('Gap-and-go strategy',
         'Target stocks gapping up >2% at open with volume confirmation; tighter stop, same target',
         'Gap-and-go is a proven intraday strategy; high hit rate in first 30 min on strong movers',
         'M', 'P2', 'IDEA'),

        ('ETF rotation strategy',
         'Rotate capital into top-performing sector ETFs weekly; use V2e sector scores to pick',
         'ETFs have lower volatility and no earnings risk; complementary to individual stock picks',
         'M', 'P3', 'BACKLOG'),

        ('A/B strategy testing',
         'Run two strategy configs simultaneously in simulation; compare weekly',
         'Currently all config changes are sequential; A/B allows parallel validation before committing',
         'L', 'P2', 'IDEA'),
    ]
)

# ── 4. Execution Quality ──────────────────────────────────────────────────────
heading('4. Execution Quality — Better Fills')
body('These features improve how orders are entered and filled in Alpaca.')

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('Time-of-day entry filter (skip first 15 min)',
         'Skip entries in first 15 min (9:30–9:45 AM ET); wide spreads and false signals in open auction',
         'Most false signals and stop-outs happen in the opening 15 min; waiting improves entry quality',
         'S', 'P1', 'SHIPPED'),

        ('Limit order entries',
         'Submit limit buy at entry_price × 1.001 instead of market order',
         'Market orders on small/mid caps can have 0.2–0.5% slippage; limits improve average entry price',
         'M', 'P1', 'SHIPPED'),

        ('Pre-market paper simulation',
         'Run a dry simulation at 8:50 AM before 9:00 AM premarket to validate pipeline without opening positions',
         'Useful during testing phases; verifies Claude picks and risk validation before live submission',
         'S', 'P3', 'BACKLOG'),

        ('Native Alpaca trailing stop',
         'StopLossRequest(trail_percent=TRAIL_PCT*100) in bracket order stop-loss leg. '
         'USE_NATIVE_TRAILING_STOP feature flag (enabled 2026-05-18). '
         'native_trail_active boolean per position. exit_mechanism column tracks NATIVE_TRAIL exits. '
         '2-week paper validation gate running — closes 2026-06-01.',
         'L', 'P0', 'SHIPPED'),
    ]
)

# ── 5. Monitoring & Alerts ────────────────────────────────────────────────────
doc.add_page_break()
heading('5. Monitoring & Alerts — Know What\'s Happening')
body('These features surface real-time events without requiring the user to check the dashboard.')

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('Per-trade email/SMS on close',
         'Send email when any position closes — ticker, reason (TARGET/STOP/TRAIL), realized P&L',
         'Biggest quality-of-life gap; currently must check dashboard to know if a target was hit',
         'S', 'P1', 'PLANNED'),

        ('Daily 4:30 PM summary email',
         'Auto-send EOD P&L summary — total P&L, winners, losers, lock-in status, win rate',
         'Replaces the manual dashboard check at end of day; one email = full picture',
         'S', 'P1', 'IDEA'),

        ('Mid-day drawdown alert',
         'Send email if realized P&L drops below -$150 intraday (before -$300 DAILY_LOSS_LIMIT fires)',
         'Early warning before the safety net — gives option to manually intervene if needed',
         'S', 'P2', 'IDEA'),

        ('Weekly performance email',
         'Auto-generated weekly summary every Friday EOD — grade, avg P&L, top/bottom trades, recommendations',
         'Replaces manual eval.py run; push delivery means nothing falls through the cracks',
         'M', 'P2', 'IDEA'),

        ('Anthropic credit balance alert',
         'Check API credit balance weekly; email if < $5 remaining',
         'No built-in Anthropic alert exists; currently requires manual console check',
         'S', 'P2', 'IDEA'),

        ('Slack / Discord integration',
         'Post position opens, closes, and daily summary to a Slack or Discord channel',
         'Team visibility if sharing with others; real-time feed without email overhead',
         'M', 'P3', 'BACKLOG'),

        ('Portfolio value milestone alert',
         'Email when portfolio crosses $110K, $125K, $150K etc.',
         'Motivation and checkpoint trigger to review/tune settings as capital compounds',
         'XS', 'P3', 'BACKLOG'),
    ]
)

# ── 6. Infrastructure & Reliability ──────────────────────────────────────────
heading('6. Infrastructure & Reliability')
body('These features make the system more robust and self-correcting.')

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('GitHub Actions retry (3x, 60s backoff)',
         'trading.yml retries up to 3 times before marking a run failed',
         'Handles transient yfinance/API failures; shipped in v5.3',
         'XS', 'P0', 'SHIPPED'),

        ('Automated capital compounding',
         'After each trading day, use ending_capital from daily_performance to update TOTAL_CAPITAL in settings',
         'Currently capital stays at $100K; compounding is manual; needs to update position sizing accordingly',
         'M', 'P1', 'IDEA'),

        ('Shadow / dry-run mode',
         'Run a full pipeline cycle in simulation alongside the live Alpaca run; compare picks',
         'Validates that simulation and live results match; useful for diagnosing drift',
         'M', 'P2', 'IDEA'),

        ('Automated parameter tuning',
         'Weekly backtest with current universe; flag if any key metric (win rate, R:R) degrades by >10%',
         'Proactive detection of strategy degradation; currently eval is reactive',
         'L', 'P2', 'BACKLOG'),

        ('DB cleanup job',
         'Monthly job to archive positions/scan_results older than 90 days to a separate table',
         'Supabase free tier has 500MB limit; months of history will eventually approach the ceiling',
         'S', 'P2', 'IDEA'),

        ('Multi-account support',
         'Run the same strategy across multiple Alpaca accounts with independent position tracking',
         'Scale-out option once live trading is validated; each account = separate capital pool',
         'L', 'P3', 'BACKLOG'),

        ('Failover to simulation on Alpaca error',
         'If Alpaca is down at premarket, auto-fall back to simulation mode and alert',
         'Currently a failed Alpaca connection crashes the run; simulation fallback keeps the pipeline running',
         'S', 'P2', 'IDEA'),
    ]
)

# ── 7. Dashboard & UX ─────────────────────────────────────────────────────────
doc.add_page_break()
heading('7. Dashboard & UX — Better Visibility')
body('These features improve the dashboard and make results easier to interpret and act on.')

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('Trailing stop dashboard annotations',
         '"Trail $X.XX ↑" on In Flight cards; "🔶 Trail Stop" vs "🔴 Stop Hit" in Today\'s Plan',
         'Makes the trailing stop feature visible and verifiable; shipped v5.3',
         'XS', 'P1', 'SHIPPED'),

        ('Pipeline funnel transparency — Today tab',
         'Step 1 shows full selection funnel: Universe → Earnings filter → Score pre-filter → ML ranking → VWAP enrichment → Sent to Claude. '
         'Orchestrator stores pipeline_counts per step in scan_results. Candidates table includes ml_score column, '
         'sorted in the exact order Claude received them.',
         'Today tab previously showed 3 metrics and described candidates going straight to Claude after earnings filter — '
         'hiding 4 intermediate steps. Full audit trail now visible for every premarket run.',
         'S', 'P1', 'SHIPPED v5.7'),

        ('Manual override — stop/restart via GitHub Actions',
         'control.py CLI: --action stop (writes halt_flag to Supabase), restart (updates to halt_flag_cleared + resumed_at), '
         'status (check current state). stop.yml workflow_dispatch: reason input + close_positions toggle (market-sells all '
         'open Alpaca positions via close_all_positions() before setting flag). restart.yml: one-click resume. '
         'Orchestrator skips every run while halt_flag is active. Red halt banner on every dashboard page.',
         'No way to stop the unsupervised agent without killing GitHub Actions jobs manually. '
         'One-click stop/restart from any device without terminal or Supabase access. '
         'close_positions toggle gives the choice between leaving Alpaca native stops active or market-selling everything.',
         'S', 'P0', 'SHIPPED v5.7'),

        ('Halt history preservation + eval chart markers',
         'restart() updates halt_flag to halt_flag_cleared with resumed_at (preserves history instead of deleting). '
         'eval.py detects halted days in the eval window from both flag types; prints in INTEGRITY CHECKS. '
         'Dashboard Performance tab: red dashed vline markers on Daily P&L and Portfolio Value charts for halted dates; '
         'Agent Scorecard shows halted-days count with active/cleared status.',
         'Without history, a halted trading day looks like a zero-P&L day with no explanation. '
         'Chart markers immediately distinguish "no trading" from "traded and lost". '
         'eval.py integrity section flags any active halt that was missed.',
         'S', 'P1', 'SHIPPED v5.7'),

        ('Intraday price chart with entry/exit markers',
         'Show a 1-day candlestick chart per position with entry price, target, stop, and trail stop marked',
         'Visual context for each trade; immediately shows whether price is trending toward target or stalling',
         'M', 'P2', 'IDEA'),

        ('Sector performance heatmap',
         'Show which sectors are performing best today across all positions',
         'Quick view of sector concentration and which bets are working; complements the trade heatmap',
         'S', 'P2', 'IDEA'),

        ('Historical trade browser',
         'Search/filter all past trades by ticker, date range, outcome (WIN/LOSS), close reason',
         'Currently hard to answer "how did IONQ do over the last 30 days?" without a SQL query',
         'M', 'P2', 'IDEA'),

        ('What-if simulator',
         'Slider to adjust TRAIL_PCT, TARGET_PCT, SCORE_THRESHOLD and see impact on historical trades',
         'Enables parameter exploration without rerunning backtest.py; visual tuning tool',
         'L', 'P3', 'BACKLOG'),

        ('Win/loss streak indicator',
         'Show current win/loss streak on Summary tab — e.g. "🔥 4W streak" or "❄️ 2L streak"',
         'Psychological context; long loss streaks may signal market regime change',
         'XS', 'P2', 'IDEA'),

        ('Anticipated vs. actual P&L trend line',
         'Add anticipated profit line to the daily P&L chart; show how often actual beats anticipated',
         'Validates whether Claude\'s estimates are calibrated or consistently over/under',
         'S', 'P2', 'IDEA'),

        ('Mobile push notifications',
         'Browser push or iOS/Android push on position close events',
         'More immediate than email; no app needed — browser push works from Streamlit via JS',
         'L', 'P3', 'BACKLOG'),
    ]
)

# ── 8. Data Sources ───────────────────────────────────────────────────────────
heading('8. Alternative Data Sources')
body('These features add new data inputs beyond yfinance to improve signal quality.')

add_table(
    ['Feature', 'What', 'Why', 'Effort', 'Priority', 'Status'],
    [
        ('Reddit/StockTwits sentiment',
         'Pull mention volume and sentiment for candidates from Reddit (pushshift) or StockTwits API',
         'Retail sentiment precedes short squeezes and momentum bursts; free API available',
         'M', 'P3', 'BACKLOG'),

        ('SEC 8-K filing monitor',
         'Alert when any universe ticker files a material 8-K (earnings, leadership change, deal)',
         'Material events often move stocks 5–15%; early detection = edge before price reaction',
         'M', 'P3', 'BACKLOG'),

        ('13-F insider buying tracker',
         'Flag tickers with recent insider purchases from SEC EDGAR 13-F/4 filings',
         'Insiders buying their own stock is a historically reliable positive signal',
         'M', 'P3', 'BACKLOG'),

        ('Macro data integration',
         'Pull yield curve slope, DXY (dollar index), oil price as additional market context for Claude',
         'Rising dollar hurts multinationals and EM; inverted yield curve = risk-off; currently not in context',
         'S', 'P3', 'BACKLOG'),

        ('Paid real-time data feed',
         'Replace yfinance (15-min delay) with a paid real-time feed (Polygon.io ~$30/month)',
         'Critical for live trading; intraday checks with 15-min delay miss fast moves; worthwhile if live capital deployed',
         'M', 'P2', 'IDEA'),
    ]
)

# ── 9. Real Edge Strategy ────────────────────────────────────────────────────
doc.add_page_break()
heading('9. Real Edge — Building Durable Competitive Advantage')
body(
    'Honest assessment of where the current system has edge — and where it does not — '
    'and what to build to create genuine, durable advantage over commodity signals.'
)

heading('Why Current Signals Have No Edge', level=2)
body(
    'RSI, MACD, and Bollinger Bands are in every trading textbook and every retail platform. '
    'Quant funds run the same signals at microsecond speed with better fills. '
    'By the time yfinance shows a signal (15-min delayed), it is already partially priced in. '
    'Commodity signals + delayed data = no structural advantage. '
    'The current system\'s real value is in its risk management discipline (3:1 R:R, stop-loss, '
    'sector caps, position sizing) — not its signal intelligence.'
)

heading('Where Retail Can Actually Win', level=2)
body(
    'Large hedge funds cannot exploit small-cap opportunities — $10M minimum position sizes '
    'make many universe tickers untradeable for institutional money. '
    'A focused system exploiting stock-specific catalysts and alternative data can find edge '
    'that institutional capital ignores by design. '
    'The goal is information asymmetry at the stock level, not speed.'
)

heading('Honest Confidence Assessment', level=2)
add_table(
    ['Scenario', 'Expected Daily P&L', 'Notes'],
    [
        ('Paper trading, bull market',    '$500–750',          'Close to 30-day backtest — favorable conditions'),
        ('Paper trading, choppy market',  '$0–200, some losses', 'Momentum signals fail in sideways/mean-reverting markets'),
        ('Real money, bull market',       '$200–400',          '~40% haircut from slippage, spreads, partial fills'),
        ('Real money, choppy market',     'Breakeven to -loss', 'No edge on non-trending days with current signals'),
        ('Real money, volatility spike',  'Reduced but protected', 'VIX gates reduce position count — good design'),
    ]
)
body('Overall confidence: 6.5/10 paper, 4.5/10 real money with ML scorer + lower target. ML scorer (AUC 0.78) adds a quantitative filter before Claude, increasing signal quality. Lower 2% target (vs 3%) is more achievable intraday — increases win rate at the cost of smaller wins per trade. 3:1 R:R maintained; break-even at 25% win rate.')

heading('Features That Create Real Edge', level=2)
body('Ranked by signal quality and buildability. Each addresses a specific gap in the current system.')

add_table(
    ['Priority', 'Feature', 'Edge Type', 'What', 'Why It Works', 'Effort', 'Status'],
    [
        ('0',
         'Native Alpaca trailing stop (shipped v5.6)',
         'Execution',
         'StopLossRequest(trail_percent=TRAIL_PCT*100) in bracket order stop-loss leg; '
         'USE_NATIVE_TRAILING_STOP feature flag; native_trail_active per position; '
         'exit_mechanism column tracks NATIVE_TRAIL / TARGET / MANUAL_TRAIL / STOP / EOD; '
         'eval.py validates cohort automatically',
         'Alpaca tracks peak price in real-time and fires stop on reversal — no 15-min polling gap; '
         'P0 blocker for real money removed; 2-week paper validation gate running (closes June 1)',
         'S', 'SHIPPED'),

        ('0',
         'ML candidate scorer (shipped v5.5)',
         'Analytical',
         'HistGradientBoosting model trained on 2y price history for 429 tickers; '
         'P(next-day high ≥ close × 1.02); 13 features; AUC 0.78 ± 0.04 (5-fold TimeSeries CV); '
         'sorts candidates before Claude; monthly auto-retrain via GitHub Actions',
         'ML pre-ranks candidates so Claude sees the highest-probability setups first — '
         'quantitative signal that complements LLM qualitative reasoning; '
         'top feature atr_pct (0.165) confirms volatility is the dominant predictor',
         'M', 'SHIPPED'),

        ('1',
         'Real-time data (Alpaca streams)',
         'Execution',
         'Replace yfinance 15-min delay with Alpaca real-time WebSocket feed for price + volume',
         'Foundational fix — stale prices cause entry slippage on every trade; free for Alpaca paper accounts',
         'M', 'IDEA'),

        ('2',
         'Post-earnings momentum agent',
         'Information',
         'Hunt stocks that beat estimates + raised guidance in last 48h; enter 1-2 days after report',
         'Post-earnings announcement drift (PEAD) is one of the most persistent anomalies in academic finance; '
         'current blackout throws this edge away',
         'M', 'IDEA'),

        ('3',
         'Market regime classifier',
         'Analytical',
         'Classify each day as Trending / Choppy / Volatile using VIX, market breadth, sector dispersion; '
         'switch strategy or skip accordingly',
         'Momentum signals work in trending markets and generate false signals in choppy ones; '
         'knowing the regime dramatically improves when to trade',
         'M', 'IDEA'),

        ('4',
         'Options flow signal',
         'Information',
         'Pull unusual options activity (large call buying relative to open interest) for candidates; '
         'boost score for tickers with informed positioning',
         'Unusual options activity consistently precedes price moves; represents informed institutional '
         'money moving before a catalyst — genuinely non-commodity information',
         'L', 'BACKLOG'),

        ('5',
         'Insider buying scan (Form 4)',
         'Information',
         'Weekly scan of SEC EDGAR Form 4 filings; flag tickers with cluster buys '
         '(3+ insiders buying in same week)',
         'Executive buying of own stock is historically one of the strongest retail-accessible signals; '
         'free public data that most retail systems never use',
         'M', 'IDEA'),

        ('6',
         'Claude on earnings transcripts',
         'Analytical',
         'Feed Claude actual earnings call text; extract tone, guidance language, management confidence '
         'beyond what the headline EPS number shows',
         'True LLM value-add — qualitative signals that quant models cannot capture; '
         'analyst surprise is often in the language, not just the numbers',
         'L', 'BACKLOG'),
    ]
)

heading('Realistic Ceiling With Real Edge Added', level=2)
body(
    'Even with all of the above, this remains a retail system. '
    'Realistic ceiling with full edge stack in a favorable market: $500–800/day real money. '
    'The regime classifier (item 3) is the highest-leverage single feature — it prevents '
    'trading on days where the signals do not work, which is where most losses come from. '
    'Validate paper trading for 30 days before deploying any real capital.'
)

# ── 10. Priority Summary ──────────────────────────────────────────────────────
doc.add_page_break()
heading('10. Priority Summary — What to Build Next')
body('Ranked by impact × effort. Focus on P1 ideas first, validate via backtest before deploying.')

add_table(
    ['Rank', 'Feature', 'Category', 'Effort', 'Status / Expected Impact'],
    [
        ('—', 'ML candidate scorer (step 1.76)', 'Signal Quality', 'M', 'SHIPPED v5.5 — sorts candidates by P(hit +2%) before Claude; AUC 0.78'),
        ('—', 'Target 2% / stop 0.67% (3:1 R:R)', 'Risk Mgmt', 'XS', 'SHIPPED v5.5 — lower achievable target; break-even at 25% win rate'),
        ('—', 'Time-of-day entry filter (9:45 AM ET)', 'Execution', 'XS', 'SHIPPED v5.4 — skips noisiest first 15 min of market'),
        ('—', 'Live price refresh (Alpaca ask prices)', 'Execution', 'S', 'SHIPPED v5.4 — real-time ask prices replace 15-min stale yfinance at step 1.8'),
        ('—', 'Native Alpaca trailing stop', 'Execution', 'L', 'SHIPPED v5.6 — trail_percent bracket leg; 2-week paper gate closes 2026-06-01'),
        ('—', 'Tiered lock-in — Thread 2', 'Position Mgmt', 'S', 'SHIPPED v5.7 — Tier 1 $716 ride with tighter trail, Tier 2 $1,000 ceiling close'),
        ('—', 'VWAP + RS vs SPY enrichment — Thread 1', 'Signal Quality', 'S', 'SHIPPED v5.7 — step 1.85 enriches candidates with live institutional signals before Claude'),
        ('—', 'Manual override stop/restart', 'Operational', 'S', 'SHIPPED v5.7 — control.py + stop.yml + restart.yml; halt banner on every dashboard page'),
        ('—', 'Halt history + eval markers', 'Operational', 'S', 'SHIPPED v5.7 — halt_flag_cleared + resumed_at; vline chart markers; INTEGRITY CHECKS in eval'),
        ('1', 'VWAP signal quality validation (June 1)', 'Validation', 'M', '~8 trading days accumulating; eval.py reports cohort deltas on June 1'),
        ('2', 'June 1 gate: python3 eval.py --days 14', 'Validation', 'S', 'Pass criteria: win rate ≥80%, avg P&L ≥$500/day, NATIVE_TRAIL confirmed'),
        ('3', 'Momentum confirmation — 15-min rule (V2f)', 'Signal Quality', 'M', 'Reduces false entries; should improve win rate by 3–5%'),
        ('4', 'Sector rotation scoring (V2e)', 'Signal Quality', 'M', 'Favors momentum sectors; expected to improve avg daily P&L'),
        ('5', 'Per-trade email on close', 'Monitoring', 'S', 'Highest quality-of-life; know immediately when a target is hit'),
        ('6', 'Partial profit taking (50% at +1%)', 'Position Mgmt', 'M', 'Locks in gain on fast movers; reduces STOP losses on reversals'),
        ('7', 'Daily 4:30 PM summary email', 'Monitoring', 'S', 'Replaces manual dashboard check at EOD'),
        ('8', 'Short selling / mean reversion', 'Strategy', 'M', 'Doubles opportunity set; hedges on down days'),
        ('9', 'Intraday price chart per position', 'Dashboard', 'M', 'Most requested UX improvement; visual trade context'),
    ]
)

body(
    'Items 1–2 are the immediate next milestone: accumulate 8 trading days of paper data '
    'and run the June 1 gate. Items 3 and 4 are the next signal-quality sprint after the gate. '
    'Item 5 (email alerts) is the highest quality-of-life gap remaining. '
    'Item 8 (short selling) requires the most design work but opens a full two-sided strategy.'
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/amitgarg/Claude Projects/trading-agent/Trading_Agent_Features.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
