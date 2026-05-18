"""Generates the Trading Agent project documentation as a Word document."""
import subprocess
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()

# ── Title Page ────────────────────────────────────────────────────────────────
title = doc.add_heading('AI Trading Agent', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)

sub = doc.add_paragraph('How It Was Built, How It Works, and What It Does')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph('Amit Garg  ·  May 2026  ·  v5.6  ·  Built with Claude Code + Anthropic API')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── 1. Overview ───────────────────────────────────────────────────────────────
heading('1. Overview')
body(
    'The AI Trading Agent is a fully autonomous stock trading simulation system built '
    'using Python, the Anthropic Claude API, Supabase (PostgreSQL), GitHub Actions, and '
    'Streamlit. It scans a dynamic universe of 450+ stocks and ETFs every trading day — '
    'refreshed weekly from S&P 500 + Nasdaq 100 components screened for liquidity and '
    'volatility — checks market conditions and news intelligence, selects the highest-conviction '
    'setups using a multi-agent AI pipeline, manages simulated positions throughout the day, and '
    'logs performance to a live password-protected dashboard — all without any human '
    'intervention after setup.'
)
body(
    'The system runs entirely in the cloud. Once deployed, it wakes up on its own schedule, '
    'makes decisions, stores results, and goes back to sleep. The dashboard shows the full '
    'workflow — market conditions, screened stocks, selected trades, and live position P&L — '
    'accessible from any device via a public URL.'
)

heading('Key Goals', 2)
bullet('Generate $750–$1,000/day in simulated paper trading profit')
bullet('Run fully autonomously Monday–Friday with no manual intervention')
bullet('Make every decision auditable — every trade has reasoning attached')
bullet('Protect against high-volatility markets, earnings binary events, and bad news')
bullet('Keep infrastructure cost at or near zero (GitHub Actions + Supabase + Streamlit Cloud free tiers)')

heading('Backtest Results (30 trading days, Apr–May 2026)', 2)

body('Configuration evolution — iterative tuning from initial 52-ticker setup to current:')
add_table(
    ['Configuration', 'Avg Daily P&L', 'Win Days', 'Annualized', 'Grade'],
    [
        ('52 tickers, 3 positions, 2% target', '$198', '50%', '50%', 'D'),
        ('408 tickers, 10 positions, 2% target', '$548', '87%', '138%', 'B'),
        ('434 tickers, 15 positions, 2.5% target', '$690', '90%', '174%', 'B'),
        ('430 tickers, 15 positions, 3% target', '$800', '93%', '202%', 'B'),
        ('429 tickers, 15 positions, 2% target + ML scorer (current)', '$716 baseline', '90%', '180%', 'B'),
    ]
)

body('V2 intelligence gate validation — baseline vs. V2-gated (Apr 6 – May 15, 2026):')
add_table(
    ['Metric', 'Baseline (no gates)', 'V2-Gated (v2.3)', 'Delta'],
    [
        ('Total P&L (30 days)', '$24,023', '$21,474', '-$2,549'),
        ('Avg daily P&L', '$801', '$716', '-$85'),
        ('Win days', '28/30 (93%)', '27/30 (90%)', '-1 day'),
        ('Win rate (per trade)', '48.2%', '48.5%', '+0.3%'),
        ('Reward:risk', '2.63x', '2.66x', '+0.02x'),
        ('Annualized return', '201.8%', '180.4%', '-21.4%'),
        ('Grade', 'B', 'B', 'same'),
    ]
)
body('Gate activity: fired on 9 of 30 days — VIX spikes (Apr 6–8), bearish futures (Apr 13, 20), '
     'FOMC (May 6), CPI/NFP (Apr 10, May 1, May 13). Apr 10 NFP and May 13 CPI were both losing '
     'days — position cap directly limited damage on both. Win rate and R:R improved slightly with '
     'gates applied.')
body('The $2,549 gap vs baseline is the cost of caution during a bull recovery window — '
     'acceptable insurance. In a genuine market decline the gates would save significantly more.')

heading('One-Month Forward Projection (22 trading days)', 2)
add_table(
    ['Scenario', 'Daily P&L Assumption', 'Monthly P&L', 'Ending Capital'],
    [
        ('Conservative', '80% of backtest avg ($573/day)', '~$12,600', '~$112,600'),
        ('Base case', 'Backtest avg ($716/day)', '~$15,750', '~$115,750'),
        ('Optimistic', 'Baseline avg ($801/day)', '~$17,600', '~$117,600'),
    ]
)
body('Compounding path to $1K/day: at $716/day on $100K starting capital, the portfolio reaches '
     '~$140K in ~45 trading days (~2 months), at which point the same % return = $1,000/day. '
     'No settings changes needed — capital growth does the work automatically.')
body('Important caveat: the 30-day backtest window was a strong bull recovery after the tariff '
     'panic selloff — one of the better environments for momentum trading. A sideways or declining '
     'month would look different. V2 gates are specifically designed for that scenario.')

# ── 2. What We Built ──────────────────────────────────────────────────────────
doc.add_page_break()
heading('2. What We Built')

heading('Project Structure', 2)
body('The project lives at: /Users/amitgarg/Claude Projects/trading-agent/')
body('GitHub repository: https://github.com/amitgarg73/trading-agent (private)')
code('trading-agent/')
code('├── orchestrator.py              # Master controller — chains all agents')
code('├── requirements.txt             # Python dependencies')
code('├── schema.sql                   # Supabase database schema')
code('├── backtest.py                  # 30-day historical backtest (no Claude API)')
code('├── eval.py                      # Eval script — grades live performance; --write saves to Supabase (auto-runs EOD)')
code('├── generate_doc.py              # Generates this Word document')
code('├── generate_prd.py              # Generates the PRD document')
code('├── train_model.py               # Standalone ML model trainer (HistGradientBoosting, 429 tickers, 2y data)')
code('├── models/')
code('│   ├── xgb_scorer.pkl              # Trained ML model — P(hit +2% next day) per candidate')
code('│   └── feature_columns.json        # Ordered feature list matching inference in ml_scorer.py')
code('├── .env                         # Local secrets (gitignored)')
code('├── .env.example                 # Template for secrets')
code('├── streamlit_secrets.toml       # Streamlit Cloud secrets template (gitignored)')
code('├── .gitignore                   # Excludes .env and secrets files')
code('├── .github/')
code('│   └── workflows/')
code('│       ├── trading.yml              # Premarket/intraday/EOD schedule')
code('│       ├── universe_refresh.yml     # Monday 8:30 AM ET — weekly universe update')
code('│       └── retrain_model.yml        # 1st of month 10 AM UTC — ML model retrain + auto-commit pkl')
code('├── scanner/')
code('│   ├── scanner.py                  # Market scanner (yfinance + TA, dynamic universe)')
code('│   └── ml_scorer.py                # ML candidate ranker — loads pkl, scores all candidates by P(hit +2%)')
code('├── agents/')
code('│   ├── market_context.py           # V2a: VIX gate + futures signal + international')
code('│   ├── news_intel.py               # V2b: earnings blackout + news headlines')
code('│   ├── universe_refresh.py         # V3a: weekly S&P500+Nasdaq100 screener')
code('│   ├── strategy.py                 # Claude picks the best trades')
code('│   ├── risk.py                     # Claude validates risk parameters')
code('│   ├── portfolio.py                # Opens/tracks simulated positions')
code('│   ├── intraday.py                 # Monitors open positions mid-day')
code('│   ├── performance.py              # Calculates EOD P&L')
code('│   ├── sector_guard.py             # V2d: sector cap (max 3 per sector, Unknown passes through)')
code('│   ├── guardrails.py               # V5: 6 safety checks before any trade executes')
code('│   └── alpaca_broker.py            # Alpaca paper trading broker (bracket orders, sync, fills)')
code('├── config/')
code('│   ├── settings.py                 # Capital, thresholds, static stock universe (fallback)')
code('│   └── company_names.py            # Static ticker→company name lookup (~430 tickers)')
code('├── core/')
code('│   └── db.py                       # Supabase database client')
code('└── dashboard/')
code('    └── app.py                      # Streamlit dashboard — full workflow view')

# ── 3. The Agent Pipeline ─────────────────────────────────────────────────────
doc.add_page_break()
heading('3. The Agent Pipeline')
body(
    'The system is a chain of specialized agents. Each agent does one job and passes its '
    'output to the next. orchestrator.py coordinates the chain. There are three daily '
    'pipeline runs: Premarket, Intraday, and End of Day.'
)

heading('3a. Premarket Pipeline (9:45 AM ET)', 2)
body('Runs once at 9:45 AM ET — skips first 15 min of market (widest spreads). Seven stages plus two new pre-entry steps:')
body('Concurrent Run Lock')
bullet('First action before any pipeline stage runs')
bullet('Checks if a premarket scan_results row already exists for today — if yes, exits immediately with a skip message')
bullet('Prevents duplicate trades if GitHub Actions fires two runs in parallel (e.g., delayed + on-time cron overlap)')
bullet('Safe for manual reruns — a re-triggered premarket run after 9AM will exit cleanly without re-scanning or re-opening positions')
doc.add_paragraph()

body('Stage 0 — Market Context Agent (V2a)')
bullet('Fetches VIX (^VIX) — fear index, measures overall market volatility')
bullet('Fetches US futures: S&P500 (ES=F), Nasdaq (NQ=F), Dow (YM=F) — pre-market direction')
bullet('Fetches international markets: Nikkei, FTSE, DAX, Hang Seng, Shanghai — global context')
bullet('VIX tiered gate: VIX >45 → 2 positions; VIX 30–45 → 3 positions; VIX 25–30 → 5 positions; VIX 20–25 → 10 positions; VIX <20 → 15 positions (no hard skip on VIX)')
bullet('Futures avg down >1.5%: skip trading (strong sell-off); down >0.5%: bearish bias, reduce positions')
bullet('Passes full market summary to Claude strategy agent as context')
bullet('Stores VIX, futures, and international data in scan_results for full audit trail')

body('Stage 1 — Market Scanner')
bullet('Loads dynamic universe via load_universe() — reads latest refresh from Supabase if ≤7 days old, else falls back to static settings.py list')
bullet('Pulls price data for 450+ stocks and ETFs using yfinance (free, no API key needed)')
bullet('Calculates technical indicators: RSI, MACD, Bollinger Bands, volume ratio, ATR, SMA 20/50')
bullet('Scores each ticker from -10 to +10 based on signal strength')
bullet('Returns candidates with score ≥ 3 (configurable in settings.py)')

body('Stage 1.5 — News Intelligence Agent (V2b)')
bullet('Checks earnings calendar for every candidate via yfinance ticker.calendar')
bullet('Removes any ticker with earnings today or tomorrow (binary event = gap risk)')
bullet('Fetches 3 most recent news headlines for each remaining candidate')
bullet('Passes earnings-filtered candidate list and news context string to strategy agent')
bullet('Logs all blocked tickers with reason for audit trail')

body('Stage 1.75 — Strategy Pre-Filter')
bullet('Runs before Claude — no API call needed (deterministic score filter)')
bullet('Filters candidates to technical_score >= STRATEGY_MIN_SCORE (default: 4) and positive score only')
bullet('Typical reduction: 150+ candidates → 30–50 passed to Claude — cuts input tokens ~60–70%')
bullet('Log signature: [ 1.75/4 ] Strategy pre-filter: X → Y candidates (score ≥ 4)')
bullet('Tuning: lower STRATEGY_MIN_SCORE to 3 if missing good setups; raise to 6–7 for further token cuts')

body('Stage 1.76 — ML Scorer  (v5.5 — SHIPPED)')
bullet('Loads HistGradientBoostingClassifier from models/xgb_scorer.pkl (trained on 2y price history, 429 tickers)')
bullet('Scores each candidate: P(next-day high ≥ close × 1.02) — probability of hitting +2% the next day')
bullet('13 features: rsi, macd_hist, bb_pct, vol_ratio, atr_pct, dist_sma20, dist_sma50, mom1, mom5, range_52w_pct, dow, vix, technical_score')
bullet('Top feature importance: atr_pct (0.165), vix (0.038), dow (0.011), vol_ratio (0.008) — AUC 0.78 ± 0.04 (5-fold TimeSeries CV)')
bullet('Sorts candidates descending by ml_score before passing to Claude — highest-probability setups ranked first')
bullet('Graceful fallback: if model file missing, sets ml_score=None and passes candidates unchanged (no pipeline interruption)')
bullet('Log signature: [ 1.76/4 ] ML scoring: X candidates ranked by P(hit +2%) — top score: Y.YY')

body('Stage 1.8 — Live Price Refresh (Alpaca mode only)')
bullet('Fetches real-time ask prices for all pre-filter candidates via Alpaca Market Data API (batch call)')
bullet('Updates candidate current_price with live ask — overrides 15-min delayed yfinance price')
bullet('10% sanity guard: rejects update if ask differs >10% from yfinance price (data anomaly protection)')
bullet('Log signature: [ 1.8/4 ] Live price refresh: X/Y tickers updated from Alpaca')
bullet('Why: Claude was setting entry prices on 15-min stale data — target price could already be eaten before order submitted')

body('Stage 2 — Strategy Agent (Claude AI)')
bullet('Receives scored candidates, market summary (VIX + futures + news), and max_positions for today')
bullet('Claude reads market context, momentum, and signal combinations')
bullet('Selects up to max_positions highest-conviction trades for the day')
bullet('Assigns entry price (hard: 2% target above entry), stop loss (hard: 0.67% below entry — maintains 3:1 R:R)')
bullet('Sets position size ($5K–$7K per trade), shares, estimated profit, confidence (HIGH/MEDIUM/LOW)')
bullet('Writes 2–3 sentence reasoning for every trade citing specific signals from the scan data')
bullet('Can select zero trades if no high-conviction setups exist — protects principal')

body('Stage 3 — Risk Agent (Claude AI)')
bullet('Reviews every proposed trade against hard-coded risk rules')
bullet('Rejects if: stop loss > 0.67% of entry, target < 2% of entry, reward:risk < 3:1 (break-even at 25% win rate)')
bullet('Rejects if: target below entry (BUY) or stop above entry (BUY)')
bullet('Position sizing: confidence-weighted — HIGH=$7,000, MEDIUM=$6,000, LOW=$5,000; risk agent _apply_confidence_sizing() overrides whatever Claude sets before running validation — ensures correct sizing even if Claude hallucinates a different value')
bullet('Returns approved and rejected trades with specific rejection reasons in plain English')

body('Stage 3.5 — Sector Correlation Guard (V2d)')
bullet('Groups approved trades by sector using yfinance sector lookup (ETFs classified via ETF_UNIVERSE set — no API call needed)')
bullet('Caps positions at MAX_PER_SECTOR=3 per sector — drops lowest-confidence excess')
bullet('Unknown sector (yfinance rate-limited or unclassifiable) passes through without cap — safe fallback')
bullet('No Claude API call — deterministic rule-based filter')

body('Stage 3.75 — Guardrails (V5)')
bullet('Action whitelist: only BUY permitted — rejects SELL, SHORT, or any other action Claude might hallucinate')
bullet('Ticker whitelist: rejects any ticker not in the current universe (static or dynamic)')
bullet('Duplicate position guard: checks both currently-open positions and positions already traded today — prevents double-entry on the same ticker')
bullet('Price sanity: fetches current market price for each trade; rejects if entry price is >5% from market price (catches hallucinated or stale prices)')
bullet('Capital check: if broker=alpaca, checks Alpaca buying_power covers the position_size before submitting (prevents over-leverage)')
bullet('Daily loss limit: if today\'s realized P&L is already below -$300, blocks ALL remaining new trades until next day')
bullet('scan_results updated in DB with sector_blocked and guardrail_blocked data for full dashboard visibility (fixes prior bug where sector_blocked was always stored as [])')

body('Stage 4 — Portfolio Agent')
bullet('Broker abstraction: broker="simulation" (default, yfinance) or broker="alpaca"')
bullet('simulation mode: writes positions to Supabase only')
bullet('alpaca mode: submits bracket order as LIMIT entry (entry_price × 1.001) + take-profit limit + stop-loss; stores alpaca_order_id in positions table')
bullet('Limit order vs market: pays slightly less than ask; still fills quickly on liquid stocks (1M+ avg volume); avoids paying full spread on entry')
bullet('Check-before-insert prevents duplicate key errors on manual reruns')
bullet('Writes trade plan summary and all individual positions to the database')

heading('3b. Intraday Pipeline (Every 15 min, 10:00 AM – 3:45 PM ET)', 2)
body('Reconciliation (alpaca mode only — runs first):')
bullet('Compares all Supabase OPEN positions against Alpaca\'s actual live positions (get_open_tickers())')
bullet('Any Supabase OPEN position not in Alpaca → marked UNFILLED (status=CLOSED, close_reason=UNFILLED, realized_pnl=0, close_price=entry_price)')
bullet('Catches the rare case where a bracket order was submitted (order_id stored) but the entry leg never filled (symbol halted mid-submission, etc.)')
bullet('Dashboard and eval both exclude UNFILLED entries alongside CLEANUP — prevents stale data surfacing')
doc.add_paragraph()
body('Position monitoring:')
bullet('simulation mode: fetches current prices via yfinance, calculates unrealized P&L')
bullet('alpaca mode: calls Alpaca API (get_position_data) to sync live price and unrealized P&L for each open position')
bullet('alpaca mode: when a position disappears from Alpaca (bracket fill triggered), fetches actual fill price from bracket order leg to compute realized P&L')
bullet('Trailing stop: effective_stop = max(original_stop_loss, high_watermark × 0.99); high_watermark updated every 15-min cycle to track the highest price seen since entry; fires if price pulls back 1% from its peak while still in profit — locks in gains on strong movers before reversal; Alpaca mode calls close_position() directly (bracket auto-cancels)')
bullet('Closes positions that hit +3% target (take profit) or effective stop (cut loss or protect gain); close_reason=TARGET or STOP')
bullet('Records close reason: TARGET, STOP, EOD, or LOCK_IN')
bullet('No overnight holds — all positions closed by end of day')
doc.add_paragraph()
body('Daily profit lock-in (runs after every position sync):')
bullet('Checks today\'s total realized P&L (excluding CLEANUP, UNFILLED, LOCK_IN positions)')
bullet('If realized P&L ≥ DAILY_LOCK_IN_TARGET ($716 — the 30-day backtest average), closes ALL remaining open positions immediately via Alpaca with close_reason=LOCK_IN')
bullet('Books the gain and stops trading for the day — does not wait for EOD')
bullet('Rationale: realized P&L only (not unrealized) — no slippage risk, gain is fully locked before closing remaining positions')
bullet('Dashboard shows LOCK_IN positions as "🎯 Day Locked" in Today\'s Plan status column')
bullet('Threshold configurable via DAILY_LOCK_IN_TARGET in config/settings.py — raise as capital compounds')

heading('3c. End of Day Pipeline (4:30 PM ET)', 2)
bullet('simulation mode: closes remaining open positions at yfinance market close price')
bullet('alpaca mode: calls close_position() on Alpaca, uses actual fill price (not yfinance) for realized P&L; retries once after 2s on failure; prints loud warning if still failing (⚠️ manual close required)')
bullet('Calculates total daily P&L, win rate, best and worst trade')
bullet('Writes daily performance record to Supabase daily_performance table')
bullet('Updates portfolio capital for compounding into the next trading day')
doc.add_paragraph()
body('Automatic eval (runs immediately after EOD close, every trading day):')
bullet('eval.py --days 30 --write runs as a separate step in the EOD GitHub Actions job')
bullet('Computes score (0–100), grade (A/B/C/D), avg daily P&L vs target, win days, trade win rate, actual reward:risk, best/worst trade, close reason breakdown, and tuning recommendations')
bullet('Saves results to scan_results table (scan_type="eval") — visible in Performance tab Agent Scorecard on dashboard immediately')
bullet('Excludes CLEANUP and UNFILLED positions from all calculations — score reflects actual trades only')
bullet('If fewer than 30 days of data exist, eval covers all available days')

# ── 4. The Technology Stack ───────────────────────────────────────────────────
doc.add_page_break()
heading('4. The Technology Stack')

add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ('AI Brain', 'Anthropic Claude API (claude-sonnet-4-6)', 'Strategy decisions and risk validation (~$0.15–0.30/day)'),
        ('Market Data', 'yfinance (free, no key needed)', 'Price history, volume, OHLC, news, earnings calendar'),
        ('Technical Analysis', 'ta (Python library)', 'RSI, MACD, Bollinger Bands, ATR, SMA'),
        ('Database', 'Supabase (PostgreSQL, free tier)', 'Stores all trades, positions, P&L history'),
        ('Scheduler', 'GitHub Actions (cron, free tier)', 'Runs the pipeline on schedule automatically'),
        ('Dashboard', 'Streamlit Cloud (free tier)', 'Live workflow view, password protected, always on'),
        ('Secrets (cloud)', 'GitHub Secrets + Streamlit Cloud Secrets', 'API keys injected at runtime, never in code'),
        ('Secrets (local)', '.env file (gitignored)', 'Local development credentials'),
        ('Version Control', 'Git + GitHub (private repo)', 'Code storage and CI/CD trigger (v1.0, v2.0, v2.1)'),
        ('Broker', 'Alpaca Paper Trading (alpaca-py)', 'Real bracket order simulation — entry, take-profit, stop-loss in one call'),
        ('Auth (GitHub)', 'SSH key pair', 'Passwordless git push from local machine'),
    ]
)

heading('Model Cost', 2)
body(
    'Switched from claude-opus-4-7 to claude-sonnet-4-6 to reduce API costs. '
    'Estimated cost: $0.15–0.30/day for two Claude calls (strategy + risk). '
    'At that rate, $10 of API credit lasts 33–66 trading days. '
    'Note: Anthropic console has no usage alert feature — monitor manually at console.anthropic.com.'
)

# ── 5. The Database Schema ────────────────────────────────────────────────────
heading('5. The Database Schema (Supabase)')
body('Five tables in PostgreSQL, hosted on Supabase free tier. Schema defined in schema.sql.')

add_table(
    ['Table', 'What It Stores', 'Key Constraint'],
    [
        ('trade_plans', 'One row per trading day — market context, estimated profit, risk note', 'UNIQUE on date'),
        ('planned_trades', 'Individual trades — ticker, entry, target, stop, confidence, reasoning, status', 'FK to trade_plans'),
        ('positions', 'Simulated open/closed positions with real-time P&L tracking', 'FK to planned_trades'),
        ('daily_performance', 'EOD summary — total P&L, win rate, best/worst trade, portfolio value', 'UNIQUE on date'),
        ('scan_results', 'Raw scanner + market context output every premarket run (full audit trail)', 'Indexed by date'),
    ]
)
body(
    'scan_results stores rich context per premarket run: VIX, futures (S&P/Nasdaq/Dow), '
    'international markets (Nikkei/FTSE/DAX/Hang Seng/Shanghai), futures_bias, screened '
    'candidates, and earnings-blocked tickers. All data is stored as JSONB.'
)

# ── 6. The Automation (GitHub Actions) ───────────────────────────────────────
doc.add_page_break()
heading('6. The Automation — GitHub Actions')
body(
    'The workflow file at .github/workflows/trading.yml runs the orchestrator on a UTC cron '
    'schedule. GitHub provides 2,000 free minutes/month — well within limits for current run '
    'frequency. The workflow also supports manual trigger for testing and reruns.'
)

add_table(
    ['Run', 'UTC Cron', 'ET Time', 'What Happens'],
    [
        ('Universe Refresh', '30 12 * * 1', '8:30 AM Mondays', 'Fetch S&P500+Nasdaq100, screen for ATR/volume, save 450+ tickers to Supabase'),
        ('Premarket', '45 13 * * 1-5', '9:45 AM Mon–Fri', 'Market check → Scan → Pre-filter → Live prices → News → Strategy → Risk → Limit orders'),
        ('Intraday', '*/15 14-19 * * 1-5', 'Every 15 min 10AM–3:45PM', 'Monitor positions, close on target/stop, trail stop check'),
        ('EOD', '30 20 * * 1-5', '4:30 PM Mon–Fri', 'Close remaining, calculate daily P&L, then auto-run eval (30-day, saves to Supabase)'),
    ]
)
body('Cron reliability: GitHub Actions scheduler can fire 5–15 min late. cron-job.org external service fires workflow_dispatch at exact UTC times — more reliable for time-sensitive premarket runs. GitHub schedule kept as backup; concurrent run lock handles any duplicates.')

body('Each run:')
bullet('Checks out the code from GitHub')
bullet('Sets up Python 3.11')
bullet('Installs dependencies from requirements.txt')
bullet('Injects secrets (ANTHROPIC_API_KEY, SUPABASE_URL, SUPABASE_KEY, ALPACA_API_KEY, ALPACA_SECRET_KEY) from GitHub Secrets')
bullet('Runs: python orchestrator.py --mode premarket|intraday|eod --broker alpaca (scheduled runs default to alpaca)')
bullet('FORCE_JAVASCRIPT_ACTIONS_TO_NODE24=true set to suppress Node.js deprecation warnings')

heading('Manual Trigger', 2)
body(
    'Any run can be triggered manually from GitHub → Actions → Trading Agent → Run workflow. '
    'Select the mode (premarket, intraday, or eod) and broker (simulation or alpaca) from the dispatch inputs and click Run. '
    'Used for testing and re-running failed jobs. URL: https://github.com/amitgarg73/trading-agent/actions'
)

# ── 7. The Dashboard ──────────────────────────────────────────────────────────
heading('7. The Dashboard — Workflow View')
body(
    'A Streamlit web app (dashboard/app.py) shows the full pipeline progression in real time. '
    'Deployed to Streamlit Community Cloud — always on, accessible from any device. '
    'Password protected. Connects directly to Supabase and reads live data. '
    'Five tabs: Summary, Today, Positions, Performance, Scan Log.'
)
body(
    'Sidebar: capital and daily target displayed at all times. Refresh button forces data reload. '
    'Staleness detection: if no premarket run exists for today, a STALE badge and warning banner appear across Summary and Today tabs, '
    'showing the most recent available data rather than a blank screen.'
)

heading('Summary Tab — Daily Cockpit', 2)
body('The Summary tab is the primary at-a-glance view for the current trading day:')
bullet('Header: date and status badge (TRADING / STALE / PENDING) color-coded green/gray')
bullet('KPI row (6 metrics): Capital · Today\'s P&L with % return delta · Realized P&L · Unrealized P&L · Anticipated profit vs target · % Return')
bullet('Trade stats row: Open Positions · Closed Today · Win Rate (W/L split) · Total Trades selected by Claude')
bullet('In Flight section: one inline card per open position — ticker + company name, entry price, current price, target/stop, unrealized P&L in green/red')
bullet('Today\'s Plan table: all planned trades with Status column (⏳ Pending → 🟢 In Flight → ✅ Target Hit / 🔴 Stop Hit / ⏰ EOD Close) plus estimated and actual P&L. Expandable Claude reasoning per trade.')
bullet('Trade Heatmap: Plotly Treemap — each block is one stock, sized by position size, colored green/red by P&L. Hover shows full detail. Gray = pending/flat.')

heading('Today Tab — Pipeline Workflow', 2)
body('Shows the complete premarket pipeline as a numbered step-by-step view:')
add_table(
    ['Step', 'What You See'],
    [
        ('0 — Market Conditions', 'VIX (green/yellow/red), Fear & Greed index, futures bias, S&P/Nasdaq/Dow % change. FOMC/CPI/NFP event banners when active. Expandable international markets (Nikkei, FTSE, DAX, Hang Seng, Shanghai). All metrics have help tooltips explaining acronyms and thresholds.'),
        ('1 — Scanner', 'Total candidates found, earnings-blocked count, full table of screened stocks with technical scores and company names'),
        ('2 — Strategy & Risk', "Claude's market read, estimated profit vs $1K target, approved trades table, expandable per-trade reasoning. Sector cap blocked expander (🏭). Guardrails blocked expander (🛑)."),
        ('3 — Live Positions', 'Open position cards (TICKER · Company, entry/current/target/stop, color-coded P&L) and closed positions table with realized P&L and close reason'),
    ]
)

heading('Positions Tab', 2)
bullet('All currently OPEN positions with live unrealized P&L (updated every 30 min by intraday agent)')
bullet('All positions CLOSED today with realized P&L, fill price, and close reason (TARGET/STOP/EOD)')
bullet('Note: this tab shows all DB positions without plan scoping — use Today tab for the cleanest current-day view')

heading('Performance Tab', 2)
body('Historical performance view with automated Agent Scorecard:')
bullet('Agent Scorecard (top, expanded by default): latest eval results from Supabase — Score (0–100), Grade (A/B/C/D), Avg Daily P&L vs target, Win Days, Trade Win Rate, Actual R:R, close reason breakdown, best/worst trade, tuning recommendations. Updated automatically every trading day after EOD.')
bullet('KPI row: Total P&L · Avg Daily P&L · Win Days · Avg Trade Win % · Portfolio Value · Annualized Return')
bullet('Daily P&L bar chart: Plotly go.Bar — green/red per day with dotted target line ($716/day)')
bullet('Portfolio value + cumulative P&L: Plotly dual-axis line chart — portfolio value (left axis) and cumulative P&L (right axis)')
bullet('Daily log table: one row per day sorted descending')

heading('Scan Log Tab', 2)
bullet('Every premarket and intraday scan entry as expandable rows')
bullet('Premarket rows show: VIX, futures bias, candidate count, sector/guardrail blocks, raw JSONB results')
bullet('Useful for auditing why the pipeline skipped, reduced positions, or blocked specific trades on any given day')

heading('Access', 2)
bullet('URL: Your Streamlit Cloud app URL (dashboard password protected)')
bullet('Run locally: python3 -m streamlit run dashboard/app.py (opens at http://localhost:8501)')
bullet('Secrets set via: App → ⋮ → Settings → Secrets (TOML format) in Streamlit Cloud')

# ── 8. Configuration ──────────────────────────────────────────────────────────
doc.add_page_break()
heading('8. Configuration — What You Can Tune')
body('All tunable parameters live in config/settings.py:')

add_table(
    ['Parameter', 'Current Value', 'What It Controls'],
    [
        ('TOTAL_CAPITAL', '$100,000', 'Starting portfolio size'),
        ('DAILY_PROFIT_TARGET', '$1,000', 'Target P&L shown in dashboard'),
        ('MAX_POSITIONS', '15', 'Max concurrent open positions (reduced by market conditions)'),
        ('MAX_POSITION_PCT', '7%', 'Max capital per single trade ($7K)'),
        ('MIN_POSITION_PCT', '5%', 'Min capital per single trade ($5K)'),
        ('TARGET_PCT', '2%', 'Profit target per trade (lowered from 3% — more achievable intraday move; hard rule)'),
        ('MAX_LOSS_PER_TRADE', '0.67%', 'Stop loss (lowered from 1% to maintain 3:1 R:R with 2% target; break-even at 25% win rate)'),
        ('MIN_REWARD_RISK', '3.0', 'Minimum reward:risk ratio to approve a trade (3:1)'),
        ('SCORE_THRESHOLD', '3', 'Minimum scanner score to be a strategy candidate'),
        ('RSI_OVERSOLD', '35', 'RSI level considered oversold (buy signal)'),
        ('RSI_OVERBOUGHT', '65', 'RSI level considered overbought (sell signal)'),
        ('MIN_VOLUME_RATIO', '1.5x', 'Volume must be 1.5x the 20-day average'),
        ('MIN_AVG_VOLUME', '1,000,000', 'Minimum average daily volume (liquidity floor — raised from 500K to reduce spreads and slippage)'),
        ('STRATEGY_MIN_SCORE', '4', 'Pre-filter: only pass candidates with technical_score >= this to Claude; reduces input tokens ~60–70%'),
        ('MIN_PRICE', '$5.00', 'Minimum stock price (filters out penny stocks)'),
        ('MAX_PER_SECTOR', '3', 'V2d: max positions in any single sector per day'),
        ('DAILY_LOSS_LIMIT', '-$300', 'V5: stop opening new trades if today\'s realized P&L drops below this'),
        ('PRICE_SANITY_PCT', '5%', 'V5: reject trade if entry price is more than 5% from current market price'),
        ('DAILY_LOCK_IN_TARGET', '$716', 'V5: close all open positions and lock in gains once today\'s realized P&L hits this threshold — set to 30-day backtest average daily P&L; raise as capital compounds'),
        ('TRAIL_PCT', '1%', 'V5.3: trailing stop — close if price drops 1% from highest point seen since entry; effective_stop = max(original_stop, high_watermark × 0.99); ratchets up as stock rises in profit'),
        ('POSITION_SIZE_BY_CONFIDENCE', 'HIGH=$7K / MED=$6K / LOW=$5K', 'V5.3: risk agent maps Claude confidence level to position size before validation; auto-corrects regardless of what Claude sets in the trade JSON'),
    ]
)

heading('Stock Universe (430+ Tickers Across 15 Sectors)', 2)
body(
    'The universe was expanded from 52 to 430+ tickers based on 30-day backtest results. '
    'Drag tickers (consistent underperformers) were removed. Top performers in backtest: '
    'IONQ, WOLF, RPD, QUBT, ACHR, ASAN, SOXS, AXTI, NET, MRNA.'
)
bullet('Mega-cap tech: AAPL, MSFT, NVDA, GOOGL, AMZN, TSLA, META')
bullet('AI / Cloud: PLTR, CRWD, SNOW, NET, DDOG, MDB, SMCI, AI, GTLB, CFLT')
bullet('Semiconductors: AMD, AVGO, QCOM (removed), AMAT, LRCX (removed), MRVL, KLAC, ONTO, WOLF, AXTI')
bullet('Quantum computing: IONQ, QUBT, RGTI, QBTS')
bullet('Fintech / Crypto: SOFI, HOOD, COIN, MSTR, RIOT, MARA, HUT8')
bullet('Biotech / Health: LLY, MRNA, ABBV, RXRX, ACHR (aerospace), BEAM, EDIT')
bullet('Energy / Clean: XOM, CVX, OXY, ENPH, FSLR, PLUG, RUN')
bullet('Defense / Industrial: LMT, RTX, KTOS, NOC, HII')
bullet('Retail / Consumer: AMZN, WMT, TGT, COST, SHOP')
bullet('Leveraged ETFs: TQQQ, SQQQ, SOXS, UPRO, SPXU, UVXY')
bullet('Removed drags: PYPL, META (moved to mega-cap), ARKK, IWM, JPM, IBM, MA, ROOT, PSA, TWLO, INTU, LRCX, ASML, HPE, UBER, TENB, ARKW, DUOL, GPN, QCOM, NKE, BROS, GLBE, EXEL')

# ── 9. V2 Intelligence Layer ──────────────────────────────────────────────────
doc.add_page_break()
heading('9. V2/V3 Intelligence Layer')
body(
    'V2 adds pre-trade intelligence gates; V3 adds dynamic data sourcing. Each phase adds '
    'a new protection, signal, or data quality improvement.'
)

add_table(
    ['Phase', 'What', 'Status'],
    [
        ('V2a', 'Volatility gate + futures signal + international markets', 'Built and deployed (v2.0)'),
        ('V2b', 'Earnings blackout + news headlines for strategy context', 'Built and deployed (v2.1)'),
        ('V2c', 'Fear & Greed Index + FOMC/CPI/NFP economic calendar gates', 'Built and deployed (v2.2)'),
        ('V2c.1', 'Tune F&G gate: confirming signal only (not standalone)', 'Built and deployed (v2.3)'),
        ('V3a', 'Dynamic universe refresh — weekly S&P500+Nasdaq100 screener', 'Built and deployed (v3.0)'),
        ('V2g', 'Alpaca paper trading integration — real order simulation', 'Built and deployed (v4.0)'),
        ('V2d', 'Sector correlation guard — max 3 positions per sector, drops lowest-confidence excess', 'Built and deployed (v4.1)'),
        ('V5', 'Guardrails — 6 safety checks before any trade executes (action, whitelist, duplicate, price, capital, loss limit)', 'Built and deployed (v5.0)'),
        ('—', 'Prompt caching (strategy SYSTEM ephemeral, 90% discount on cached tokens)', 'Built and deployed (v5.4)'),
        ('—', 'Strategy pre-filter (score ≥ 4 before Claude call, ~60–70% token reduction)', 'Built and deployed (v5.4)'),
        ('—', 'Execution friction: 1M volume floor, live Alpaca prices, limit orders, 9:45 AM start, 15-min intraday', 'Built and deployed (v5.4)'),
        ('—', 'Reliable cron via cron-job.org; mode detection time windows; eval date filter fix', 'Built and deployed (v5.4)'),
        ('—', 'Target 3% → 2%, stop 1% → 0.67% — maintains 3:1 R:R, break-even at 25% win rate', 'Built and deployed (v5.5)'),
        ('—', 'ML scorer (step 1.76): HistGradientBoosting trained on 2y/429 tickers, AUC 0.78 ± 0.04', 'Built and deployed (v5.5)'),
        ('—', 'Monthly ML retrain workflow (retrain_model.yml) — auto-commits updated pkl to repo', 'Built and deployed (v5.5)'),
        ('V2e', 'Sector rotation scoring — favor sectors showing relative strength', 'Planned'),
        ('V2f', 'Momentum confirmation — 15-minute rule before entry', 'Planned'),
    ]
)

heading('V2a — Market Context Agent (agents/market_context.py)', 2)
body('Runs as Step 0 of premarket, before the scanner. Checks three things:')
bullet('VIX tiered gate: ^VIX from yfinance. VIX >45 → 2 pos; VIX 30–45 → 3 pos; VIX 25–30 → 5 pos; VIX 20–25 → 10 pos; VIX <20 → 15 pos. No hard skip on VIX alone — futures gate handles genuine crash days.')
bullet('Futures gate (tiered — two separate thresholds): ES=F, NQ=F, YM=F average change. '
       'Below -1.5% → SKIP entire day (strong pre-market sell-off, no trades at all). '
       'Between -0.5% and -1.5% → CAUTION, reduce to 8 positions max, futures_bias=BEARISH. '
       'Above +0.5% → BULLISH bias. '
       'These are distinct tiers, not contradictory: caution band still trades (with fewer positions), hard skip fires only on a genuine crash.')
bullet('International markets: Nikkei, FTSE, DAX, Hang Seng, Shanghai. Context only — no hard gate. Majority positive/negative noted in summary.')
body(
    'Returns: decision (GO/CAUTION/SKIP), max_positions (dynamic), summary string for Claude, '
    'vix, futures dict, intl_markets dict, futures_bias (BULLISH/BEARISH/NEUTRAL). '
    'On SKIP, records reason to scan_results and exits — no Claude API calls made.'
)

heading('V2b — News Intelligence Agent (agents/news_intel.py)', 2)
body('Runs as Step 1.5, between scanner and strategy. Two jobs:')
bullet('Earnings blackout: checks yfinance ticker.calendar for each candidate. Removes any with earnings today or tomorrow. Earnings = binary event = unacceptable gap risk for a day-trading system.')
bullet('News context: fetches 3 most recent headlines per remaining candidate from yfinance ticker.news. Passes them to Claude strategy agent as additional context in the prompt.')
body(
    'Returns: filtered_candidates (earnings-removed), blackout_tickers (with reason), '
    'news_context (formatted string for Claude prompt), news_by_ticker (dict for logging). '
    'Handles both DataFrame and dict formats of yfinance calendar (version compatibility).'
)

heading('V2c — Fear & Greed + Economic Calendar (agents/market_context.py)', 2)
body('Added to market_context.py as part of V2.2 and tuned in V2.3:')
bullet('Fear & Greed Index: fetches from alternative.me (free, no API key). Extreme Fear <25 reduces positions when confirmed by VIX>20 or bearish futures. Extreme Greed >80 caps positions at 12. F&G alone is informational — lagging indicator that reads low during bull recoveries.')
bullet('Economic calendar: FOMC, CPI, NFP dates hardcoded for 2025+2026. FOMC day → cap at 8 positions. CPI/NFP day → cap at 10 positions.')
bullet('V2.3 tuning: removed standalone F&G gate — gate cost dropped from -$9,596 to -$2,549 vs baseline (30-day backtest). Both grade B.')

heading('V2d — Sector Correlation Guard (agents/sector_guard.py)', 2)
body('Runs as Step 3.5 between risk agent and portfolio agent — no Claude API call needed. Prevents Claude from over-concentrating in one sector on a given day:')
bullet('Fetches sector from yfinance for each approved trade (10–15 calls max, runs fast)')
bullet('ETFs classified as "ETF" using ETF_UNIVERSE set — no API call needed for those')
bullet('Unknown sector (yfinance rate-limited or unclassifiable) → passes through without cap — safe fallback, never blocks trades due to data gaps')
bullet('Caps at MAX_PER_SECTOR=3 per sector — drops lowest-confidence excess trades')
bullet('Confidence ranking for tiebreak: HIGH > MEDIUM > LOW, then by estimated_profit')
bullet('Dashboard shows sector-blocked trades in Step 2 (Strategy & Risk) as a collapsible expander')
bullet('Configured via MAX_PER_SECTOR in config/settings.py (default: 3)')

heading('V5 — Guardrails (agents/guardrails.py)', 2)
body('Runs as Step 3.75, after sector guard and before portfolio. No Claude API call — deterministic rule checks only. Six checks applied in order:')
add_table(
    ['Check', 'What It Prevents', 'Config'],
    [
        ('Daily loss limit', 'If today\'s realized P&L < -$300, blocks ALL new trades for the rest of the day', 'DAILY_LOSS_LIMIT = -300'),
        ('Action whitelist', 'Rejects any trade action that is not BUY (prevents rogue SELL/SHORT from Claude hallucination)', 'Hard-coded: BUY only'),
        ('Ticker whitelist', 'Rejects any ticker not in the current universe (static or dynamic Supabase list)', 'universe passed from orchestrator'),
        ('Duplicate position guard', 'Blocks a ticker already open or already traded today — prevents doubling up on same name', 'Checks positions table (status=OPEN + closed today)'),
        ('Price sanity', 'Rejects if entry price is >5% from actual market price — catches hallucinated or stale prices from Claude', 'PRICE_SANITY_PCT = 0.05'),
        ('Cumulative capital check', 'Tracks total committed capital across the full approved batch — rejects any trade that would push total deployed capital past available buying_power. Applies to both Alpaca and simulation modes. Prevents margin trading and over-deployment.', 'Calls alpaca_broker.get_buying_power() for Alpaca; uses TOTAL_CAPITAL for simulation'),
    ]
)
body('Additional safeguards in this release:')
bullet('Concurrent run lock: premarket checks for existing scan_results for today at the very start — duplicate GitHub Actions runs exit immediately without opening positions')
bullet('EOD close retry: if Alpaca close_position() fails, retries once after 2 seconds; prints loud warning if still failing (⚠️ manual close required in Alpaca dashboard)')
bullet('sector_blocked and guardrail_blocked now properly persisted to scan_results JSONB in DB after both steps — fixes prior bug where sector_blocked was always stored as [] and never visible on dashboard')

heading('V3a — Dynamic Universe Refresh (agents/universe_refresh.py)', 2)
body('Runs every Monday 8:30 AM ET via a separate GitHub Actions workflow. Replaces static universe with a live-screened list:')
bullet('Fetches S&P 500 (~503 tickers) and Nasdaq 100 (~101 tickers) from Wikipedia using requests with User-Agent header')
bullet('Adds ~41 curated high-momentum tickers: quantum (IONQ, QUBT, QBTS), crypto miners (MARA, RIOT, CLSK), space/eVTOL (ASTS, RKLB, ACHR), AI (SOUN, BBAI), leveraged ETFs (SOXL, SOXS, TQQQ)')
bullet('Screens all ~553 combined tickers: price $5–$500, avg daily volume ≥500K, ATR% ≥2%')
bullet('Saves passing tickers sorted by ATR% descending to Supabase as scan_type="universe_refresh"')
bullet('orchestrator.py load_universe(): reads Supabase if ≤7 days old, falls back to static settings.py')
bullet('First run (May 17 2026): 553 screened → 458 passed. Top movers: SOXS, SOXL, QUBT, OKLO, IREN, LUNR, QBTS, IONQ, HUT')

# ── 10. Backtesting ───────────────────────────────────────────────────────────
heading('10. Backtesting (backtest.py)')
body(
    'A deterministic 30-day historical backtest was built to validate strategy configurations '
    'without spending Claude API credits. It replays historical trading days using real yfinance '
    'price data and a rule-based scorer (no Claude).'
)
bullet('Scores each ticker per historical day using the same technical indicators as the live scanner')
bullet('Selects top N tickers by score as simulated trades')
bullet('Simulates entry at open price, checks daily high/low for 3% target or 1% stop hit')
bullet('Falls back to close price at end of day if neither target nor stop is hit')
bullet('Runs on 30 trading days of actual market data — no synthetic data')
bullet('Usage: python3 backtest.py --days 30 --top 15')
body(
    'Key fix required: yfinance returns multi-level column headers on batch downloads. '
    'Fixed with: df.columns = df.columns.get_level_values(0)'
)

# ── 11. Secrets and Security ──────────────────────────────────────────────────
doc.add_page_break()
heading('11. Secrets and Security')
add_table(
    ['Secret', 'Where Stored', 'What It Does'],
    [
        ('ANTHROPIC_API_KEY', 'GitHub Secrets + Streamlit Secrets + .env', 'Authenticates Claude API calls'),
        ('SUPABASE_URL', 'GitHub Secrets + Streamlit Secrets + .env', 'Points to your Supabase project'),
        ('SUPABASE_KEY', 'GitHub Secrets + Streamlit Secrets + .env', 'Service role key for DB read/write'),
        ('DASHBOARD_PASSWORD', 'Streamlit Secrets + .env', 'Protects the Streamlit dashboard login'),
        ('ALPACA_API_KEY', 'GitHub Secrets + .env', 'Alpaca paper trading API key'),
        ('ALPACA_SECRET_KEY', 'GitHub Secrets + .env', 'Alpaca paper trading secret key'),
    ]
)
body('Security principles applied:')
bullet('.env is gitignored — never pushed to GitHub')
bullet('streamlit_secrets.toml is gitignored — reference file only, never committed')
bullet('GitHub Secrets are encrypted and never visible in logs')
bullet('Streamlit Cloud secrets are encrypted at rest and injected at runtime only')
bullet('SSH key authentication used for git push — no passwords stored')
bullet('Supabase service role key used server-side only — never exposed to browser')

# ── 12. How It Was Built ──────────────────────────────────────────────────────
heading('12. How It Was Built')
body(
    'This system was built entirely using Claude Code — Anthropic\'s AI-powered CLI for software '
    'engineering. The process was conversational: describing what the system should do in plain '
    'English, and Claude Code generating the code, debugging errors, and wiring everything together. '
    'No prior Python expertise required for the builder — Claude handled implementation while the '
    'builder focused on decisions and direction.'
)

heading('Build Steps', 2)
add_table(
    ['Step', 'What', 'Details'],
    [
        ('1', 'Designed the architecture', 'Multi-agent pipeline — scanner → strategy → risk → portfolio'),
        ('2', 'Built the market scanner', 'yfinance + ta library, scoring 60+ tickers initially'),
        ('3', 'Built the strategy agent', 'Claude selects trades with entry/target/stop/reasoning'),
        ('4', 'Built the risk agent', 'Claude validates every trade against hard risk rules'),
        ('5', 'Built the portfolio agent', 'Opens and tracks simulated positions in Supabase'),
        ('6', 'Built the intraday agent', 'Monitors positions every 30 min, closes on target/stop'),
        ('7', 'Built the EOD agent', 'Calculates daily P&L, updates portfolio capital'),
        ('8', 'Wrote the orchestrator', 'Chains all agents, callable with --mode flag'),
        ('9', 'Created the database schema', 'Five Supabase tables with proper indexes and constraints'),
        ('10', 'Set up GitHub Actions', 'Cron workflow running premarket/intraday/EOD on weekdays'),
        ('11', 'Set up Git + GitHub (SSH)', 'Generated SSH key, added to GitHub, switched remote to SSH'),
        ('12', 'Set up Supabase', 'Created project, ran schema.sql, obtained URL and service role key'),
        ('13', 'Set GitHub Secrets', 'Added API keys to repo secrets for runtime injection'),
        ('14', 'End-to-end test', 'Manual premarket run via GitHub Actions — confirmed full pipeline working'),
        ('15', 'Built the Streamlit dashboard', '4-tab app with Today, Positions, Performance, Scan Log'),
        ('16', 'Deployed to Streamlit Cloud', 'Connected repo, set secrets in TOML format, got public URL'),
        ('17', 'Added password protection', 'Login screen added — only authorized users can view data'),
        ('18', 'Built backtest.py', '30-day historical replay with real yfinance data, no Claude API'),
        ('19', 'Expanded universe to 430+ tickers', 'Removed drags based on backtest, added momentum names'),
        ('20', 'Tuned to 3% target, 15 positions, 3:1 R:R', 'Backtest showed $750/day avg at B grade (93% win days)'),
        ('21', 'Built eval.py', 'Weekly grading script: A/B/C/D based on P&L, win rate, reward:risk'),
        ('22', 'Built V2a — Market Context Agent', 'VIX gate + futures signal + international markets'),
        ('23', 'Built V2b — News Intelligence Agent', 'Earnings blackout + news headlines context for Claude'),
        ('24', 'Redesigned dashboard to workflow view', 'Today tab shows all 4 steps: market → scan → plan → positions'),
        ('25', 'Built V2c — Fear & Greed + economic calendar', 'alternative.me F&G Index + FOMC/CPI/NFP hardcoded dates'),
        ('26', 'Tuned V2c to confirming signal (v2.3)', 'F&G standalone gate was too aggressive — now requires VIX or futures confirmation'),
        ('27', 'Built V3a — Universe Refresh Agent', 'Weekly S&P500+Nasdaq100 screener; 458 tickers saved to Supabase on first run'),
        ('28', 'Added company names to dashboard tables', 'config/company_names.py static dict, company column in all ticker tables'),
        ('29', 'Generated documentation and PRD', 'Word documents updated to v3.0'),
        ('30', 'Built V2g — Alpaca broker', 'agents/alpaca_broker.py: bracket orders, position sync, fill price on close'),
        ('31', 'Updated portfolio.py with broker abstraction', 'simulation default, alpaca mode with order submission and sync'),
        ('32', 'Updated trading.yml with Alpaca secrets and --broker flag', 'default broker=alpaca for scheduled runs; manual dispatch has broker input option'),
        ('33', 'Fixed risk agent floating point false rejects', 'stop loss and reward:risk comparison precision; 2dp display in rejection messages'),
        ('34', 'Generated v4.0 documentation and architecture diagrams', 'Word docs updated to v4.0; two PNG architecture diagrams embedded'),
        ('35', 'Built V2d — Sector correlation guard', 'agents/sector_guard.py: fetches sector via yfinance, caps at 3 per sector, Unknown passes through; MAX_PER_SECTOR=3 in settings.py; step 3.5 in orchestrator; dashboard shows sector-blocked trades'),
        ('36', 'Fixed strategy.py JSON extraction', 'Regex on outermost { } or [ ] — handles markdown fence edge cases that produced empty string on json.loads'),
        ('37', 'Built V5 guardrails', 'agents/guardrails.py: 6 safety checks — action whitelist, ticker whitelist, duplicate guard, price sanity (±5%), capital check, daily loss limit (-$300)'),
        ('38', 'Added concurrent run lock', 'orchestrator.py checks for existing premarket scan_results at startup — exits immediately if already ran today, preventing duplicate positions from overlapping GitHub Actions runs'),
        ('39', 'EOD close retry for Alpaca', 'portfolio.py close_all_positions(): retries Alpaca close once after 2s on failure; warns loudly if still failing'),
        ('40', 'Fixed sector_blocked/guardrail_blocked persistence', 'scan_results JSONB updated with actual values after both sector guard and guardrails (previously always stored as [])'),
        ('41', 'Updated docs to v5.0', 'Trading_Agent_Documentation.docx and Trading_Agent_PRD.docx regenerated with V5 guardrails'),
        ('42', 'Redesigned Summary tab', 'Replaced Trades by Sector with three sections: In Flight (live position cards), Today\'s Plan (deduped trade table with status labels), Trade Heatmap (Plotly Treemap — green/red by P&L, sized by position size)'),
        ('43', 'Added automated EOD eval', 'eval.py --write flag saves score/grade/recommendations to Supabase (scan_results, scan_type=eval); EOD GitHub Actions job runs eval.py --days 30 --write automatically after every EOD close'),
        ('44', 'Added Agent Scorecard to Performance tab', 'Reads latest eval from Supabase and displays Score, Grade, Avg Daily P&L, Win Days, Trade Win Rate, Actual R:R, close reason breakdown, best/worst trade, and recommendations — updated every trading day without manual action'),
        ('45', 'Added Alpaca position reconciliation', 'intraday.py _reconcile_with_alpaca() runs first on every 30-min cycle; any Supabase OPEN position not in Alpaca is marked UNFILLED (P&L=0) — catches the rare case where a bracket order was submitted but the entry leg never filled; dashboard and eval exclude UNFILLED alongside CLEANUP'),
        ('46', 'Updated docs to v5.1', 'Trading_Agent_Documentation.docx and Trading_Agent_PRD.docx regenerated with Summary redesign, eval automation, Agent Scorecard, and Alpaca reconciliation'),
        ('47', 'Added no-margin cumulative capital check and daily profit lock-in', 'guardrails.py: committed_capital tracked across full approved batch for both Alpaca and simulation — rejects any trade that would push total deployed capital past buying_power (no margin); intraday.py: LOCK_IN trigger closes all open positions when realized P&L ≥ $716; dashboard shows LOCK_IN as "🎯 Day Locked"; settings.py: DAILY_LOCK_IN_TARGET = 716; docs updated to v5.2'),
        ('48', 'Added GitHub Actions retry logic', 'trading.yml: Run agent step retries up to 3 times with 60s backoff before failing — handles transient yfinance/Alpaca/API failures without manual rerun; concurrent run lock ensures retries after partial success exit cleanly'),
        ('49', 'Trailing stops + confidence-weighted position sizing', 'portfolio.py: high_watermark per position updated each intraday cycle; eff_stop = max(original_stop, peak × 0.99) ratchets as stock rises; Alpaca mode: close_position() fires when trail triggers (bracket auto-cancels); risk.py: _apply_confidence_sizing() maps HIGH→$7K, MEDIUM→$6K, LOW→$5K before validation; strategy.py prompt updated; settings.py: TRAIL_PCT=0.01, POSITION_SIZE_BY_CONFIDENCE; schema.sql: high_watermark column (ALTER TABLE run in Supabase)'),
        ('50', 'Dashboard trailing stop annotations + docs v5.3', 'fmt_stop() helper shows "Trail $X.XX ↑" on In Flight cards (Summary, Today, Positions tabs) when stop has ratcheted; trade_status() shows "🔶 Trail Stop" for STOP closes with positive P&L vs "🔴 Stop Hit" for losses; docs updated to v5.3'),
        ('51', 'Prompt caching', 'strategy.py SYSTEM prompt expanded to 1,243 tokens and marked cache_control: ephemeral; API call uses anthropic-beta: prompt-caching-2024-07-31 header; logs 💾 WRITE (first call) and ⚡ HIT (subsequent calls); ~90% discount on cached tokens'),
        ('52', 'Strategy pre-filter', 'orchestrator step 1.75: filters candidates to technical_score >= STRATEGY_MIN_SCORE (default 4) and positive only before Claude call; combined with caching cuts input tokens ~70–75% per premarket run'),
        ('53', 'Liquidity floor raised 500K → 1M avg volume', 'MIN_AVG_VOLUME in settings.py raised from 500,000 to 1,000,000 — filters thinly-traded tickers, reduces spread and slippage on entries'),
        ('54', 'Real-time Alpaca price refresh', 'get_live_prices() added to alpaca_broker.py using StockLatestQuoteRequest; orchestrator step 1.8 updates candidate current_price with live ask before Claude call; 10% sanity guard rejects anomalous data'),
        ('55', 'Limit order entries', 'submit_bracket_order() changed from MarketOrderRequest to LimitOrderRequest at entry_price × 1.001; avoids paying full spread on entry; portfolio.py updated to pass entry_price'),
        ('56', 'Premarket moved 9:00 → 9:45 AM ET + intraday every 15 min', 'trading.yml cron: premarket 0 13 → 45 13; intraday 0,30 14-19 → */15 14-19; mode detection rewritten to use time windows (not exact minute match) so late runs still classified correctly'),
        ('57', 'cron-job.org external triggers', 'GitHub Actions scheduler fires 5–15 min late — replaced with cron-job.org workflow_dispatch triggers; GitHub schedule kept as backup; concurrent run lock handles duplicates; PAT needs repo+workflow scopes and must be updated in both git credentials AND cron-job.org on regeneration'),
        ('58', 'eval.py date filter fix', 'positions query had no date filter — fetched all history regardless of --days; fixed by building eval_dates set from perf_rows and filtering closed_at date'),
        ('59', 'Futures unavailable on Mondays', 'period="2d" returns only 1 row after weekend (needs 2 for % change calc); fixed to period="5d" in market_context.py; added error logging for insufficient data'),
        ('60', 'Updated docs to v5.4', 'Trading_Agent_Documentation.docx and Trading_Agent_PRD.docx regenerated with all v5.4 friction fixes'),
        ('61', 'Lowered profit target 3% → 2%, stop 1% → 0.67%', 'Maintains 3:1 R:R (break-even at 25% win rate); 2% intraday move is more achievable than 3%; both TARGET_PCT and MAX_LOSS_PER_TRADE in settings.py; strategy.py prompt uses values dynamically'),
        ('62', 'Built ML scorer (train_model.py + scanner/ml_scorer.py)', 'HistGradientBoostingClassifier (sklearn, no libomp dependency); trained on 2y price history for all 429 universe tickers; 13 features; AUC 0.78 ± 0.04 (5-fold TimeSeriesSplit); step 1.76 in orchestrator sorts candidates by P(hit +2%) before Claude call; pkl committed to repo (~2MB)'),
        ('63', 'Monthly ML retrain workflow (.github/workflows/retrain_model.yml)', 'GitHub Actions: triggers 1st of each month at 10 AM UTC; downloads 2y data, retrains model, commits updated xgb_scorer.pkl + feature_columns.json back to main; no manual step required; SUPABASE_URL/KEY via GitHub Secrets'),
        ('64', 'Updated docs and architecture to v5.5', 'Trading_Agent_Documentation.docx, Trading_Agent_PRD.docx, Trading_Agent_Features.docx regenerated; architecture diagrams (high-level + low-level) updated with full 13-step pipeline, ML feedback loop, cron-job.org triggers, interdependencies'),
        ('65', 'Native Alpaca trailing stop (v5.6)', 'USE_NATIVE_TRAILING_STOP feature flag (True); submit_bracket_order() uses StopLossRequest(trail_percent=TRAIL_PCT*100); native_trail_active boolean per position; refresh_positions() skips manual trail check when active; dashboard fmt_stop() shows "Trail 1% ↑ (native)"'),
        ('66', 'Exit mechanism tracking', 'exit_mechanism column added to positions (NATIVE_TRAIL, TARGET, MANUAL_TRAIL, STOP, EOD); all close paths in portfolio.py and alpaca_broker.get_order_fill() populate it; trailing_stop leg type distinguished from fixed stop leg'),
        ('67', 'eval.py — VERDICT summary', 'Plain-language What\'s working / Watch / Action required section at top of eval output; synthesizes all metrics into a 10-second read with specific actionable calls'),
        ('68', 'eval.py — annotated metrics', 'Inline ✅/⚠️/❌ flags with benchmark targets on every number; grade score broken into P&L (40pts) / win-day (30pts) / win-rate (30pts) components; exit reason distribution with healthy mix interpretation'),
        ('69', 'eval.py — integrity checks', 'UNFILLED rate, orphaned open positions, duplicate ticker detection, missing exit_mechanism count, loss-limit day count, lock-in trigger count — all previously invisible'),
        ('70', 'eval.py — Claude quality checks', 'R:R integrity per planned trade (guardrails don\'t enforce R:R — Claude can slip <3:1 trades); position size bounds check; confidence cohort table: HIGH/MEDIUM/LOW win rate, avg P&L, total P&L with signal on whether HIGH outperforms LOW'),
        ('71', 'eval.py — trailing stop validation', 'Native vs manual cohort comparison; delta on win rate and avg P&L; explicit confirmation when NATIVE_TRAIL exits are clean with no double-sells; ⏳ flag when no stop exits yet'),
        ('72', 'Updated docs and architecture to v5.6', 'All generate_*.py scripts updated; Trading_Agent_Changelog.docx created as living session-by-session record of what was built, why, and impact'),
    ]
)

# ── 13. Issues Encountered ────────────────────────────────────────────────────
doc.add_page_break()
heading('13. Issues Encountered and How They Were Fixed')

add_table(
    ['Issue', 'Root Cause', 'Fix'],
    [
        ('git push failed', 'HTTPS auth not configured on Mac', 'Generated SSH key, added to GitHub, switched remote to SSH'),
        ('SSH permission denied after key setup', 'ssh-agent not running in shell session', 'Ran eval "$(ssh-agent -s)" && ssh-add ~/.ssh/id_ed25519'),
        ('Push worked in terminal but not Claude Code', 'SSH agent is session-scoped — not shared across processes', 'User runs git push directly from their terminal'),
        ('Python 3.9 type hint crash (X|Y syntax)', 'GitHub runner uses Python 3.9; X|Y union syntax requires 3.10+', 'Added from __future__ import annotations + used Optional[] from typing'),
        ('Duplicate trade plan DB error (23505)', 'Manual reruns tried to insert same date twice (UNIQUE constraint)', 'Check-before-insert logic in orchestrator before db.insert'),
        ('Dashboard showed no data', 'GitHub runner UTC date ≠ local machine date (one day ahead)', 'Dashboard falls back to most recent plan if today\'s is missing'),
        ('Node.js 20 deprecation warning in Actions', 'Default runner using Node 20, deprecated June 2026', 'Set FORCE_JAVASCRIPT_ACTIONS_TO_NODE24=true in workflow env'),
        ('Streamlit Cloud password not working', 'Secrets not set in Streamlit Cloud — app used empty string as password', 'Added all 4 secrets via Settings → Secrets in TOML format'),
        ('Streamlit secrets not loading via os.getenv()', 'st.secrets not accessible at module import time in settings.py', 'Moved secret injection to top of dashboard/app.py before imports'),
        ('pip/streamlit not found locally', 'Python not in PATH on Mac', 'Used pip3 and python3 -m streamlit instead'),
        ('backtest.py crash — 1D array error', 'yfinance returns multi-level column headers on batch download', 'Fixed with df.columns = df.columns.get_level_values(0)'),
        ('orchestrator undefined variable error', 'db.insert referenced intel["blackout_tickers"] before intel was defined', 'Moved db.insert to after news_intel.run() call'),
        ('yfinance calendar format varies by version', 'ticker.calendar returns DataFrame or dict depending on yfinance version', 'Added isinstance checks in news_intel._get_earnings_date()'),
        ('git commit blocked in Claude Code', 'Claude Code permission mode requires approval for git commit', 'User approves in permission prompt or runs git push from own terminal'),
        ('Risk agent false rejects on exact thresholds', 'Float arithmetic on 2dp stock prices produces 1.000000002 instead of 1.0 — fails strict > comparison', 'round(potential_loss, 4) + tolerance; round(rr, 2) for reward:risk; 2dp display in rejection message'),
        ('GitHub Actions scheduler 5–15 min late', 'GitHub cron scheduler is not guaranteed to fire at exact times', 'cron-job.org external service fires workflow_dispatch at exact UTC times; GitHub schedule kept as backup'),
        ('Mode detection bug on late runs', 'Exact minute match (HOUR=13 AND MIN=00) fails when run fires 10 min late — classified as intraday instead of premarket', 'Rewritten to time window arithmetic: if TIME >= 810 && TIME < 840 → premarket'),
        ('cron-job.org 401 Unauthorized', 'Classic GitHub PAT missing workflow scope (had repo only)', 'Regenerate PAT with both repo + workflow scopes; update in git credential store AND all 3 cron-job.org jobs'),
        ('Streamlit ImportError on TRAIL_PCT', 'Stale Streamlit Cloud deployment had old settings.py without TRAIL_PCT', 'Reboot app from Streamlit Cloud dashboard (⋮ → Reboot app)'),
        ('Futures unavailable on Mondays', 'period="2d" returns 1 row after weekend — not enough for % change calc', 'Changed to period="5d" in market_context.py; added error logging for insufficient data'),
        ('eval.py showing wrong trade count', 'Positions query had no date filter — fetched all historical positions regardless of --days window', 'Build eval_dates set from perf_rows; filter positions by closed_at date'),
        ('XGBoost libomp.dylib error on Mac', 'XGBoost requires OpenMP (libomp.dylib) which is not installed by default on macOS; install fails or import crashes', 'Switched to sklearn HistGradientBoostingClassifier — no C++ deps, pure Python, comparable performance'),
        ('HistGradientBoostingClassifier subsample param error', 'HistGradientBoosting does not expose a subsample parameter (unlike XGBoost)', 'Removed subsample parameter from constructor'),
        ('feature_importances_ AttributeError', 'HistGradientBoosting does not expose feature_importances_ directly — requires permutation_importance from sklearn.inspection', 'Used permutation_importance; moved evaluation display to main() after model save so pkl always persists'),
    ]
)

# ── 14. Daily Operations ──────────────────────────────────────────────────────
heading('14. Daily Operations')
body('Once deployed, the system requires zero daily intervention. What happens automatically:')

add_table(
    ['Time (ET)', 'What Happens', 'Where to See It'],
    [
        ('9:45 AM Mon–Fri', 'Market check → Scan → Pre-filter → Live prices → News → Strategy → Risk → Limit orders', 'GitHub Actions logs + Dashboard Today tab'),
        ('Every 15 min 10AM–3:45PM', 'Intraday position monitoring, close on target/stop, trail stop check', 'Dashboard Positions tab'),
        ('4:30 PM Mon–Fri', 'EOD close + daily P&L + auto eval (saves Agent Scorecard to Supabase)', 'Dashboard Performance tab'),
        ('Anytime', 'Manual trigger via GitHub Actions → Run workflow', 'GitHub Actions'),
        ('Anytime', 'View live workflow dashboard (Summary, Today, Positions, Performance, Scan Log)', 'Streamlit Cloud URL'),
    ]
)

heading('How to Check on Things', 2)
bullet('Dashboard: open your Streamlit Cloud URL → login → Summary tab shows today\'s cockpit; Today tab shows full pipeline workflow')
bullet('Agent Scorecard: Performance tab — updated automatically after every EOD close; shows score, grade, recommendations')
bullet('Logs: github.com/amitgarg73/trading-agent → Actions → click any run for full output')
bullet('Raw data: Supabase → Table Editor → browse any of the 5 tables directly')
bullet('Manual eval: python3 eval.py --days 5 (console only, no --write flag needed unless saving to Supabase manually)')
bullet('Rerun manually: GitHub Actions → Trading Agent → Run workflow → select mode')

# ── 15. What's Next ───────────────────────────────────────────────────────────
heading('15. What\'s Next')
body('The system is live and running at v5.6. Native trailing stop is enabled on paper. 2-week validation gate open (2026-05-18 → 2026-06-01). Next steps:')

add_table(
    ['Phase', 'What', 'Priority'],
    [
        ('2-week paper validation gate', 'Run python3 eval.py --days 14 on June 1; confirm NATIVE_TRAIL exits firing correctly, win rate ≥80%, avg P&L ≥$500/day, no integrity flags', 'P0 — gate: 2026-06-01'),
        ('Post-fix backtest', 'python3 backtest.py --days 30 --top 15; compare to $716/day pre-fix baseline', 'After June 1'),
        ('Real money capital sizing', 'Decide capital amount; rescale POSITION_SIZE_BY_CONFIDENCE and DAILY_LOCK_IN_TARGET proportionally', 'After June 1'),
        ('ML model live validation', 'Compare win rate vs baseline (no scorer) over 30 days paper; validate AUC 0.78 holds live', 'Next sprint'),
        ('Post-earnings momentum agent', 'Scan for stocks 1–3 days post-positive earnings — momentum window before analyst upgrades', 'Backlog'),
        ('Options flow signal', 'Unusual options activity as leading indicator for next-day momentum plays', 'Backlog'),
        ('Insider buying (Form 4)', 'SEC Form 4 filings as conviction signal; weight Claude selection toward insider-bought tickers', 'Backlog'),
    ]
)

# ── Architecture Diagrams ─────────────────────────────────────────────────────
doc.add_page_break()
heading('16. Architecture Diagrams')
body(
    'The following diagrams are auto-generated by generate_architecture.py. '
    'Re-run that script to regenerate if the architecture changes.'
)

_project_dir = "/Users/amitgarg/Claude Projects/trading-agent"
_hl_png = os.path.join(_project_dir, "architecture_high_level.png")
_ll_png = os.path.join(_project_dir, "architecture_low_level.png")

# Generate diagrams if not already present
if not os.path.exists(_hl_png) or not os.path.exists(_ll_png):
    subprocess.run(
        ["python3", os.path.join(_project_dir, "generate_architecture.py")],
        check=True
    )

heading('High-Level Architecture', 2)
body('Shows the trigger sources, orchestrator, full agent pipeline, and external services.')
if os.path.exists(_hl_png):
    doc.add_picture(_hl_png, width=Inches(6.5))
    caption = doc.add_paragraph('Figure 1: High-Level Architecture — triggers, agents, data stores, and external APIs')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

heading('Detailed Pipeline (Low-Level)', 2)
body('Step-by-step premarket pipeline with Alpaca broker abstraction, intraday sync, and EOD flow.')
if os.path.exists(_ll_png):
    doc.add_picture(_ll_png, width=Inches(6.5))
    caption = doc.add_paragraph('Figure 2: Low-Level Pipeline — detailed data flow from market context through EOD with broker modes')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/amitgarg/Claude Projects/trading-agent/Trading_Agent_Documentation.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
