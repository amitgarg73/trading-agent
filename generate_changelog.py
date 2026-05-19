"""
Generates Trading_Agent_Changelog.docx — a living session-by-session record
of what was built, why, the impact, and confidence score at each version.

Update this file every sprint when new features ship. Run alongside the
other generate_*.py scripts when bumping the version.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

NAVY   = RGBColor(0x1A, 0x3A, 0x6A)
ORANGE = RGBColor(0xF4, 0x7B, 0x20)
GRAY   = RGBColor(0x44, 0x44, 0x44)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = NAVY
    return p

def subheading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AI Trading Agent — Session Changelog', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = NAVY

meta = doc.add_paragraph('Amit Garg  ·  May 2026  ·  v5.8  ·  Living document — update each sprint')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(11)
meta.runs[0].font.color.rgb = GRAY

doc.add_paragraph()
body(
    'This document is the authoritative session-by-session record of what was built, '
    'why it was built, the impact on the trading agent, and the real-money confidence score '
    'at each point. Update this file every time a new version ships. '
    'Run python3 generate_changelog.py to regenerate Trading_Agent_Changelog.docx.'
)


# ══════════════════════════════════════════════════════════════════════════════
# v5.8 — 2026-05-18
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('v5.8 — 2026-05-18')
body('Confidence score: 6.5/10 (unchanged)')
body(
    'Theme: Dashboard transparency + observability — Performance tab rebuild, '
    'live scorecard, daily AI summary, and EOD double-run bug fix.'
)

divider()
subheading('1. Daily EOD Summary — Claude Haiku Narrative')
body(
    'What: New agents/daily_summary.py generates a plain-English 3–4 sentence narrative after '
    'every EOD run. Claude Haiku (claude-haiku-4-5-20251001) receives today\'s performance record, '
    'exit reason breakdown, and 5-day rolling context. Prompt explicitly requests: direct interpretation '
    '(not raw number repetition), one actionable observation for tomorrow, plain text only — '
    'no markdown, no headings, no backticks. Stored in scan_results with scan_type=daily_summary '
    '(upserted by date — re-runs overwrite cleanly). Orchestrator EOD function calls generate() '
    'after performance.run(). Dashboard Performance tab shows the latest summary in a light card '
    'at the top of the tab.'
)
body(
    'Why: The Agent Scorecard is comprehensive but requires scrolling through a dense expander. '
    'A 3-sentence plain-English narrative gives an immediate qualitative take — '
    '"momentum worked today, stop-outs on two names dragged the win rate, consider tighter '
    'pre-filter score threshold tomorrow" — without needing to interpret numbers first. '
    'Haiku is used (not Sonnet) because summarization is a lightweight task with a fixed output cap.'
)
body('Impact: Every EOD close now produces a human-readable session brief. Zero additional latency cost.')

divider()
subheading('2. Live Scorecard — eval.py perf_rows Parameter')
body(
    'What: _compute_metrics() in eval.py now accepts an optional perf_rows parameter. '
    'When provided, it skips the Supabase query and computes all metrics from the passed list. '
    'latest_cap uses max(perf_rows, key=lambda r: r["date"]) instead of index 0 to handle any ordering. '
    'Dashboard imports _compute_metrics from eval.py, loads all performance data, applies the '
    'selected date range filter, then passes df.to_dict("records") to compute a live scorecard '
    '— no eval snapshot in Supabase required.'
)
body(
    'Why: The old Scorecard read a pre-computed snapshot saved by eval.py --write. '
    'This meant the scorecard was always fixed to the most recent EOD eval window (30 days) '
    'regardless of what date range the user selected. '
    'With live computation, the scorecard updates dynamically when the user switches date ranges: '
    '"Last 7 days" grade and metrics are computed from the last 7 days of data only.'
)
body('Impact: Scorecard is always consistent with the selected date range. Eval snapshot dependency removed.')

divider()
subheading('3. Dynamic Date Range Selector — Performance Tab')
body(
    'What: Performance tab now has a horizontal radio selector: Last 7 days / Last 30 days / '
    'Last 90 days / All time. Options are shown only when enough data exists: '
    '"Last 30 days" is hidden until total_days >= 30, "Last 90 days" hidden until total_days >= 90. '
    'Filters all three charts (Daily P&L bar, Portfolio Value line, Cumulative P&L line), '
    'the Agent Scorecard, and the daily log table to the selected range.'
)
body(
    'Why: With only 1–5 days of data, showing "Last 30 days" as an option is confusing — '
    'it implies there is 30-day history when there isn\'t. '
    'Hiding options until the data threshold is reached avoids misleading date range labels.'
)
body('Impact: Clean progressive disclosure — date ranges appear as they become meaningful.')

divider()
subheading('4. Dashboard UX — Total Return, Tooltips, Side-by-Side Charts')
body(
    'What: Multiple Performance tab quality-of-life improvements shipped together. '
    '(a) Total Return metric card added (alongside Portfolio Value and Cumulative P&L) — '
    'shows (portfolio_value − starting_capital) / starting_capital as a percentage with help '
    'text showing the absolute dollar gain on $100K since Day 1. '
    '(b) All Agent Scorecard metric cards now have help= tooltip text explaining the formula, '
    'targets, and interpretation: Score (P&L 40pts + win-day 30pts + win-rate 30pts), '
    'Actual R:R, Annualized Return formula, integrity check meaning, VWAP/RS interpretation. '
    '(c) Portfolio Value and Cumulative P&L split from one chart into two side-by-side charts '
    'to avoid scale conflicts. '
    '(d) Verdict section replaced from three-column layout (Action, Watch, Wins) to a single '
    'plain-text narrative paragraph combining all three signals in readable prose. '
    '(e) Grade formula caption added under the Verdict heading: '
    '"Score = P&L vs target (up to 40 pts) + Win day rate (30 pts) + Trade win rate (30 pts)." '
    '(f) VWAP section renamed to "VWAP & Relative Strength Signal Quality", '
    'Thread 1 validation language removed, RS explainer added in collapsible section with '
    'plain-English definition and interpretation thresholds.'
)
body(
    'Why: Multiple independent transparency requests addressed in one pass. '
    'The Total Return metric answered Amit\'s question "how much has the portfolio grown from $100K?" '
    'The tooltips expose the math behind every number. '
    'Side-by-side charts prevent the large portfolio value (e.g. $100,128) from making the daily '
    'P&L ($128) look flat on the same scale. '
    'The Verdict narrative is easier to read than three columns of action items.'
)
body('Impact: Performance tab is now self-explanatory — every metric has context and the chart layout is clear.')

divider()
subheading('5. EOD Double-Run Bug Fix — Starting Capital Calculation')
body(
    'What: performance.py was fetching prev = db.select("daily_performance", order="date", limit=1). '
    'On an EOD re-run, this returned today\'s own record (most recent by date) and used its '
    'ending_capital as starting_capital — inflating the starting capital and producing a wrong portfolio value. '
    'Fix: fetch 2 rows, filter to r["date"] < today, so today\'s own record is excluded on any re-run. '
    'Bad record in Supabase also corrected manually: starting_capital=$100,000, ending_capital=$100,128.29.'
)
body(
    'Why: The bug was invisible until Day 1 was re-run for testing. '
    'The ending_capital shown on the dashboard was $100,385 instead of $100,128.29 — '
    'a $257 discrepancy from the actual $128.29 P&L. '
    'The root cause was the ascending/descending ordering ambiguity: with only one row, '
    'the most recent record and today\'s record are the same object.'
)
body('Impact: Portfolio value now accurate on both first-run and re-run of EOD. Capital carries forward correctly.')

add_table(
    ['Item', 'Status'],
    [
        ('agents/daily_summary.py — Claude Haiku EOD narrative', '✅ Done'),
        ('orchestrator.py — daily_summary.generate() called from EOD', '✅ Done'),
        ('eval.py — perf_rows parameter for live computation', '✅ Done'),
        ('Dashboard — live scorecard from selected date range', '✅ Done'),
        ('Dashboard — dynamic date range selector (hides until data exists)', '✅ Done'),
        ('Dashboard — Total Return metric card', '✅ Done'),
        ('Dashboard — metric help= tooltips on all Scorecard cards', '✅ Done'),
        ('Dashboard — plain-text Verdict narrative paragraph', '✅ Done'),
        ('Dashboard — grade formula caption under Verdict', '✅ Done'),
        ('Dashboard — side-by-side Portfolio Value / Cumulative P&L charts', '✅ Done'),
        ('Dashboard — VWAP section cleanup + RS plain-English explainer', '✅ Done'),
        ('Dashboard — daily EOD summary card at top of Performance tab', '✅ Done'),
        ('agents/performance.py — EOD double-run starting_capital bug fix', '✅ Done'),
        ('VWAP signal quality validation (June 1 gate)', '⏳ ~8 trading days of data accumulating'),
        ('June 1 gate: python3 eval.py --days 14', '⏳ Gate date: 2026-06-01'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# v5.7 — 2026-05-18
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('v5.7 — 2026-05-18')
body('Confidence score: 6.5/10 (was 6/10)')
body(
    'Theme: Stock selection quality (Thread 1 — VWAP + RS vs SPY) + '
    'win-more-on-good-days mechanics (Thread 2 — tiered lock-in) + '
    'validation analytics in eval and dashboard.'
)

divider()
subheading('1. Thread 2 — Tiered Lock-In: Let Winners Ride Past $716')
body(
    'What: Two-tier daily P&L management replacing the old all-or-nothing close. '
    'Tier 1 ($716 realized): stop forcing position closes — let open positions ride '
    'with a tighter 0.5% trail (simulation) or Alpaca native trail (paper). '
    'Tier 2 ($1,000 realized+unrealized): close everything and protect the exceptional day. '
    'New settings: DAILY_LOCK_IN_TARGET=716, DAILY_BONUS_TARGET=1000, LOCK_IN_TRAIL_PCT=0.005. '
    'effective_trail computed before simulation loop — switches to tighter trail when Tier 1 is hit.'
)
body(
    'Why: The old system closed all positions the moment realized P&L hit $716, '
    'leaving open positions that were still moving in the right direction. '
    'On strong momentum days, the best trades are still running at $716 — '
    'a 1% pull from a peak that had not yet reversed. '
    'Closing at $716 on a $1,200 day meant giving away ~$484 of captured value.'
)
body(
    'Impact: On tailwind days (realized ≥ $716), the system now lets winners ride '
    'to the $1,000 ceiling while protecting against reversals with the tighter trail. '
    'eval.py and Agent Scorecard track tailwind days, riders, and extra captured above the floor — '
    'gives a direct before/after comparison in the June 1 eval.'
)

divider()
subheading('2. Thread 1 — VWAP + Relative Strength Signal Enrichment (Step 1.85)')
body(
    'What: New get_intraday_signals(tickers) in alpaca_broker.py uses StockSnapshotRequest '
    'to fetch per-ticker: above_vwap (bool), vwap (price level), today_pct_change (% from open), '
    'rs_vs_spy (stock % move ÷ SPY % move since open). '
    'Step 1.85 in orchestrator (after live prices at 1.8, before Claude call at 2) enriches '
    'each candidate dict and re-sorts the list — above-VWAP candidates first, ties broken by RS descending. '
    'Strategy SYSTEM prompt updated: explains all four fields and teaches Claude to prefer '
    '"above VWAP + RS > 1.5× = ideal momentum setup".'
)
body(
    'Why: The existing scanner uses overnight technical signals — RSI, MACD, Bollinger Bands — '
    'which can be stale by the time the market opens. A stock can be oversold on the 1-day chart '
    'but already selling off at open (below VWAP) with negative relative strength. '
    'VWAP and RS vs SPY provide live confirmation that momentum is intact at entry time. '
    'These are the signals institutional desks use to decide whether a setup is "working" today.'
)
body(
    'Impact: Claude now selects from a pre-sorted list where confirmed momentum setups appear first. '
    'VWAP data persisted to scan_results (vwap_signals dict) for dashboard and eval consumption. '
    'Alpaca mode only — gracefully absent in simulation runs.'
)

divider()
subheading('3. Dashboard — Tailwind Mode + VWAP Badges')
body(
    'Tailwind mode: Blue info banner + progress bar when Tier 1 is active '
    '(realized ≥ $716, positions still open). Green success banner when Tier 2 ceiling is hit. '
    'Position cards show "🚀 tailwind" badge and switch to "Tight Trail" stop label '
    '(0.5% instead of 1%). Applied to Summary In Flight, Today Live Positions.'
)
body(
    'VWAP badges: Each open position card shows an inline ▲ VWAP · RS N× badge (dark blue) '
    'or ▼ VWAP (gray) reflecting the entry-time signal. '
    'Collapsible ℹ️ VWAP & RS signals explained expander defines both metrics with '
    'interpretation thresholds. Candidates table in Today tab gains above_vwap and rs_vs_spy columns. '
    'Applied to Summary, Today, and Positions pages.'
)

divider()
subheading('4. eval.py — Tailwind Analysis + VWAP Signal Quality')
body(
    'Tailwind analysis: _tailwind_analysis() reconstructs the intraday close timeline by '
    'sorting positions by closed_at, walks cumulative P&L to find the exact Tier 1 trigger point, '
    'identifies "riders" (positions that closed after the trigger), and computes extra P&L captured '
    'above the $716 floor. Prints day-by-day breakdown with per-rider exit context. '
    'Available in both console output and Agent Scorecard.'
)
body(
    'VWAP signal quality: _vwap_signal_analysis() loads vwap_signals from each day\'s premarket '
    'scan_result and cross-references every closed position. Produces four cohort tables: '
    'Above VWAP, Below VWAP, RS ≥ 1.5×, RS < 1.5×. '
    'Delta metrics (avg P&L, win rate) measure whether Thread 1 actually adds alpha. '
    'Pass/warn/fail verdict by June 1 gate with ~8 trading days of data.'
)

divider()
subheading('5. Pipeline Funnel Transparency — Today Tab')
body(
    'What: Orchestrator now stores pipeline_counts in scan_results after each filter step: '
    'post_blackout (after earnings filter), post_prefilter + prefilter_dropped (after score ≥4 filter), '
    'ml_scored (candidates with ML probability), live_price_updated (Alpaca price refreshes), '
    'vwap_enriched + above_vwap (VWAP enrichment stats), final_count (what Claude actually saw). '
    'Dashboard Today tab Step 1 replaced with a pipeline funnel expander showing each intermediate '
    'count and a plain-English selection funnel (Universe → Passed scanner → Earnings clear → '
    'Score filter → ML ranked → VWAP enriched → Sent to Claude). '
    'Candidates table now includes ml_score column and is sorted in the exact order Claude received them. '
    '"Passed to Strategy" metric renamed "Sent to Claude" with accurate help text.'
)
body(
    'Why: The Today tab previously showed only three metrics (Candidates Found, Earnings Blocked, '
    'Passed to Strategy) with help text claiming candidates went straight to Claude after earnings filter. '
    'In reality four more filters run: score pre-filter at 1.75, ML ranking at 1.76, live prices at 1.8, '
    'VWAP enrichment at 1.85. A user reading the dashboard had no visibility into why certain stocks '
    'appeared or what order Claude saw them in — making the "why trades were selected" section '
    'structurally incomplete.'
)
body(
    'Impact: Complete audit trail from raw universe to Claude\'s exact input. '
    'Works immediately with historical data for the non-pipeline fields; '
    'pipeline_counts populates from the next live premarket run onward.'
)

divider()
subheading('6. Manual Override — Stop/Restart via GitHub Actions')
body(
    'What: New control.py CLI script and two GitHub Actions workflows for manual override. '
    'control.py --action stop writes a halt_flag row to scan_results (reason, halted_at, '
    'positions_closed). control.py --action restart clears the flag. control.py --action status '
    'checks current state. '
    'stop.yml (workflow_dispatch) accepts a reason string and a close_positions choice (false/true). '
    'When close_positions=true: cancel_all_orders() then close_all_positions() run before the halt flag is set — '
    'every open Alpaca position is market-sold. When false: positions left open with Alpaca native stops protecting them. '
    'restart.yml (workflow_dispatch) runs control.py --action restart. '
    'close_all_positions() added to alpaca_broker.py — fetches all positions via get_all_positions() '
    'and calls close_position() on each; returns list of {ticker, success, fill_price}. '
    'Orchestrator premarket(), intraday(), and eod() each call _is_halted() at the top — '
    'skips the entire run if a halt_flag row exists, prints reason and halted_at.'
)
body(
    'Why: The agent runs unsupervised across market hours via GitHub Actions cron. '
    'If a market circuit breaker fires, strategy degradation is detected, or a technical issue occurs, '
    'there is no way to stop it mid-run without killing the GitHub Actions job manually. '
    'A simple workflow_dispatch "Stop Trading Agent" button in GitHub gives one-click manual control '
    'from any device without needing terminal access or Supabase credentials.'
)
body(
    'Dashboard: Red st.error() halt banner appears on every page after auth — '
    'shows reason, halted_at timestamp, and positions_closed list (or note that Alpaca native stops are active). '
    'Trigger message to restart via GitHub Actions. Applies regardless of which tab is active.'
)

divider()
subheading('7. Halt History Preservation + eval.py Detection + Dashboard Chart Markers')
body(
    'What: restart() in control.py now updates the halt_flag row to scan_type=halt_flag_cleared '
    'with a resumed_at timestamp rather than deleting the record. This preserves the full halt history '
    'in Supabase — both the stop event (halted_at, reason, positions_closed) and the restart event '
    '(resumed_at) are retained in one row per halt cycle. '
    'stop() still hard-deletes any existing halt_flag before writing a fresh one '
    '(guards against stale double-flag state if stop is triggered twice).'
)
body(
    'eval.py: _compute_metrics() now queries both halt_flag and halt_flag_cleared records. '
    'Any record whose halted_at date falls within the eval window is included in halted_days list '
    'with: date, reason, halted_at, resumed_at, and active (bool). '
    '_print_metrics() prints halted_days in [ INTEGRITY CHECKS ] with icon (🛑 if active, ✅ if cleared), '
    'per-halt detail line showing halt time, resume time, and reason. '
    'halted_days is also included in the metrics dict written to Supabase on --write runs.'
)
body(
    'Dashboard Performance tab: Agent Scorecard integrity section shows halted days count '
    'with status icon and per-day captions. '
    'Before the Daily P&L bar chart and Portfolio Value line chart, halt dates are loaded '
    'from scan_results (both halt_flag and halt_flag_cleared). For any halt date that appears '
    'in the chart\'s date range, a red dashed vertical line (add_vline) is drawn on both charts '
    'with a "🛑 Halted" annotation — makes it immediately visible which trading days were paused '
    'and avoids misattributing a flat/missing bar to poor performance.'
)

add_table(
    ['Item', 'Status'],
    [
        ('Thread 2: Tiered lock-in (Tier 1 $716 / Tier 2 $1,000)', '✅ Done'),
        ('Thread 1: VWAP + RS enrichment at step 1.85', '✅ Done'),
        ('Dashboard tailwind mode (banners, badges, tighter trail)', '✅ Done'),
        ('Dashboard VWAP badges + explanation on position cards', '✅ Done'),
        ('eval.py tailwind analysis section', '✅ Done'),
        ('eval.py VWAP signal quality section', '✅ Done'),
        ('Agent Scorecard tailwind + VWAP sections', '✅ Done'),
        ('Pipeline funnel transparency in Today tab', '✅ Done'),
        ('Manual override — stop.yml + restart.yml + control.py', '✅ Done'),
        ('Halt banner on every dashboard page', '✅ Done'),
        ('Halt history preservation (halt_flag_cleared + resumed_at)', '✅ Done'),
        ('eval.py halt detection + INTEGRITY CHECKS section', '✅ Done'),
        ('Dashboard chart vline markers for halted dates', '✅ Done'),
        ('VWAP signal quality validation (June 1 gate)', '⏳ ~8 trading days of data accumulating'),
        ('June 1 gate: python3 eval.py --days 14', '⏳ Gate date: 2026-06-01'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# v5.6 — 2026-05-18
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('v5.6 — 2026-05-18')
body('Confidence score: 6/10 (was 5/10 entering this session)')
body('Theme: P0 real-money blocker resolved + automated validation infrastructure')

divider()
subheading('1. Native Alpaca Trailing Stop')
body(
    'What: StopLossRequest(trail_percent=TRAIL_PCT*100) wired into the bracket order '
    'stop-loss leg. USE_NATIVE_TRAILING_STOP feature flag (enabled as of 2026-05-18). '
    'native_trail_active boolean stored per position. refresh_positions() skips manual '
    'high-watermark check when flag is active. Dashboard shows "Trail 1% ↑ (native)".'
)
body(
    'Why: The manual trailing stop polls every 15 minutes. A stock that peaks and drops '
    '1.5% in 8 minutes isn\'t caught until the next cycle — that\'s real money left on the '
    'table on real capital. Alpaca\'s native trail fires the instant the reversal threshold '
    'is crossed, in real-time, with no polling gap.'
)
body(
    'Impact: Closes the most dangerous gap between paper and real money. Fast reversals — '
    'the hardest scenario for an intraday momentum strategy — are now handled by Alpaca\'s '
    'infrastructure at millisecond resolution. This was the explicit P0 blocker before going live.'
)

divider()
subheading('2. Exit Mechanism Tracking')
body(
    'What: exit_mechanism column added to positions (NATIVE_TRAIL, TARGET, MANUAL_TRAIL, '
    'STOP, EOD). Every close path in portfolio.py and alpaca_broker.get_order_fill() '
    'populates it. get_order_fill() now distinguishes trailing_stop leg type from fixed '
    'stop leg to return NATIVE_TRAIL vs STOP.'
)
body(
    'Why: Without per-position exit tracking there was no way to programmatically confirm '
    'the native trailing stop was firing correctly. Manual Alpaca log inspection doesn\'t scale.'
)
body(
    'Impact: Every closed position now has a complete exit story. '
    'Makes the 2-week validation automatic — no manual log-checking required.'
)

divider()
subheading('3. eval.py — Full Rebuild')
body('Four layers added on top of the existing eval:')
bullet(
    'VERDICT summary — plain-language What\'s working / Watch / Action required at the top. '
    'Synthesizes every metric into a 10-second read with specific action calls.'
)
bullet(
    'Annotated metrics — inline ✅/⚠️/❌ flags with benchmark targets on every number. '
    'Grade score broken into P&L (40pts) / win-day (30pts) / win-rate (30pts) components. '
    'Exit reason distribution with healthy mix interpretation.'
)
bullet(
    'Integrity checks — UNFILLED rate, orphaned open positions, duplicate ticker detection, '
    'missing exit_mechanism count, loss-limit day count, lock-in trigger count. '
    'All previously invisible.'
)
bullet(
    'Claude quality checks — R:R integrity per planned trade (guardrails don\'t enforce R:R), '
    'position size bounds, confidence cohort table: HIGH/MEDIUM/LOW win rate, avg P&L, '
    'total P&L with signal on whether HIGH actually outperforms LOW.'
)
bullet(
    'Trailing stop validation — native vs manual cohort comparison; delta on win rate and avg P&L; '
    'explicit confirmation when NATIVE_TRAIL exits are clean with no double-sells.'
)
body(
    'Why: The old eval answered "are we making money?" The new eval answers "is the system '
    'behaving correctly, is Claude respecting constraints, is the trailing stop working, '
    'and are we ready for real money?" — all in one run.'
)
body(
    'Impact: June 1 python3 eval.py --days 14 produces a complete pass/fail verdict on '
    'every real-money readiness dimension. No spreadsheets, no manual log-checking.'
)

divider()
subheading('4. 2-Week Paper Validation Gate')
body(
    'What: Native trailing stop enabled on paper 2026-05-18. Gate date: 2026-06-01. '
    'PLAN.md updated with pass criteria: win rate ≥80%, avg P&L ≥$500/day, '
    'NATIVE_TRAIL exits confirmed, no integrity flags.'
)
body(
    'Why: Enabling a new exit mechanism on real capital without validation is reckless. '
    'Two weeks of paper trades with automated exit tracking provides statistical confidence.'
)
body('Impact: Converts open-ended "validate before going live" into a concrete gated milestone.')

add_table(
    ['Item', 'Status'],
    [
        ('All 5 friction fixes live', '✅ Done'),
        ('ML scorer live (AUC 0.78)', '✅ Done'),
        ('Native trailing stop built and enabled', '✅ Done'),
        ('Trailing stop validated on paper', '⏳ In progress — gate: 2026-06-01'),
        ('Post-fix backtest rerun', '⬜ After June 1'),
        ('Real capital sizing decided', '⬜ After June 1'),
        ('Confidence assessment ≥7/10', '⬜ After June 1'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# v5.5 — 2026-05-18 (morning)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('v5.5 — 2026-05-18 (morning)')
body('Confidence score: 5/10')
body('Theme: ML scoring layer + target tuning to make daily goal reachable')

divider()
subheading('1. Profit Target Tuning (3% → 2%, Stop 1% → 0.67%)')
body(
    'What: TARGET_PCT lowered from 3% → 2%. MAX_LOSS_PER_TRADE lowered from 1% → 0.67%. '
    '3:1 reward:risk maintained. Break-even win rate drops from 25% to 25% (maintained). '
    'Strategy prompt uses values dynamically.'
)
body(
    'Why: A 3% intraday move is rare — most strong momentum stocks move 1.5–2.5% before '
    'reversing. Lowering to 2% dramatically increases the number of days the target is reachable.'
)
body('Impact: More trades hitting target each day. Wider reachable range without sacrificing R:R.')

divider()
subheading('2. ML Scorer (Step 1.76)')
body(
    'What: HistGradientBoostingClassifier trained on 2y price history for all 429 universe '
    'tickers. 13 features. AUC 0.78 ± 0.04 (5-fold TimeSeriesSplit). Step 1.76 in orchestrator '
    'sorts candidates by P(hit +2%) before Claude call. Top feature: atr_pct (0.165).'
)
body(
    'Why: The scanner produces 50–150 candidates daily. Claude has to pick 10–15. Without '
    'ranking, Claude sees weak and strong setups randomly. The ML scorer surfaces the '
    'highest-probability candidates first, improving selection quality.'
)
body(
    'Impact: Expected win rate improvement by filtering low-probability setups before they '
    'consume position slots. AUC 0.78 is a meaningful predictor above random (0.5).'
)

divider()
subheading('3. Monthly ML Retrain Workflow')
body(
    'What: .github/workflows/retrain_model.yml — fires 1st of each month at 10 AM UTC. '
    'Downloads 2y data, retrains model, commits updated pkl + feature_columns.json to main. '
    'No manual step required.'
)
body('Why: A model trained on 2-year-old data drifts as market regimes change.')
body('Impact: Model stays current automatically. No manual intervention required.')


# ══════════════════════════════════════════════════════════════════════════════
# v5.4 — 2026-05-18 (early morning)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('v5.4 — 2026-05-18 (early morning)')
body('Confidence score: 4/10')
body('Theme: All 5 execution friction fixes + live run stabilization')

divider()
subheading('All 5 Execution Friction Fixes')
add_table(
    ['Fix', 'What', 'Why', 'Impact'],
    [
        ('Real-time price refresh',
         'Alpaca live ask prices replace 15-min stale yfinance before Claude call (step 1.8)',
         'Claude was setting entry/target/stop on prices 15 min old — target often already eaten',
         'Accurate entry prices; Claude decisions based on actual market state'),
        ('Limit order entries',
         'Market orders → limit at entry_price × 1.001; unfilled orders auto-detected',
         'Market orders pay full spread; limit orders reduce slippage on liquid stocks',
         'Better fill quality; UNFILLED positions tracked and excluded from P&L'),
        ('Skip first 15 min',
         'Premarket run 9:00 → 9:45 AM ET; trading.yml cron updated',
         '9:30–9:45 AM has widest spreads; market makers still finding price',
         'Cleaner entries; avoids gap-and-fade setups from opening auction noise'),
        ('15-min intraday checks',
         'Intraday cycle halved: 30 min → 15 min',
         'Trailing stop miss window too wide — fast reversals missed in 30-min gap',
         'Stop miss window cut in half; faster response to reversals'),
        ('cron-job.org scheduler',
         'External HTTP triggers replace GitHub-native cron (fires 5–15 min late)',
         'GitHub cron is not guaranteed to fire on time — premarket ran late every day',
         'Runs fire at exact scheduled times; no more late premarket entries'),
    ]
)

divider()
subheading('Live Run Bug Fixes')
add_table(
    ['Bug', 'Fix'],
    [
        ('STRATEGY_MIN_SCORE=5 collapsed 93→4 candidates', 'Lowered to 4'),
        ('Futures unavailable on Mondays (period="2d")', 'Changed to period="5d"'),
        ('Mode detection failed on late runs', 'Switched to time-window arithmetic'),
        ('eval.py pulling all-time positions regardless of --days', 'Scoped to eval_dates from perf_rows'),
    ]
)


# ══════════════════════════════════════════════════════════════════════════════
# v5.3 and earlier — summary
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('v5.0 – v5.3 — Earlier Versions (Summary)')
body('Detailed commit history in generate_doc.py version history table (commits 1–60).')

add_table(
    ['Version', 'Theme', 'Key Features'],
    [
        ('v5.3', 'Trailing stops + confidence sizing',
         'high_watermark trailing stop (1% trail); HIGH $7K / MEDIUM $6K / LOW $5K sizing; '
         'dashboard trail stop annotations; fmt_stop() helper'),
        ('v5.2', 'Capital safety',
         'No-margin cumulative capital check; daily profit lock-in at $716 (LOCK_IN); '
         'dashboard LOCK_IN display'),
        ('v5.1', 'Dashboard + eval automation',
         'Summary tab redesign (In Flight / Plan / Heatmap); automated EOD eval; '
         'Agent Scorecard; Alpaca position reconciliation (UNFILLED)'),
        ('v5.0', 'Guardrails',
         '6 safety checks: action whitelist, ticker whitelist, duplicate guard, '
         'price sanity (±5%), capital check, daily loss limit (-$300)'),
        ('v4.0', 'Alpaca paper trading',
         'Bracket orders (entry + take-profit + stop-loss); simulation mode retained; '
         'broker= parameter throughout'),
        ('v3.0', 'Supabase + dashboard',
         'PostgreSQL via Supabase; Streamlit Cloud dashboard; 5 tables; '
         'positions, planned_trades, daily_performance'),
        ('v2.0', 'Claude strategy agent',
         'Anthropic API integration; SYSTEM prompt with universe + risk constraints; '
         'structured JSON trade plan output'),
        ('v1.0', 'Scanner foundation',
         'yfinance-based technical scanner; RSI, MACD, Bollinger, volume ratio; '
         'score -10 to +10; configurable universe'),
    ]
)


# ── How to Update This Document ───────────────────────────────────────────────
doc.add_page_break()
heading('How to Update This Document')
body('At the end of every feature sprint or version bump:')
bullet('Add a new top-level section (vX.Y — YYYY-MM-DD) above the previous version')
bullet('Fill in: confidence score, theme, one subsection per major feature shipped')
bullet('Each subsection: What / Why / Impact')
bullet('Add a status table showing pre-real-money checklist items')
bullet('Run python3 generate_changelog.py to regenerate Trading_Agent_Changelog.docx')
bullet('Run all other generate_*.py scripts to keep docs in sync')
bullet('Commit: git add *.docx *.png && git commit -m "docs: update all to vX.Y"')
body('')
body('Documents to regenerate on every version bump:')
add_table(
    ['Script', 'Output'],
    [
        ('python3 generate_changelog.py',   'Trading_Agent_Changelog.docx'),
        ('python3 generate_doc.py',         'Trading_Agent_Documentation.docx'),
        ('python3 generate_prd.py',         'Trading_Agent_PRD.docx'),
        ('python3 generate_features.py',    'Trading_Agent_Features.docx'),
        ('python3 generate_architecture.py','architecture_high_level.png + architecture_low_level.png'),
        ('python3 generate_assessment.py',  'Trading_Agent_Brutal_Assessment.docx'),
        ('python3 generate_vision.py',      'Trading_Agent_Vision.docx'),
    ]
)


# ── Save ──────────────────────────────────────────────────────────────────────
_project_dir = "/Users/amitgarg/Claude Projects/trading-agent"
out = os.path.join(_project_dir, "Trading_Agent_Changelog.docx")
doc.save(out)
print(f"Saved: {out}")
